"""
格式转换器 - 将小说内容转换为多种格式
支持: TXT / HTML / EPUB / PDF / DOCX / Markdown
"""

import json
import re
import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime


class FormatConverter:
    """小说格式转换器"""
    
    SUPPORTED_FORMATS = {
        "txt": {"name": "TXT纯文本", "ext": ".txt", "desc": "纯文本格式，通用性最强"},
        "html": {"name": "HTML网页", "ext": ".html", "desc": "带样式的网页格式"},
        "markdown": {"name": "Markdown", "ext": ".md", "desc": "Markdown标记格式"},
        "epub": {"name": "EPUB电子书", "ext": ".epub", "desc": "电子书阅读器格式"},
        "pdf": {"name": "PDF文档", "ext": ".pdf", "desc": "PDF文档格式"},
        "docx": {"name": "Word文档", "ext": ".docx", "desc": "Microsoft Word格式"},
    }
    
    def __init__(self, novel_dir: Path = None):
        self.novel_dir = novel_dir
    
    def get_formats(self) -> Dict:
        return self.SUPPORTED_FORMATS
    
    def convert(self, content: str, title: str, format_type: str, 
                chapters: List[Dict] = None, metadata: Dict = None,
                images: Dict[str, bytes] = None) -> Optional[str]:
        """转换格式，返回输出文件路径"""
        if format_type not in self.SUPPORTED_FORMATS:
            return None
        
        ext = self.SUPPORTED_FORMATS[format_type]["ext"]
        output_dir = self.novel_dir / "export" if self.novel_dir else Path("export")
        output_dir.mkdir(exist_ok=True)
        
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)
        output_file = output_dir / f"{safe_title}{ext}"
        
        try:
            if format_type == "txt":
                return self._to_txt(content, title, output_file, chapters)
            elif format_type == "html":
                return self._to_html(content, title, output_file, chapters, metadata, images)
            elif format_type == "markdown":
                return self._to_markdown(content, title, output_file, chapters, metadata)
            elif format_type == "epub":
                return self._to_epub(content, title, output_file, chapters, metadata)
            elif format_type == "pdf":
                return self._to_pdf(content, title, output_file, chapters, metadata)
            elif format_type == "docx":
                return self._to_docx(content, title, output_file, chapters, metadata)
        except Exception as e:
            print(f"格式转换失败: {e}")
            return None
        
        return None
    
    def _to_txt(self, content: str, title: str, output_file: Path, 
                chapters: List[Dict] = None) -> str:
        """转换为TXT"""
        lines = [f"{'='*50}", f"  {title}", f"{'='*50}", ""]
        
        if chapters:
            for ch in chapters:
                ch_title = ch.get("title", f"第{ch.get('num', '?')}章")
                ch_content = ch.get("content", "")
                lines.append(f"\n{'='*30}")
                lines.append(f"  {ch_title}")
                lines.append(f"{'='*30}\n")
                lines.append(ch_content)
                lines.append("")
        else:
            lines.append(content)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return str(output_file)
    
    def _to_html(self, content: str, title: str, output_file: Path,
                 chapters: List[Dict] = None, metadata: Dict = None,
                 images: Dict[str, bytes] = None) -> str:
        """转换为HTML"""
        meta = metadata or {}
        author = meta.get("author", "未知作者")
        genre = meta.get("genre", "")
        
        html_parts = [f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: "Microsoft YaHei", "PingFang SC", sans-serif;
            background: #f5f0e8; color: #2c2c2c;
            max-width: 800px; margin: 0 auto; padding: 40px 20px;
            line-height: 1.8;
        }}
        .cover {{ 
            text-align: center; padding: 60px 0; 
            border-bottom: 2px solid #ddd; margin-bottom: 40px;
        }}
        .cover h1 {{ font-size: 2.5em; color: #1a1a2e; margin-bottom: 10px; }}
        .cover .meta {{ color: #666; font-size: 1.1em; }}
        .chapter {{ margin-bottom: 40px; }}
        .chapter h2 {{ 
            font-size: 1.5em; color: #1a1a2e; 
            border-bottom: 1px solid #ddd; padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .chapter p {{ 
            text-indent: 2em; margin-bottom: 1em; 
            font-size: 1.1em;
        }}
        .chapter img {{
            max-width: 100%; height: auto;
            display: block; margin: 20px auto;
            border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        }}
        .footer {{ 
            text-align: center; color: #999; 
            margin-top: 60px; padding-top: 20px;
            border-top: 1px solid #ddd;
        }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>{title}</h1>
        <div class="meta">
            {f'<p>作者: {author}</p>' if author != '未知作者' else ''}
            {f'<p>类型: {genre}</p>' if genre else ''}
            <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        </div>
    </div>
"""]
        
        # 保存图片并生成引用
        img_dir = output_file.parent / "images"
        if images:
            img_dir.mkdir(exist_ok=True)
        
        if chapters:
            for i, ch in enumerate(chapters):
                ch_title = ch.get("title", f"第{ch.get('num', i+1)}章")
                ch_content = ch.get("content", "")
                
                # 处理内容中的图片标记
                ch_html = self._content_to_html(ch_content, images, img_dir)
                
                html_parts.append(f"""
    <div class="chapter" id="chapter-{i+1}">
        <h2>{ch_title}</h2>
        {ch_html}
    </div>
""")
        else:
            content_html = self._content_to_html(content, images, img_dir)
            html_parts.append(f"""
    <div class="chapter">
        {content_html}
    </div>
""")
        
        html_parts.append("""
    <div class="footer">
        <p>由 AI小说创作工坊 生成</p>
    </div>
</body>
</html>""")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(''.join(html_parts))
        
        return str(output_file)
    
    def _content_to_html(self, content: str, images: Dict = None, img_dir: Path = None) -> str:
        """将内容转换为HTML段落"""
        paragraphs = content.split('\n')
        html_parts = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 检查图片标记 ![alt](path)
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', para)
            if img_match:
                alt = img_match.group(1)
                img_path = img_match.group(2)
                
                # 如果有图片数据，保存到文件
                if images and img_path in images and img_dir:
                    img_file = img_dir / f"{len(html_parts)}.png"
                    with open(img_file, 'wb') as f:
                        f.write(images[img_path])
                    img_path = f"images/{img_file.name}"
                
                html_parts.append(f'<p><img src="{img_path}" alt="{alt}"></p>')
            else:
                html_parts.append(f'<p>{para}</p>')
        
        return '\n        '.join(html_parts)
    
    def _to_markdown(self, content: str, title: str, output_file: Path,
                     chapters: List[Dict] = None, metadata: Dict = None) -> str:
        """转换为Markdown"""
        meta = metadata or {}
        
        md_parts = [
            f"# {title}\n",
        ]
        
        if meta.get("author"):
            md_parts.append(f"**作者**: {meta['author']}\n")
        if meta.get("genre"):
            md_parts.append(f"**类型**: {meta['genre']}\n")
        
        md_parts.append(f"*生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n")
        md_parts.append("---\n")
        
        if chapters:
            for i, ch in enumerate(chapters):
                ch_title = ch.get("title", f"第{ch.get('num', i+1)}章")
                ch_content = ch.get("content", "")
                md_parts.append(f"\n## {ch_title}\n")
                md_parts.append(ch_content)
                md_parts.append("")
        else:
            md_parts.append(content)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_parts))
        
        return str(output_file)
    
    def _to_epub(self, content: str, title: str, output_file: Path,
                 chapters: List[Dict] = None, metadata: Dict = None) -> str:
        """转换为EPUB"""
        try:
            from ebooklib import epub
        except ImportError:
            # 如果没有ebooklib，生成HTML版本作为替代
            return self._to_html(content, title, 
                                output_file.with_suffix('.html'), chapters, metadata)
        
        book = epub.EpubBook()
        book.set_identifier(f'novel_{int(datetime.now().timestamp())}')
        book.set_title(title)
        book.set_language('zh')
        
        meta = metadata or {}
        if meta.get("author"):
            book.add_author(meta["author"])
        
        # 添加样式
        style = '''
        body { font-family: "Microsoft YaHei", sans-serif; line-height: 1.8; padding: 20px; }
        h1 { text-align: center; margin-bottom: 30px; }
        h2 { border-bottom: 1px solid #ddd; padding-bottom: 10px; margin: 30px 0 20px; }
        p { text-indent: 2em; margin-bottom: 1em; }
        img { max-width: 100%; height: auto; }
        '''
        css = epub.EpubItem(uid="style", file_name="style/default.css", 
                           media_type="text/css", content=style.encode())
        book.add_item(css)
        
        epub_chapters = []
        toc = []
        
        if chapters:
            for i, ch in enumerate(chapters):
                ch_title = ch.get("title", f"第{ch.get('num', i+1)}章")
                ch_content = ch.get("content", "")
                
                # 转换换行为段落
                paragraphs = ch_content.split('\n')
                html_content = '\n'.join(f'<p>{p.strip()}</p>' for p in paragraphs if p.strip())
                
                epub_ch = epub.EpubHtml(title=ch_title, file_name=f'chapter_{i+1}.xhtml', lang='zh')
                epub_ch.content = f'<html><head><link rel="stylesheet" href="style/default.css"/></head><body><h2>{ch_title}</h2>{html_content}</body></html>'
                epub_ch.add_item(css)
                book.add_item(epub_ch)
                epub_chapters.append(epub_ch)
                toc.append(epub_ch)
        else:
            paragraphs = content.split('\n')
            html_content = '\n'.join(f'<p>{p.strip()}</p>' for p in paragraphs if p.strip())
            
            ch = epub.EpubHtml(title='正文', file_name='content.xhtml', lang='zh')
            ch.content = f'<html><head><link rel="stylesheet" href="style/default.css"/></head><body>{html_content}</body></html>'
            ch.add_item(css)
            book.add_item(ch)
            epub_chapters.append(ch)
            toc.append(ch)
        
        book.toc = toc
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + epub_chapters
        
        epub.write_epub(str(output_file), book)
        return str(output_file)
    
    def _to_pdf(self, content: str, title: str, output_file: Path,
                chapters: List[Dict] = None, metadata: Dict = None) -> str:
        """转换为PDF"""
        try:
            from fpdf import FPDF
        except ImportError:
            return self._to_txt(content, title, output_file.with_suffix('.txt'), chapters)
        
        pdf = FPDF()
        pdf.add_page()
        
        # 尝试添加中文字体
        font_paths = [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf", 
            "C:/Windows/Fonts/simsun.ttc",
        ]
        
        font_added = False
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdf.add_font("Chinese", "", fp, uni=True)
                    pdf.set_font("Chinese", size=12)
                    font_added = True
                    break
                except Exception:
                    continue
        
        if not font_added:
            pdf.set_font("Helvetica", size=12)
        
        # 标题
        pdf.set_font_size(24)
        pdf.cell(0, 20, title, ln=True, align='C')
        pdf.ln(10)
        
        # 元信息
        meta = metadata or {}
        pdf.set_font_size(10)
        if meta.get("author"):
            pdf.cell(0, 8, f"作者: {meta['author']}", ln=True, align='C')
        if meta.get("genre"):
            pdf.cell(0, 8, f"类型: {meta['genre']}", ln=True, align='C')
        pdf.cell(0, 8, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align='C')
        pdf.ln(20)
        
        # 内容
        pdf.set_font_size(12)
        
        if chapters:
            for ch in chapters:
                ch_title = ch.get("title", "")
                ch_content = ch.get("content", "")
                
                pdf.set_font_size(16)
                pdf.cell(0, 12, ch_title, ln=True)
                pdf.ln(5)
                
                pdf.set_font_size(12)
                for para in ch_content.split('\n'):
                    para = para.strip()
                    if para:
                        pdf.multi_cell(0, 8, para)
                        pdf.ln(3)
                
                pdf.ln(10)
        else:
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    pdf.multi_cell(0, 8, para)
                    pdf.ln(3)
        
        pdf.output(str(output_file))
        return str(output_file)
    
    def _to_docx(self, content: str, title: str, output_file: Path,
                 chapters: List[Dict] = None, metadata: Dict = None) -> str:
        """转换为DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return self._to_txt(content, title, output_file.with_suffix('.txt'), chapters)
        
        doc = Document()
        
        # 标题
        doc.add_heading(title, 0)
        
        # 元信息
        meta = metadata or {}
        if meta.get("author"):
            p = doc.add_paragraph(f"作者: {meta['author']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta.get("genre"):
            p = doc.add_paragraph(f"类型: {meta['genre']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        
        if chapters:
            for ch in chapters:
                ch_title = ch.get("title", "")
                ch_content = ch.get("content", "")
                
                doc.add_heading(ch_title, level=1)
                
                for para in ch_content.split('\n'):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
                
                doc.add_paragraph("")
        else:
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
        
        doc.save(str(output_file))
        return str(output_file)


class ImageManager:
    """图片管理器 - 在编辑区插入和管理图片"""
    
    def __init__(self, novel_dir: Path = None):
        self.novel_dir = novel_dir
        self.images_dir = novel_dir / "images" if novel_dir else None
        if self.images_dir:
            self.images_dir.mkdir(exist_ok=True)
        self.images: Dict[str, Path] = {}  # name -> path
    
    def import_image(self, source_path: str, name: str = None) -> Optional[str]:
        """导入图片到项目"""
        source = Path(source_path)
        if not source.exists():
            return None
        
        if not name:
            name = f"img_{int(datetime.now().timestamp())}_{source.name}"
        
        if not self.images_dir:
            return None
        
        dest = self.images_dir / name
        
        import shutil
        shutil.copy2(source, dest)
        
        self.images[name] = dest
        return str(dest)
    
    def get_image_path(self, name: str) -> Optional[str]:
        """获取图片路径"""
        if name in self.images:
            return str(self.images[name])
        return None
    
    def list_images(self) -> List[Dict]:
        """列出所有图片"""
        if not self.images_dir:
            return []
        
        images = []
        for f in self.images_dir.glob("*"):
            if f.suffix.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                images.append({
                    "name": f.name,
                    "path": str(f),
                    "size": f.stat().st_size,
                    "size_str": f"{f.stat().st_size / 1024:.1f}KB",
                })
        
        return images
    
    def insert_image_marker(self, content: str, position: int, image_name: str, 
                           alt_text: str = "插图") -> str:
        """在内容中插入图片标记"""
        marker = f"![{alt_text}]({image_name})"
        return content[:position] + marker + content[position:]
    
    def get_images_data(self) -> Dict[str, bytes]:
        """获取所有图片的二进制数据"""
        data = {}
        if not self.images_dir:
            return data
        
        for f in self.images_dir.glob("*"):
            if f.suffix.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                with open(f, 'rb') as fp:
                    data[f.name] = fp.read()
        
        return data
