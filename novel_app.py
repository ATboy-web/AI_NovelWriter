"""
AI自动写小说系统 - 桌面应用程序
集成AI API、长上下文记忆、智能体机制
"""

import json
import time
import shutil
import threading
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime

# GUI
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext, simpledialog

# HTTP客户端
import httpx

# 结构化日志
from loguru import logger

# 文件格式支持（条件导入）
try:
    from ebooklib import epub
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# 小说工具集
from novel_toolkit import (ElementLibrary, BridgeLibrary, DescriptionLibrary,
                           DialogueEngine, StoryFlowEngine, StyleTransferEngine, AdaptEngine,
                           WebSearchAdaptEngine)
from character_system import CharacterSystem
from format_converter import FormatConverter, ImageManager
from cloud_storage import CloudStorageManager
from app.navigation import NavigationManager


# ==================== 从 app 包导入核心类 ====================
from app import (AppConfig, AIClient, ImageGenerator, SceneDetector,
                 MemoryManager, NoteManager, FullscreenWriter,
                 NovelAgent, ReadingManager, UIStyle)

from app.panels import (ElementsPanelMixin, BridgesPanelMixin,
                        DescriptionsPanelMixin, DialoguePanelMixin,
                        StoryFlowPanelMixin, StylePanelMixin,
                        AdaptPanelMixin, WebSearchPanelMixin,
                        MemoryVizPanelMixin, SummaryMgmtPanelMixin,
                        BatchOpsPanelMixin, ChapterAnalysisPanelMixin)


class NovelWriterApp(
    ElementsPanelMixin, BridgesPanelMixin, DescriptionsPanelMixin,
    DialoguePanelMixin, StoryFlowPanelMixin, StylePanelMixin,
    AdaptPanelMixin, WebSearchPanelMixin, MemoryVizPanelMixin,
    SummaryMgmtPanelMixin, BatchOpsPanelMixin, ChapterAnalysisPanelMixin
):
    """AI自动写小说主应用"""
    
    def __init__(self):
        self.config = AppConfig()
        self.ai_client = AIClient(self.config)
        self.image_gen = ImageGenerator(self.config)
        self.note_manager = NoteManager(config=self.config)
        self.reading_manager = ReadingManager(self.config)
        
        # 工具集
        self.element_lib = ElementLibrary()
        self.bridge_lib = BridgeLibrary()
        self.desc_lib = DescriptionLibrary()
        self.dialogue_engine = None
        self.story_flow_engine = None
        self.style_engine = None
        self.adapt_engine = None
        self.web_search_engine = None
        self.character_system = None
        self.format_converter = None
        self.image_manager = None
        self.cloud_storage = CloudStorageManager()
        
        self.current_novel_dir = None
        self.memory = None
        self.agent = None
        self.outline = []
        self.current_chapter = 0
        self.is_modified = False  # 文档是否已修改
        self._state_lock = threading.Lock()  # 保护共享状态的锁
        
        # 创建GUI
        self.root = tk.Tk()
        self.root.title("AI小说创作工坊 v2.0")
        self.root.geometry("1400x900")
        self.root.minsize(1100, 700)
        
        # 设置图标
        try:
            icon_path = Path(__file__).parent / "icon.ico"
            if icon_path.exists():
                self.root.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"图标加载失败（非致命）: {e}")
        
        # 应用主题
        UIStyle.apply_theme(self.root)
        
        self.nav_manager = NavigationManager(self)
        self.nav_manager.create_menu()
        self._create_widgets()
        self._update_status()
        
        # 绑定快捷键
        self.root.bind('<Control-s>', lambda e: self._save_chapter())
        self.root.bind('<Control-n>', lambda e: self._new_novel())
        self.root.bind('<Control-o>', lambda e: self._open_novel())
        self.root.bind('<F11>', lambda e: self._open_fullscreen_writer())
    
    def _run_async(self, task_func, success_callback=None, error_prefix="操作"):
        """在后台线程执行任务，自动处理异常和UI线程回调
        
        Args:
            task_func: 要执行的函数，返回结果传给 success_callback
            success_callback: 成功回调 function(result)，在UI线程中执行
            error_prefix: 错误消息前缀
        """
        def wrapper():
            try:
                result = task_func()
                if success_callback:
                    self.root.after(0, lambda: success_callback(result))
            except Exception as e:
                self._log(f"{error_prefix}失败: {e}")
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        threading.Thread(target=wrapper, daemon=True).start()
    
    def _create_widgets(self):
        """创建主界面 - 深色主题美化版"""
        C = UIStyle.COLORS
        
        # ===== 顶部标题栏 =====
        header = tk.Frame(self.root, bg=C['accent'], height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        # Logo和标题
        title_frame = tk.Frame(header, bg=C['accent'])
        title_frame.pack(side=tk.LEFT, padx=20, fill=tk.Y)
        
        tk.Label(title_frame, text="AI", font=('Arial', 18, 'bold'), 
                bg=C['accent'], fg='white').pack(side=tk.LEFT)
        tk.Label(title_frame, text=" 小说创作工坊", font=('微软雅黑', 14), 
                bg=C['accent'], fg='white').pack(side=tk.LEFT, padx=(5, 0))
        
        # 顶部按钮
        btn_frame = tk.Frame(header, bg=C['accent'])
        btn_frame.pack(side=tk.RIGHT, padx=20, fill=tk.Y)
        
        for text, cmd in [("新建", self._new_novel), ("打开", self._open_novel), 
                         ("导出", self._export_txt), ("格式转换", self._show_format_converter),
                         ("插入图片", self._insert_image), ("云端同步", self._cloud_sync),
                         ("设置", self._show_settings)]:
            tk.Button(btn_frame, text=text, font=('微软雅黑', 10),
                     bg=C['accent_hover'], fg='white', relief=tk.FLAT,
                     padx=15, pady=5, cursor='hand2',
                     activebackground=C['accent_light'],
                     command=cmd).pack(side=tk.LEFT, padx=3)
        
        # 状态指示
        self.status_indicator = tk.Label(btn_frame, text="未配置", 
                                        font=('微软雅黑', 9), bg=C['accent'], fg='#ffd700')
        self.status_indicator.pack(side=tk.RIGHT, padx=10)
        
        # ===== 主内容区 - 三栏布局 =====
        main_container = tk.Frame(self.root, bg=C['bg_dark'])
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # 快速操作工具栏 (顶部)
        toolbar = tk.Frame(main_container, bg=C['bg_medium'], height=48)
        toolbar.pack(fill=tk.X, padx=0, pady=0)
        toolbar.pack_propagate(False)
        
        # 工具栏左侧 - 标题
        tk.Label(toolbar, text="AI小说创作工坊", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(side=tk.LEFT, padx=15)
        
        # 工具栏中间 - 快速操作按钮
        quick_btn_frame = tk.Frame(toolbar, bg=C['bg_medium'])
        quick_btn_frame.pack(side=tk.LEFT, padx=20)
        
        for text, cmd, color in [
            ("新建小说", lambda: self._new_novel(), C['accent']),
            ("打开小说", lambda: self._open_novel(), C['bg_light']),
            ("续写新章", lambda: self._continue_novel(), C['success']),
        ]:
            btn = tk.Button(quick_btn_frame, text=text, font=('微软雅黑', 9),
                          bg=color, fg='white' if color == C['accent'] else C['text_primary'],
                          relief=tk.FLAT, padx=12, pady=4, cursor='hand2',
                          activebackground=C['accent_hover'],
                          command=cmd)
            btn.pack(side=tk.LEFT, padx=3)
        
        # 工具栏右侧 - 状态信息
        status_frame = tk.Frame(toolbar, bg=C['bg_medium'])
        status_frame.pack(side=tk.RIGHT, padx=15)
        
        self.status_indicator = tk.Label(status_frame, text="未连接AI", font=('微软雅黑', 9),
                                        bg=C['bg_medium'], fg=C['warning'])
        self.status_indicator.pack(side=tk.RIGHT, padx=10)
        
        # 左侧面板 - 可滚动 (280px)
        left_container = tk.Frame(main_container, bg=C['bg_dark'], width=280)
        left_container.pack(side=tk.LEFT, fill=tk.BOTH, padx=0, pady=0)
        left_container.pack_propagate(False)
        
        # 左侧滚动画布
        left_canvas = tk.Canvas(left_container, bg=C['bg_medium'], highlightthickness=0, bd=0, width=280)
        left_scrollbar = tk.Scrollbar(left_container, orient=tk.VERTICAL, command=left_canvas.yview)
        left_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        left_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        left_canvas.configure(yscrollcommand=left_scrollbar.set)
        
        left_panel = tk.Frame(left_canvas, bg=C['bg_medium'])
        left_canvas.create_window((0, 0), window=left_panel, anchor=tk.NW)
        
        def on_left_canvas_configure(event):
            left_canvas.itemconfig(1, width=event.width)
        left_canvas.bind("<Configure>", on_left_canvas_configure)
        
        # 鼠标滚轮绑定
        def on_left_scroll(event):
            left_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        left_canvas.bind("<MouseWheel>", on_left_scroll)
        left_panel.bind("<MouseWheel>", on_left_scroll)
        left_scrollbar.bind("<MouseWheel>", on_left_scroll)
        
        # 左侧 - 小说信息卡片
        info_card = tk.Frame(left_panel, bg=C['bg_card'], padx=15, pady=15)
        info_card.pack(fill=tk.X, padx=10, pady=10)
        
        self.title_var = tk.StringVar(value="未创建小说")
        tk.Label(info_card, textvariable=self.title_var, font=('微软雅黑', 12, 'bold'),
                bg=C['bg_card'], fg=C['text_primary']).pack(anchor=tk.W)
        
        info_grid = tk.Frame(info_card, bg=C['bg_card'])
        info_grid.pack(fill=tk.X, pady=(10, 0))
        
        self.genre_var = tk.StringVar(value="-")
        self.chapter_var = tk.StringVar(value="0/0")
        self.word_count_var = tk.StringVar(value="0")
        
        for i, (label, var) in enumerate([("类型", self.genre_var), ("进度", self.chapter_var), ("字数", self.word_count_var)]):
            tk.Label(info_grid, text=label, font=('微软雅黑', 9), bg=C['bg_card'], fg=C['text_muted']).grid(row=i, column=0, sticky=tk.W, pady=2)
            tk.Label(info_grid, textvariable=var, font=('微软雅黑', 9), bg=C['bg_card'], fg=C['text_secondary']).grid(row=i, column=1, sticky=tk.W, padx=(10, 0), pady=2)
        
        # 左侧 - 模式切换
        mode_frame = tk.Frame(left_panel, bg=C['bg_medium'], padx=10, pady=5)
        mode_frame.pack(fill=tk.X, padx=10)
        
        tk.Label(mode_frame, text="创作模式", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W, pady=(0, 5))
        
        # 自动创作 + 停止 同一行
        btn_row = tk.Frame(mode_frame, bg=C['bg_medium'])
        btn_row.pack(fill=tk.X, pady=2)
        
        self.auto_btn = tk.Button(btn_row, text="自动创作", font=('微软雅黑', 10),
                                 bg=C['accent'], fg='white', relief=tk.FLAT,
                                 padx=10, pady=6, cursor='hand2',
                                 activebackground=C['accent_hover'],
                                 command=self._auto_generate)
        self.auto_btn.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 2))
        
        self.stop_btn = tk.Button(btn_row, text="停止", font=('微软雅黑', 10),
                                 bg=C['error'], fg='white', relief=tk.FLAT,
                                 padx=10, pady=6, cursor='hand2',
                                 activebackground='#dc2626',
                                 command=self._stop_generate)
        self.stop_btn.pack(side=tk.RIGHT)
        
        # AI辅助写作按钮
        self.assist_btn = tk.Button(mode_frame, text="AI辅助写作 (F11)", font=('微软雅黑', 10),
                                   bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT,
                                   padx=10, pady=6, cursor='hand2',
                                   activebackground=C['hover'],
                                   command=self._open_fullscreen_writer)
        self.assist_btn.pack(fill=tk.X, pady=2)
        
        # 续写按钮
        self.continue_btn = tk.Button(mode_frame, text="续写新章", font=('微软雅黑', 10),
                                     bg=C['success'], fg='white', relief=tk.FLAT,
                                     padx=10, pady=6, cursor='hand2',
                                     activebackground='#059669',
                                     command=self._continue_novel)
        self.continue_btn.pack(fill=tk.X, pady=2)
        
        # 左侧 - 智能体步骤
        agent_frame = tk.Frame(left_panel, bg=C['bg_medium'], padx=10, pady=5)
        agent_frame.pack(fill=tk.X, padx=10, pady=(10, 0))
        
        tk.Label(agent_frame, text="创作流程", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W, pady=(0, 5))
        
        # 基础步骤
        basic_label = tk.Label(agent_frame, text="基础设置", font=('微软雅黑', 8),
                              bg=C['bg_medium'], fg=C['text_muted'])
        basic_label.pack(anchor=tk.W, pady=(0, 2))
        
        basic_steps = [
            ("1. 世界观设定", self._gen_settings),
            ("2. 角色创建", self._gen_characters),
            ("3. 故事大纲", self._gen_outline),
        ]
        
        for text, cmd in basic_steps:
            btn = tk.Button(agent_frame, text=text, font=('微软雅黑', 9),
                          bg=C['bg_light'], fg=C['text_secondary'], relief=tk.FLAT,
                          padx=8, pady=4, cursor='hand2', anchor=tk.W,
                          activebackground=C['hover'],
                          command=cmd)
            btn.pack(fill=tk.X, pady=1)
        
        # 创作步骤
        write_label = tk.Label(agent_frame, text="章节创作", font=('微软雅黑', 8),
                              bg=C['bg_medium'], fg=C['text_muted'])
        write_label.pack(anchor=tk.W, pady=(8, 2))
        
        write_steps = [
            ("4. 生成初稿", self._gen_chapter),
            ("5. AI审校", self._review_chapter),
            ("6. 风格优化", self._style_optimize),
            ("7. 保存章节", self._save_chapter),
        ]
        
        for text, cmd in write_steps:
            btn = tk.Button(agent_frame, text=text, font=('微软雅黑', 9),
                          bg=C['bg_light'], fg=C['text_secondary'], relief=tk.FLAT,
                          padx=8, pady=4, cursor='hand2', anchor=tk.W,
                          activebackground=C['hover'],
                          command=cmd)
            btn.pack(fill=tk.X, pady=1)
        
        # 导入与分析
        import_label = tk.Label(agent_frame, text="导入分析", font=('微软雅黑', 8),
                               bg=C['bg_medium'], fg=C['text_muted'])
        import_label.pack(anchor=tk.W, pady=(8, 2))
        
        import_steps = [
            ("导入文档", self._import_document),
            ("AI分析建议", self._ai_analyze_content),
        ]
        
        for text, cmd in import_steps:
            btn = tk.Button(agent_frame, text=text, font=('微软雅黑', 9),
                          bg=C['accent_bg'], fg=C['accent_light'], relief=tk.FLAT,
                          padx=8, pady=4, cursor='hand2', anchor=tk.W,
                          activebackground=C['hover'],
                          command=cmd)
            btn.pack(fill=tk.X, pady=1)
        
        # 左侧 - 大纲列表
        outline_frame = tk.Frame(left_panel, bg=C['bg_medium'], padx=10, pady=5)
        outline_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 10))
        
        # 大纲类型选择
        outline_header = tk.Frame(outline_frame, bg=C['bg_medium'])
        outline_header.pack(fill=tk.X, pady=(0, 5))
        
        tk.Label(outline_header, text="大纲", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(side=tk.LEFT)
        
        # 大纲类型下拉框
        self.outline_type_var = tk.StringVar(value="章节大纲")
        outline_type_combo = ttk.Combobox(outline_header, textvariable=self.outline_type_var,
                                         values=["整体大纲", "章节大纲", "故事大纲"],
                                         state="readonly", width=10, font=('微软雅黑', 8))
        outline_type_combo.pack(side=tk.RIGHT)
        outline_type_combo.bind('<<ComboboxSelected>>', self._on_outline_type_change)
        
        # 大纲操作按钮
        outline_btn_frame = tk.Frame(outline_frame, bg=C['bg_medium'])
        outline_btn_frame.pack(fill=tk.X, pady=2)
        
        tk.Button(outline_btn_frame, text="生成", font=('微软雅黑', 8),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=4,
                 command=self._gen_outline).pack(side=tk.LEFT, padx=1)
        tk.Button(outline_btn_frame, text="添加", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=4,
                 command=self._add_outline_item).pack(side=tk.LEFT, padx=1)
        tk.Button(outline_btn_frame, text="编辑", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=4,
                 command=self._edit_outline_item).pack(side=tk.LEFT, padx=1)
        tk.Button(outline_btn_frame, text="删除", font=('微软雅黑', 8),
                 bg=C['error'], fg='white', relief=tk.FLAT, padx=4,
                 command=self._delete_outline_item).pack(side=tk.RIGHT, padx=1)
        
        # 大纲列表框
        list_frame = tk.Frame(outline_frame, bg=C['bg_dark'])
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.outline_list = tk.Listbox(list_frame, bg=C['bg_dark'], fg=C['text_secondary'],
                                      font=('微软雅黑', 9), selectbackground=C['accent'],
                                      selectforeground='white', relief=tk.FLAT,
                                      highlightthickness=0, borderwidth=0)
        scrollbar = tk.Scrollbar(list_frame, command=self.outline_list.yview)
        self.outline_list.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.outline_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.outline_list.bind('<<ListboxSelect>>', self._on_outline_select)
        
        # 左侧 - 角色卡片 (新增)
        char_cards_frame = tk.Frame(left_panel, bg=C['bg_medium'], padx=10, pady=5)
        char_cards_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        tk.Label(char_cards_frame, text="角色", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W, pady=(0, 5))
        
        # 角色卡片容器
        self.char_cards_container = tk.Frame(char_cards_frame, bg=C['bg_medium'])
        self.char_cards_container.pack(fill=tk.X)
        
        # 角色操作按钮
        char_btn_frame = tk.Frame(char_cards_frame, bg=C['bg_medium'])
        char_btn_frame.pack(fill=tk.X, pady=5)
        tk.Button(char_btn_frame, text="新建", font=('微软雅黑', 8),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=6,
                 command=self._create_character_dialog).pack(side=tk.LEFT, padx=1)
        tk.Button(char_btn_frame, text="AI生成", font=('微软雅黑', 8),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=6,
                 command=self._ai_create_character).pack(side=tk.LEFT, padx=1)
        tk.Button(char_btn_frame, text="传记", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=6,
                 command=self._gen_char_biography).pack(side=tk.LEFT, padx=1)
        
        # ===== 右侧主内容区 =====
        right_panel = tk.Frame(main_container, bg=C['bg_dark'])
        right_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 右侧 - 标签页
        self.notebook = ttk.Notebook(right_panel, style='Dark.TNotebook')
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # === 章节内容页 ===
        chapter_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(chapter_frame, text=" 章节内容 ")
        
        # 章节标题和导航
        title_bar = tk.Frame(chapter_frame, bg=C['bg_dark'])
        title_bar.pack(fill=tk.X, padx=15, pady=(10, 5))
        
        self.chapter_title_var = tk.StringVar(value="选择或生成章节")
        tk.Label(title_bar, textvariable=self.chapter_title_var, 
                font=('微软雅黑', 13, 'bold'), bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        
        # 章节导航按钮
        nav_frame = tk.Frame(title_bar, bg=C['bg_dark'])
        nav_frame.pack(side=tk.RIGHT)
        
        self.prev_chapter_btn = tk.Button(nav_frame, text="◀ 上一章", font=('微软雅黑', 9),
                                         bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT,
                                         padx=8, pady=2, cursor='hand2',
                                         command=self._prev_chapter)
        self.prev_chapter_btn.pack(side=tk.LEFT, padx=2)
        
        self.chapter_select_var = tk.StringVar(value="")
        self.chapter_select = ttk.Combobox(nav_frame, textvariable=self.chapter_select_var,
                                          state="readonly", width=12, font=('微软雅黑', 9))
        self.chapter_select.pack(side=tk.LEFT, padx=2)
        self.chapter_select.bind('<<ComboboxSelected>>', self._on_chapter_select)
        
        self.next_chapter_btn = tk.Button(nav_frame, text="下一章 ▶", font=('微软雅黑', 9),
                                         bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT,
                                         padx=8, pady=2, cursor='hand2',
                                         command=self._next_chapter)
        self.next_chapter_btn.pack(side=tk.LEFT, padx=2)
        
        self.save_chapter_btn = tk.Button(nav_frame, text="💾 保存", font=('微软雅黑', 9),
                                          bg=C['accent'], fg='white', relief=tk.FLAT,
                                          padx=8, pady=2, cursor='hand2',
                                          command=self._save_chapter)
        self.save_chapter_btn.pack(side=tk.LEFT, padx=2)
        
        # 写作区域
        text_frame = tk.Frame(chapter_frame, bg=C['bg_dark'])
        text_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 10))
        
        self.content_text = tk.Text(text_frame, wrap=tk.WORD, font=('微软雅黑', 11),
                                   bg=C['bg_card'], fg=C['text_primary'],
                                   insertbackground=C['accent_light'],
                                   selectbackground=C['accent'],
                                   relief=tk.FLAT, padx=15, pady=15,
                                   spacing1=3, spacing3=3, undo=True)
        
        content_scrollbar = tk.Scrollbar(text_frame, command=self.content_text.yview)
        self.content_text.configure(yscrollcommand=content_scrollbar.set)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.content_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 绑定文本变化事件
        self.content_text.bind('<KeyRelease>', self._on_text_change)
        self.content_text.bind('<Button-3>', self._show_editor_context_menu)
        
        # === 日志页 ===
        log_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(log_frame, text=" 运行日志 & 角色 ")
        
        # 左侧 - 角色面板（在日志左边）
        char_frame = tk.Frame(log_frame, bg=C['bg_medium'], width=240)
        char_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=(10, 5), pady=10)
        char_frame.pack_propagate(False)
        
        tk.Label(char_frame, text="角色面板", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W, pady=(5, 5), padx=5)
        
        # 角色选择下拉框
        char_select_frame = tk.Frame(char_frame, bg=C['bg_medium'])
        char_select_frame.pack(fill=tk.X, padx=5, pady=2)
        self.char_select_var = tk.StringVar(value="无角色")
        self.char_select_combo = ttk.Combobox(char_select_frame, textvariable=self.char_select_var, 
                                              state="readonly", width=16)
        self.char_select_combo.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.char_select_combo.bind('<<ComboboxSelected>>', self._on_char_select)
        
        # 传记按钮
        tk.Button(char_select_frame, text="传", font=('微软雅黑', 8),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=5,
                 command=self._gen_char_biography).pack(side=tk.RIGHT, padx=2)
        
        # 角色信息滚动显示
        char_canvas = tk.Canvas(char_frame, bg=C['bg_medium'], highlightthickness=0, bd=0)
        char_scrollbar = tk.Scrollbar(char_frame, orient=tk.VERTICAL, command=char_canvas.yview)
        char_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        char_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        char_canvas.configure(yscrollcommand=char_scrollbar.set)
        
        self.char_detail_frame = tk.Frame(char_canvas, bg=C['bg_medium'])
        char_canvas.create_window((0, 0), window=self.char_detail_frame, anchor=tk.NW)
        
        def on_char_canvas_configure(event):
            char_canvas.itemconfig(1, width=event.width)
        char_canvas.bind("<Configure>", on_char_canvas_configure)
        for w in [char_canvas, self.char_detail_frame, char_scrollbar]:
            w.bind("<MouseWheel>", lambda e: char_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # 角色操作按钮
        char_btn_frame = tk.Frame(char_frame, bg=C['bg_medium'])
        char_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        tk.Button(char_btn_frame, text="新建", font=('微软雅黑', 8),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=5,
                 command=self._create_character_dialog).pack(side=tk.LEFT, padx=1)
        tk.Button(char_btn_frame, text="AI", font=('微软雅黑', 8),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=5,
                 command=self._ai_create_character).pack(side=tk.LEFT, padx=1)
        tk.Button(char_btn_frame, text="武器", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=5,
                 command=self._equip_weapon).pack(side=tk.LEFT, padx=1)
        tk.Button(char_btn_frame, text="技能", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=5,
                 command=self._learn_skill).pack(side=tk.LEFT, padx=1)
        
        # 右侧 - 日志
        self.log_text = tk.Text(log_frame, wrap=tk.WORD, font=('Consolas', 10),
                               bg=C['bg_card'], fg=C['text_muted'],
                               relief=tk.FLAT, padx=15, pady=15, state=tk.DISABLED)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10), pady=10)
        
        # === 审校结果页 ===
        review_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(review_frame, text=" 审校结果 ")
        
        self.review_text = tk.Text(review_frame, wrap=tk.WORD, font=('微软雅黑', 11),
                                  bg=C['bg_card'], fg=C['text_primary'],
                                  relief=tk.FLAT, padx=15, pady=15)
        self.review_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # === 笔记页 ===
        notes_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(notes_frame, text=" 笔记 ")
        
        # 笔记类型选择
        note_type_frame = tk.Frame(notes_frame, bg=C['bg_dark'])
        note_type_frame.pack(fill=tk.X, padx=15, pady=(15, 5))
        
        self.note_type_var = tk.StringVar(value="project")
        for val, label in [("project", "工程笔记"), ("doc", "文档笔记"), ("sticky", "便笺本")]:
            tk.Radiobutton(note_type_frame, text=label, variable=self.note_type_var, value=val,
                          font=('微软雅黑', 9), bg=C['bg_dark'], fg=C['text_secondary'],
                          selectcolor=C['accent'], activebackground=C['bg_dark'],
                          command=self._refresh_notes).pack(side=tk.LEFT, padx=10)
        
        tk.Button(note_type_frame, text="+ 新建", font=('微软雅黑', 9),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=10,
                 command=self._add_note).pack(side=tk.RIGHT)
        
        # 笔记列表和内容
        note_content_frame = tk.Frame(notes_frame, bg=C['bg_dark'])
        note_content_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 15))
        
        self.notes_list = tk.Listbox(note_content_frame, bg=C['bg_card'], fg=C['text_secondary'],
                                    font=('微软雅黑', 9), selectbackground=C['accent'],
                                    relief=tk.FLAT, highlightthickness=0, width=30)
        self.notes_list.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        self.notes_list.bind('<<ListboxSelect>>', self._on_note_select)
        
        self.note_content = tk.Text(note_content_frame, wrap=tk.WORD, font=('微软雅黑', 11),
                                   bg=C['bg_card'], fg=C['text_primary'],
                                   relief=tk.FLAT, padx=15, pady=15)
        self.note_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 笔记操作按钮
        note_btn_frame = tk.Frame(notes_frame, bg=C['bg_dark'])
        note_btn_frame.pack(fill=tk.X, padx=15, pady=(0, 15))
        
        tk.Button(note_btn_frame, text="保存笔记", font=('微软雅黑', 9),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=10,
                 command=self._save_note).pack(side=tk.LEFT, padx=5)
        tk.Button(note_btn_frame, text="删除笔记", font=('微软雅黑', 9),
                 bg=C['error'], fg='white', relief=tk.FLAT, padx=10,
                 command=self._delete_note).pack(side=tk.LEFT, padx=5)
        tk.Button(note_btn_frame, text="发送到工程", font=('微软雅黑', 9),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=10,
                 command=self._send_sticky_to_project).pack(side=tk.RIGHT, padx=5)
        
        # === 创作工具页 ===
        toolkit_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(toolkit_frame, text=" 创作工具 ")
        
        # 工具选择
        tool_select_frame = tk.Frame(toolkit_frame, bg=C['bg_dark'])
        tool_select_frame.pack(fill=tk.X, padx=15, pady=(15, 5))
        
        self.tool_type_var = tk.StringVar(value="elements")
        tools = [("elements", "元素库"), ("bridges", "桥段库"), ("descriptions", "描写库"),
                 ("dialogue", "对话推演"), ("story_flow", "故事流"), ("style", "风格转换"),
                 ("adapt", "智能改编"), ("websearch", "热点改编"), ("chapters", "章节分析"),
                 ("memory_viz", "记忆可视化"), ("summary_mgmt", "摘要管理"), ("batch_ops", "批量操作")]
        
        for val, label in tools:
            tk.Radiobutton(tool_select_frame, text=label, variable=self.tool_type_var, value=val,
                          font=('微软雅黑', 9), bg=C['bg_dark'], fg=C['text_secondary'],
                          selectcolor=C['accent'], activebackground=C['bg_dark'],
                          command=self._refresh_toolkit).pack(side=tk.LEFT, padx=8)
        
        # 工具内容区
        self.tool_content_frame = tk.Frame(toolkit_frame, bg=C['bg_dark'])
        self.tool_content_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 15))
        
        # 初始化工具集界面
        self._refresh_toolkit()
        
        # 更新左侧面板滚动区域
        left_panel.update_idletasks()
        left_canvas.configure(scrollregion=left_canvas.bbox("all"))
        
        # === 阅读管理器页 ===
        reader_frame = tk.Frame(self.notebook, bg=C['bg_dark'])
        self.notebook.add(reader_frame, text=" 阅读管理器 ")
        
        # 阅读管理器内容
        self._build_reader_ui(reader_frame)
    
    def _build_reader_ui(self, parent):
        """构建阅读管理器界面"""
        C = UIStyle.COLORS
        
        # 工具栏
        toolbar = tk.Frame(parent, bg=C['bg_medium'])
        toolbar.pack(fill=tk.X, padx=15, pady=(15, 5))
        
        tk.Button(toolbar, text="导入书籍", font=('微软雅黑', 9),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=10,
                 command=self._import_book).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="刷新书库", font=('微软雅黑', 9),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=10,
                 command=self._refresh_library).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="搜索", font=('微软雅黑', 9),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=10,
                 command=self._search_in_book).pack(side=tk.LEFT, padx=5)
        
        # 搜索框
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(toolbar, textvariable=self.search_var, 
                               font=('微软雅黑', 9), width=20,
                               bg=C['bg_medium'], fg=C['text_primary'])
        search_entry.pack(side=tk.LEFT, padx=5)
        
        # 书库列表和阅读区域
        main_frame = tk.Frame(parent, bg=C['bg_dark'])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 15))
        
        # 左侧书库列表
        left_frame = tk.Frame(main_frame, bg=C['bg_dark'], width=250)
        left_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        left_frame.pack_propagate(False)
        
        tk.Label(left_frame, text="书库", font=('微软雅黑', 11, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(anchor=tk.W, pady=(0, 10))
        
        # 书库列表框
        list_frame = tk.Frame(left_frame, bg=C['bg_dark'])
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.library_list = tk.Listbox(list_frame, bg=C['bg_card'], fg=C['text_secondary'],
                                      font=('微软雅黑', 9), selectbackground=C['accent'],
                                      relief=tk.FLAT, highlightthickness=0)
        scrollbar = tk.Scrollbar(list_frame, command=self.library_list.yview)
        self.library_list.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.library_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.library_list.bind('<<ListboxSelect>>', self._on_book_select)
        
        # 书签列表
        bookmark_frame = tk.Frame(left_frame, bg=C['bg_dark'])
        bookmark_frame.pack(fill=tk.X, pady=(10, 0))
        
        tk.Label(bookmark_frame, text="书签", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(anchor=tk.W, pady=(0, 5))
        
        self.bookmark_list = tk.Listbox(bookmark_frame, bg=C['bg_card'], fg=C['text_secondary'],
                                       font=('微软雅黑', 8), height=4,
                                       relief=tk.FLAT, highlightthickness=0)
        self.bookmark_list.pack(fill=tk.X)
        self.bookmark_list.bind('<<ListboxSelect>>', self._on_bookmark_select)
        
        # 书签操作按钮
        bookmark_btn_frame = tk.Frame(bookmark_frame, bg=C['bg_dark'])
        bookmark_btn_frame.pack(fill=tk.X, pady=(5, 0))
        tk.Button(bookmark_btn_frame, text="导入书签", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=5,
                 command=self._import_bookmarks).pack(side=tk.LEFT, padx=2)
        tk.Button(bookmark_btn_frame, text="导出书签", font=('微软雅黑', 8),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=5,
                 command=self._export_bookmarks).pack(side=tk.LEFT, padx=2)
        
        # 右侧阅读区域
        right_frame = tk.Frame(main_frame, bg=C['bg_dark'])
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 阅读工具栏
        read_toolbar = tk.Frame(right_frame, bg=C['bg_medium'])
        read_toolbar.pack(fill=tk.X, pady=(0, 10))
        
        # 字体大小
        tk.Label(read_toolbar, text="字体:", font=('微软雅黑', 9),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(side=tk.LEFT, padx=5)
        self.font_size_var = tk.StringVar(value="16")
        font_size_spin = tk.Spinbox(read_toolbar, from_=10, to=36, width=5,
                                   textvariable=self.font_size_var,
                                   command=self._update_reader_font)
        font_size_spin.pack(side=tk.LEFT, padx=5)
        
        # 主题选择
        tk.Label(read_toolbar, text="主题:", font=('微软雅黑', 9),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(side=tk.LEFT, padx=5)
        self.reader_theme_var = tk.StringVar(value="light")
        themes = [("浅色", "light"), ("深色", "dark"), ("护眼", "sepia")]
        for text, value in themes:
            tk.Radiobutton(read_toolbar, text=text, variable=self.reader_theme_var, value=value,
                          font=('微软雅黑', 8), bg=C['bg_medium'], fg=C['text_secondary'],
                          selectcolor=C['accent'], command=self._change_reader_theme).pack(side=tk.LEFT, padx=2)
        
        # 添加书签按钮
        tk.Button(read_toolbar, text="添加书签", font=('微软雅黑', 9),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=10,
                 command=self._add_bookmark).pack(side=tk.RIGHT, padx=5)
        
        # 阅读内容区域
        self.reader_text = tk.Text(right_frame, wrap=tk.WORD, font=('微软雅黑', 16),
                                  bg='#f5f0e8', fg='#2c2c2c',
                                  padx=20, pady=20, spacing1=3, spacing3=3,
                                  relief=tk.FLAT, state=tk.DISABLED)
        
        reader_scrollbar = tk.Scrollbar(right_frame, command=self.reader_text.yview)
        self.reader_text.configure(yscrollcommand=reader_scrollbar.set)
        reader_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.reader_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 当前书籍信息
        self.current_book_path = None
        self.current_book_content = None
        
        # 初始化书库
        self._refresh_library()
        self._refresh_bookmarks()
    
    def _import_book(self):
        """导入书籍"""
        file_path = filedialog.askopenfilename(
            title="选择书籍文件",
            filetypes=[
                ("所有支持格式", "*.txt *.epub *.pdf *.docx *.md"),
                ("TXT文件", "*.txt"),
                ("EPUB文件", "*.epub"),
                ("PDF文件", "*.pdf"),
                ("Word文档", "*.docx"),
                ("Markdown文件", "*.md"),
                ("所有文件", "*.*"),
            ]
        )
        
        if file_path:
            meta = self.reading_manager.import_book(file_path)
            if meta:
                self._log(f"已导入书籍: {meta['title']}")
                self._refresh_library()
                messagebox.showinfo("成功", f"已导入《{meta['title']}》")
            else:
                messagebox.showerror("错误", "导入失败，不支持该文件格式")
    
    def _refresh_library(self):
        """刷新书库列表"""
        self.library_list.delete(0, tk.END)
        books = self.reading_manager.get_library_books()
        for book in books:
            self.library_list.insert(tk.END, f"{book['title']} ({book['format']})")
    
    def _refresh_bookmarks(self):
        """刷新书签列表"""
        self.bookmark_list.delete(0, tk.END)
        bookmarks = self.reading_manager.get_bookmarks()
        for bm in bookmarks:
            self.bookmark_list.insert(tk.END, f"{bm['title']} - {bm.get('position', 0)}%")
    
    def _on_book_select(self, event):
        """书籍选中事件"""
        selection = self.library_list.curselection()
        if not selection:
            return
        
        books = self.reading_manager.get_library_books()
        idx = selection[0]
        if idx < len(books):
            book = books[idx]
            self._load_book(book['file_path'])
    
    def _on_bookmark_select(self, event):
        """书签选中事件"""
        selection = self.bookmark_list.curselection()
        if not selection:
            return
        
        bookmarks = self.reading_manager.get_bookmarks()
        idx = selection[0]
        if idx < len(bookmarks):
            bookmark = bookmarks[idx]
            # 加载对应的书籍并跳转到位置
            if bookmark.get('file_path'):
                self._load_book(bookmark['file_path'], bookmark.get('position', 0))
    
    def _load_book(self, file_path: str, position: int = 0):
        """加载书籍内容"""
        content = self.reading_manager.read_book(file_path)
        if content:
            self.current_book_path = file_path
            self.current_book_content = content
            
            # 显示内容
            self.reader_text.config(state=tk.NORMAL)
            self.reader_text.delete("1.0", tk.END)
            self.reader_text.insert("1.0", content)
            self.reader_text.config(state=tk.DISABLED)
            
            # 跳转到指定位置
            if position > 0:
                # 计算字符位置
                char_pos = int(len(content) * position / 100)
                self.reader_text.see(f"1.0+{char_pos}c")
            
            self._log(f"已加载书籍: {Path(file_path).name}")
        else:
            messagebox.showerror("错误", "无法读取书籍内容")
    
    def _update_reader_font(self):
        """更新阅读字体大小"""
        try:
            size = int(self.font_size_var.get())
            self.reader_text.configure(font=('微软雅黑', size))
        except ValueError:
            pass
    
    def _change_reader_theme(self):
        """改变阅读主题"""
        theme = self.reader_theme_var.get()
        themes = {
            'light': {'bg': '#f5f0e8', 'fg': '#2c2c2c'},
            'dark': {'bg': '#1e1e2e', 'fg': '#f8fafc'},
            'sepia': {'bg': '#f4f0e8', 'fg': '#5c4b37'},
        }
        
        if theme in themes:
            self.reader_text.configure(bg=themes[theme]['bg'], fg=themes[theme]['fg'])
    
    def _add_bookmark(self):
        """添加书签"""
        if not self.current_book_path:
            messagebox.showinfo("提示", "请先打开一本书")
            return
        
        # 获取当前位置（百分比）
        content = self.current_book_content or ""
        # 简单估算：基于滚动位置
        try:
            first_visible = self.reader_text.index("@0,0")
            # 解析行号
            line_num = int(first_visible.split('.')[0])
            total_lines = int(self.reader_text.index(tk.END).split('.')[0])
            position = int(line_num / total_lines * 100) if total_lines > 0 else 0
        except (tk.TclError, ValueError, ZeroDivisionError):
            position = 0
        
        title = f"书签 - {Path(self.current_book_path).stem} - {position}%"
        self.reading_manager.add_bookmark(self.current_book_path, position, title)
        self._refresh_bookmarks()
        self._log(f"已添加书签: {title}")
    
    def _import_bookmarks(self):
        """导入书签"""
        file_path = filedialog.askopenfilename(
            title="导入书签",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")]
        )
        if not file_path:
            return
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                bookmarks = json.load(f)
            
            if isinstance(bookmarks, list):
                for bm in bookmarks:
                    if 'book_path' in bm and 'position' in bm:
                        self.reading_manager.add_bookmark(
                            bm['book_path'],
                            bm['position'],
                            bm.get('title', '导入的书签')
                        )
                self._refresh_bookmarks()
                self._log(f"已导入 {len(bookmarks)} 个书签")
                messagebox.showinfo("成功", f"已导入 {len(bookmarks)} 个书签")
            else:
                messagebox.showerror("错误", "无效的书签文件格式")
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")
    
    def _export_bookmarks(self):
        """导出书签"""
        if not self.reading_manager.bookmarks:
            messagebox.showinfo("提示", "没有可导出的书签")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="导出书签",
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json")]
        )
        if not file_path:
            return
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(self.reading_manager.bookmarks, f, indent=2, ensure_ascii=False)
            self._log(f"已导出 {len(self.reading_manager.bookmarks)} 个书签")
            messagebox.showinfo("成功", f"已导出 {len(self.reading_manager.bookmarks)} 个书签")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")
    
    def _search_in_book(self):
        """在当前书籍中搜索"""
        if not self.current_book_path:
            messagebox.showinfo("提示", "请先打开一本书")
            return
        
        keyword = self.search_var.get().strip()
        if not keyword:
            messagebox.showinfo("提示", "请输入搜索关键词")
            return
        
        results = self.reading_manager.search_in_book(self.current_book_path, keyword)
        if results:
            # 高亮显示搜索结果
            self.reader_text.tag_remove('search_highlight', '1.0', tk.END)
            self.reader_text.tag_configure('search_highlight', background='#ffff00', foreground='#000000')
            
            for result in results[:10]:  # 最多显示10个结果
                line_num = result['line_number']
                # 高亮该行中的关键词
                start = f"{line_num}.0"
                end = f"{line_num}.end"
                self.reader_text.tag_add('search_highlight', start, end)
            
            # 跳转到第一个结果
            self.reader_text.see(f"{results[0]['line_number']}.0")
            self._log(f"找到 {len(results)} 个匹配结果")
        else:
            messagebox.showinfo("搜索", "未找到匹配内容")
    
    def _on_text_change(self, event=None):
        """文本变化事件 - 更新字数统计"""
        content = self.content_text.get("1.0", tk.END).strip()
        self.word_count_var.set(str(len(content)))
        self.is_modified = True
    
    def _show_editor_context_menu(self, event):
        """编辑器右键菜单 - 用选中文字跳转到创作工具"""
        menu = tk.Menu(self.root, tearoff=0, bg=UIStyle.COLORS['bg_card'], 
                      fg=UIStyle.COLORS['text_primary'])
        
        try:
            selected = self.content_text.selection_get()
        except tk.TclError:
            selected = ""
        
        if not selected.strip():
            selected = self.content_text.get("1.0", tk.END).strip()[:500]
            menu.add_command(label="使用全文内容", state=tk.DISABLED)
        else:
            menu.add_command(label=f"已选中 {len(selected)} 字", state=tk.DISABLED)
        
        menu.add_separator()
        menu.add_command(label="事物描写库",
                        command=lambda: self._open_tool_with_text("description", selected))
        menu.add_command(label="角色桥段库",
                        command=lambda: self._open_tool_with_text("bridge", selected))
        menu.add_command(label="情景对话推演",
                        command=lambda: self._open_tool_with_text("dialogue", selected))
        menu.add_command(label="用选中内容仿写",
                        command=lambda: self._style_imitation_with_text(selected))
        menu.add_command(label="生成图片提示词",
                        command=lambda: self._gen_prompt_from_text(selected))
        
        menu.post(event.x_root, event.y_root)
    
    def _open_tool_with_text(self, tool_type: str, text: str):
        """用选中文字跳转到创作工具"""
        self._selected_context_text = text
        tab_mapping = {"description": "创作工具 & 描写", "bridge": "创作工具 & 描写",
                       "dialogue": "情景对话推演"}
        tab_name = tab_mapping.get(tool_type, "创作工具 & 描写")
        
        for i in range(self.notebook.index("end")):
            if self.notebook.tab(i, "text") == tab_name:
                self.notebook.select(i)
                break
        
        self._log(f"已跳转到 {tab_name}，选中 {len(text)} 字作为上下文")
    
    def _style_imitation_with_text(self, text: str):
        """用选中文字进行仿写"""
        self._log("已获取选中内容作为仿写参考")
        self._style_imitation()
    
    def _gen_prompt_from_text(self, text: str):
        """从选中文字生成图片提示词"""
        if not self.current_novel_dir:
            messagebox.showinfo("提示", "请先打开小说")
            return
        scenes = SceneDetector.detect(text)
        if not scenes:
            messagebox.showinfo("提示", "未检测到适合生成图片的场景")
            return
        img_dir = self.current_novel_dir / "scene_prompts"
        img_dir.mkdir(exist_ok=True)
        ts = int(time.time())
        for i, scene in enumerate(scenes):
            f = img_dir / f"manual_{ts}_{i+1}_prompt.txt"
            f.write_text(f"场景: {scene.get('text','')[:200]}\n\n提示词:\n{scene.get('prompt','')}", encoding='utf-8')
        self._log(f"已保存 {len(scenes)} 个手动提示词到 scene_prompts/")
        messagebox.showinfo("成功", f"已生成 {len(scenes)} 个提示词\n保存到 scene_prompts/")
    
    def _log(self, message: str):
        """添加日志（线程安全）"""
        # 使用loguru记录日志
        logger.info(message)
        
        # 同时更新UI日志区域
        def _do_log():
            self.log_text.config(state=tk.NORMAL)
            timestamp = time.strftime("%H:%M:%S")
            self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)
        try:
            self.root.after(0, _do_log)
        except RuntimeError:
            pass
    
    def _update_status(self):
        """更新状态栏"""
        if self.ai_client.is_configured():
            provider = self.config.get("api_provider", "ollama")
            model = self.config.get("model", "")
            img_status = " + 文生图" if self.image_gen.is_configured() else ""
            self.status_indicator.config(text=f"{provider}/{model}{img_status}", fg='#10b981')
        else:
            self.status_indicator.config(text="未配置AI", fg='#ffd700')
    
    def _new_novel(self):
        """新建小说"""
        dialog = tk.Toplevel(self.root)
        dialog.title("新建小说")
        sw, sh = self.root.winfo_screenwidth(), self.root.winfo_screenheight()
        w, h = min(720, sw - 60), int(sh * 0.85)
        x, y = (sw - w) // 2, (sh - h) // 2
        dialog.geometry(f"{w}x{h}+{x}+{y}")
        dialog.resizable(True, True)
        dialog.minsize(550, 600)
        dialog.transient(self.root)
        dialog.grab_set()
        
        C = UIStyle.COLORS
        
        # ===== 顶部固定区域 =====
        top = tk.Frame(dialog, bg=C['bg_dark'])
        top.pack(side=tk.TOP, fill=tk.X, padx=12, pady=(8, 0))
        
        tk.Label(top, text="小说标题:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).grid(row=0, column=0, sticky=tk.W, pady=3)
        title_entry = tk.Entry(top, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'], insertbackground=C['text_primary'], relief=tk.FLAT, width=40)
        title_entry.grid(row=0, column=1, sticky=tk.EW, padx=(5, 0), pady=3)
        
        # 快速模板
        tk.Label(top, text="快速模板:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).grid(row=1, column=0, sticky=tk.W, pady=3)
        template_var = tk.StringVar(value="无")
        template_combo = ttk.Combobox(top, textvariable=template_var, state="readonly", width=35)
        template_combo['values'] = ["无", "穿越异世", "重生归来", "系统流", "都市异能", "修仙问道"]
        template_combo.grid(row=1, column=1, sticky=tk.EW, padx=(5, 0), pady=3)
        
        # 模板变量输入区域
        template_vars_frame = tk.Frame(top, bg=C['bg_dark'])
        template_vars_frame.grid(row=2, column=0, columnspan=2, sticky=tk.EW, pady=3)
        
        template_entries = {}
        
        TEMPLATES = {
            "穿越异世": {
                "genre": "玄幻-异世大陆",
                "tags": ["穿越", "异世界"],
                "template": "一道白光闪过，{name}睁开眼发现自己躺在陌生的{place}，身边站着一位{creature}。",
                "vars": {
                    "name": ("主角名", "李云"),
                    "place": ("地点", "竹林"),
                    "creature": ("人物", "白发老者")
                }
            },
            "重生归来": {
                "genre": "都市-都市异能",
                "tags": ["重生", "复仇"],
                "template": "当{name}再次睁开眼，发现自己回到了{year}年，一切还来得及改变。",
                "vars": {
                    "name": ("主角名", "张伟"),
                    "year": ("年份", "2010")
                }
            },
            "系统流": {
                "genre": "玄幻-东方玄幻",
                "tags": ["系统流", "升级"],
                "template": "叮！恭喜宿主{name}激活{system_name}系统，当前等级：Lv.1。",
                "vars": {
                    "name": ("主角名", "王明"),
                    "system_name": ("系统名", "万界商城")
                }
            },
            "都市异能": {
                "genre": "都市-都市异能",
                "tags": ["异能", "都市"],
                "template": "一场意外让{name}获得了{ability}的能力，从此生活发生了翻天覆地的变化。",
                "vars": {
                    "name": ("主角名", "林风"),
                    "ability": ("能力", "透视")
                }
            },
            "修仙问道": {
                "genre": "仙侠-古典仙侠",
                "tags": ["修仙", "问道"],
                "template": "少年{name}偶得{treasure}，踏上漫漫修仙路。",
                "vars": {
                    "name": ("主角名", "陈轩"),
                    "treasure": ("宝物", "上古功法")
                }
            }
        }
        
        def on_template_change(event=None):
            # 清除旧的模板变量输入框
            for widget in template_vars_frame.winfo_children():
                widget.destroy()
            template_entries.clear()
            
            template_name = template_var.get()
            if template_name == "无":
                return
            
            template = TEMPLATES.get(template_name, {})
            vars_def = template.get("vars", {})
            
            row = 0
            col = 0
            for var_name, (label, default) in vars_def.items():
                lbl = tk.Label(template_vars_frame, text=f"{label}:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 8))
                lbl.grid(row=row, column=col*2, sticky=tk.W, padx=(0, 3), pady=1)
                entry = tk.Entry(template_vars_frame, font=('微软雅黑', 8), bg=C['bg_card'], fg=C['text_primary'], width=12)
                entry.insert(0, default)
                entry.grid(row=row, column=col*2+1, padx=(0, 10), pady=1)
                template_entries[var_name] = entry
                col += 1
                if col >= 3:
                    col = 0
                    row += 1
        
        template_combo.bind('<<ComboboxSelected>>', on_template_change)
        
        tk.Label(top, text="小说频道:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).grid(row=3, column=0, sticky=tk.W, pady=3)
        channel_var = tk.StringVar(value="male")
        ch_frame = tk.Frame(top, bg=C['bg_dark'])
        ch_frame.grid(row=3, column=1, sticky=tk.W, padx=(5, 0), pady=3)
        
        # 男女频类型列表
        MALE_GENRES = [
            "玄幻-东方玄幻", "玄幻-异世大陆", "玄幻-高武世界",
            "仙侠-古典仙侠", "仙侠-现代修真", "仙侠-洪荒封神",
            "都市-都市生活", "都市-都市异能", "都市-青春校园",
            "历史-架空历史", "历史-两宋元明", "历史-三国争霸",
            "科幻-星际文明", "科幻-末世危机", "科幻-时空穿梭",
            "悬疑-灵异恐怖", "悬疑-侦探推理", "悬疑-探险揭秘",
            "游戏-电子竞技", "游戏-虚拟网游", "游戏-游戏异界",
            "军事-抗战烽火", "军事-谍战特工", "军事-战争幻想",
            "武侠-传统武侠", "武侠-国术古武", "武侠-武侠幻想",
            "体育-篮球风云", "体育-足球天下", "体育-综合竞技",
            "轻小说-原生幻想", "轻小说-搞笑吐槽", "轻小说-恋爱日常",
            "二次元-青春日常", "二次元-变身入替", "二次元-同人衍生",
        ]
        FEMALE_GENRES = [
            "古代言情-女尊王朝", "古代言情-宫闱宅斗", "古代言情-穿越奇情",
            "现代言情-豪门总裁", "现代言情-都市婚恋", "现代言情-职场丽人",
            "幻想言情-异世恋歌", "幻想言情-快穿攻略", "幻想言情-魔法幻情",
            "纯爱-古代纯爱", "纯爱-现代纯爱", "纯爱-幻想纯爱",
            "浪漫青春-青春校园", "浪漫青春-疼痛成长", "浪漫青春-纯爱唯美",
            "仙侠奇缘-古典仙缘", "仙侠奇缘-修仙情劫", "仙侠奇缘-洪荒情缘",
            "悬疑灵异-推理侦探", "悬疑灵异-恐怖惊悚", "悬疑灵异-灵异鬼怪",
            "游戏竞技-电子竞技", "游戏竞技-全息网游", "游戏竞技-电竞爱情",
            "短篇-短篇言情", "短篇-微小说", "短篇-轻小说",
        ]
        
        tk.Label(top, text="小说类型:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).grid(row=4, column=0, sticky=tk.W, pady=3)
        genre_var = tk.StringVar(value=MALE_GENRES[0])
        genre_combo = ttk.Combobox(top, textvariable=genre_var, values=MALE_GENRES, state="readonly", width=35)
        genre_combo.grid(row=4, column=1, sticky=tk.EW, padx=(5, 0), pady=3)
        top.columnconfigure(1, weight=1)
        
        # 男生标签（8类 80+标签）
        MALE_TAGS = {
            "角色设定": ["废材崛起", "扮猪吃虎", "杀伐果断", "智商在线", "低调男主", "独行侠", "狠人大帝", "稳健型", "腹黑型", "热血少年", "冷面高手", "逍遥自在", "护短", "不圣母", "有底线", "重生者"],
            "情节元素": ["系统流", "穿越大军", "重生复仇", "无敌流", "升级流", "种田流", "争霸流", "诸天流", "无限流", "签到流", "数据化", "聊天群", "直播流", "召唤流", "转生流", "模拟器"],
            "世界观": ["异界大陆", "王朝争霸", "宗门林立", "末世废土", "星空宇宙", "灵气复苏", "赛博朋克", "求生冒险", "东方神话", "洪荒封神", "修真文明", "巫师世界"],
            "爽点标签": ["越级挑战", "越阶杀敌", "装逼打脸", "逆天改命", "一人成军", "万古不朽", "超神之路", "武道巅峰", "碾压全场", "秀翻天", "骚操作", "神级操作"],
            "成长路线": ["废柴逆袭", "天才陨落再起", "散修崛起", "赘婿翻身", "上门女婿", "退婚打脸", "回归都市", "隐世归来", "退役兵王", "回归豪门"],
            "战斗风格": ["肉身成圣", "剑道独尊", "拳拳到肉", "法术流", "武技流", "炼丹大师", "阵法宗师", "器道大师", "驭兽师", "暗杀流", "群战之王"],
            "感情线": ["单女主", "多女主", "后宫流", "无女主", "暧昧流", "青梅竹马", "天降系", "傲娇女主", "御姐型", "萝莉型", "病娇女主"],
            "特殊设定": ["万界穿梭", "时间回溯", "读心术", "透视眼", "隐身术", "空间戒指", "金手指", "老爷爷", "神级血脉", "远古传承", "神器认主", "神兽伙伴"],
        }
        FEMALE_TAGS = {
            "角色设定": ["甜宠女主", "女强逆袭", "马甲大佬", "团宠担当", "万人迷", "病娇偏执", "霸总老公", "白月光", "替身前妻", "软萌娇妻", "女王御姐", "萌宝来袭", "戏精女主", "佛系女主", "毒舌女主", "学霸女主"],
            "情节元素": ["先婚后爱", "追妻火葬场", "带球跑", "契约婚姻", "养成系", "宅斗宫斗", "真假千金", "失忆重逢", "假戏真做", "隐婚密爱", "替身文学", "重生虐渣", "闪婚闪离", "破镜重圆", "日久生情", "强取豪夺"],
            "气氛风格": ["虐恋情深", "欢喜冤家", "温馨治愈", "爆笑甜宠", "暗恋成真", "虐渣打脸", "逆袭爽文", "甜到齁", "虐到哭", "轻松欢脱", "高甜无虐", "玻璃渣里找糖"],
            "甜宠类型": ["一见钟情", "日久生情", "暗恋成真", "宠妻狂魔", "双向奔赴", "青梅竹马", "师生恋", "姐弟恋", "大叔宠", "萌宝助攻", "豪门恩怨", "总裁文"],
            "身份设定": ["豪门千金", "落魄千金", "穿越女主", "重生女主", "系统女主", "异能女主", "修仙女主", "古代女主", "现代女主", "末世女主", "娱乐圈女主", "军嫂文"],
            "男主人设": ["霸道总裁", "冷面军少", "腹黑王爷", "温柔竹马", "傲娇少爷", "冰山校草", "禁欲系", "病娇男主", "忠犬男主", "渣男回头", "高冷学长", "阳光少年"],
            "感情模式": ["甜宠", "先虐后甜", "先甜后虐", "甜虐交织", "高甜", "暗恋", "明恋", "单箭头", "双箭头", "三角恋", "四角恋", "骨科"],
            "特殊元素": ["萌宝", "双胞胎", "龙凤胎", "穿越", "重生", "系统", "空间", "异能", "修仙", "娱乐圈", "豪门", "校园"],
        }
        
        # ===== 中间可滚动标签区域 =====
        tag_outer = tk.LabelFrame(dialog, text=" 附加标签（可多选，滚动查看） ", padx=5, pady=5,
                                  bg=C['bg_dark'], fg=C['accent_light'], font=('微软雅黑', 9))
        tag_outer.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        
        tag_canvas = tk.Canvas(tag_outer, bg=C['bg_dark'], highlightthickness=0, bd=0, height=200)
        tag_scrollbar = tk.Scrollbar(tag_outer, orient=tk.VERTICAL, command=tag_canvas.yview)
        tag_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tag_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tag_canvas.configure(yscrollcommand=tag_scrollbar.set)
        
        tags_container = tk.Frame(tag_canvas, bg=C['bg_dark'])
        tag_canvas_window = tag_canvas.create_window((0, 0), window=tags_container, anchor=tk.NW)
        
        # 鼠标滚轮滚动标签
        def on_tag_mousewheel(event):
            tag_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        tag_canvas.bind("<MouseWheel>", on_tag_mousewheel)
        tags_container.bind("<MouseWheel>", on_tag_mousewheel)
        tag_outer.bind("<MouseWheel>", on_tag_mousewheel)
        tag_scrollbar.bind("<MouseWheel>", on_tag_mousewheel)
        
        # 使tags_container宽度跟随canvas
        def on_canvas_configure(event):
            tag_canvas.itemconfig(tag_canvas_window, width=event.width)
        tag_canvas.bind("<Configure>", on_canvas_configure)
        
        self.tag_vars = {}
        
        def update_tags(channel):
            """更新标签显示"""
            genre_combo['values'] = MALE_GENRES if channel == "male" else FEMALE_GENRES
            genre_var.set(MALE_GENRES[0] if channel == "male" else FEMALE_GENRES[0])
            for w in tags_container.winfo_children():
                w.destroy()
            self.tag_vars.clear()
            tags = MALE_TAGS if channel == "male" else FEMALE_TAGS
            for cat_name, cat_tags in tags.items():
                cat_label = tk.Label(tags_container, text=cat_name, font=("微软雅黑", 9, "bold"),
                                    bg=C['bg_dark'], fg=C['accent_light'], anchor=tk.W)
                cat_label.pack(fill=tk.X, pady=(6, 1), padx=5)
                tag_line = tk.Frame(tags_container, bg=C['bg_dark'])
                tag_line.pack(fill=tk.X, padx=5)
                for tag in cat_tags:
                    var = tk.BooleanVar(value=False)
                    self.tag_vars[tag] = var
                    cb = tk.Checkbutton(tag_line, text=tag, variable=var, font=("微软雅黑", 8),
                                       bg=C['bg_dark'], fg=C['text_secondary'],
                                       selectcolor=C['bg_card'],
                                       activebackground=C['bg_dark'],
                                       activeforeground=C['accent_light'],
                                       relief=tk.FLAT, padx=1, pady=0)
                    cb.pack(side=tk.LEFT, padx=2, pady=1)
            tags_container.update_idletasks()
            tag_canvas.configure(scrollregion=tag_canvas.bbox("all"))
            tag_canvas.yview_moveto(0)
        
        # 频道选择
        for text, val in [("男生频道", "male"), ("女生频道", "female")]:
            rb = tk.Radiobutton(ch_frame, text=text, variable=channel_var, value=val,
                               command=lambda v=val: update_tags(v),
                               bg=C['bg_dark'], fg=C['text_primary'],
                               selectcolor=C['bg_card'],
                               activebackground=C['bg_dark'],
                               activeforeground=C['accent_light'],
                               font=('微软雅黑', 9))
            rb.pack(side=tk.LEFT, padx=8)
        
        # 初始化标签
        update_tags("male")
        
        # 自定义标签输入
        custom_frame = tk.Frame(tag_outer, bg=C['bg_dark'])
        custom_frame.pack(fill=tk.X, padx=5, pady=(3, 0))
        tk.Label(custom_frame, text="自定义:", bg=C['bg_dark'], fg=C['text_secondary'], font=('微软雅黑', 8)).pack(side=tk.LEFT)
        custom_tag_entry = tk.Entry(custom_frame, font=('微软雅黑', 9), bg=C['bg_card'], fg=C['text_primary'],
                                    insertbackground=C['text_primary'], relief=tk.FLAT, width=20)
        custom_tag_entry.pack(side=tk.LEFT, padx=3)
        
        def add_custom_tag():
            tag = custom_tag_entry.get().strip()
            if not tag:
                return
            if tag in self.tag_vars:
                messagebox.showinfo("提示", f"标签 '{tag}' 已存在")
                return
            var = tk.BooleanVar(value=True)
            self.tag_vars[tag] = var
            # 添加到最后一行
            last_line = None
            for w in tags_container.winfo_children():
                if isinstance(w, tk.Frame):
                    last_line = w
            if last_line is None:
                last_line = tk.Frame(tags_container, bg=C['bg_dark'])
                last_line.pack(fill=tk.X, padx=5)
            cb = tk.Checkbutton(last_line, text=tag, variable=var, font=("微软雅黑", 8),
                               bg=C['bg_card'], fg=C['accent_light'],
                               selectcolor=C['bg_dark'],
                               activebackground=C['bg_card'],
                               activeforeground=C['accent_light'],
                               relief=tk.RAISED, padx=3, pady=1)
            cb.pack(side=tk.LEFT, padx=2, pady=1)
            custom_tag_entry.delete(0, tk.END)
            self._log(f"添加自定义标签: {tag}")
        
        tk.Button(custom_frame, text="添加", command=add_custom_tag, font=('微软雅黑', 8),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=5).pack(side=tk.LEFT, padx=3)
        custom_tag_entry.bind("<Return>", lambda e: add_custom_tag())
        
        # ===== 底部固定区域 =====
        bottom = tk.Frame(dialog, bg=C['bg_dark'])
        bottom.pack(fill=tk.X, padx=15, pady=(0, 5), side=tk.BOTTOM)
        
        # 章节数 + 每章字数 + 创建按钮（同一行）
        action_row = tk.Frame(bottom, bg=C['bg_dark'])
        action_row.pack(fill=tk.X, pady=3)
        tk.Label(action_row, text="章节数:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).pack(side=tk.LEFT)
        chapters_var = tk.StringVar(value="20")
        tk.Spinbox(action_row, from_=1, to=500, textvariable=chapters_var, width=6,
                  font=('微软雅黑', 9), bg=C['bg_card'], fg=C['text_primary']).pack(side=tk.LEFT, padx=5)
        
        tk.Label(action_row, text="每章字数:", bg=C['bg_dark'], fg=C['text_primary'], font=('微软雅黑', 9)).pack(side=tk.LEFT, padx=(15, 0))
        word_count_var = tk.StringVar(value="3000")
        word_count_combo = ttk.Combobox(action_row, textvariable=word_count_var, width=8, font=('微软雅黑', 9))
        word_count_combo['values'] = ["1000", "2000", "3000", "5000", "8000", "10000"]
        word_count_combo.pack(side=tk.LEFT, padx=5)
        
        def confirm():
            title = title_entry.get().strip()
            if not title:
                messagebox.showwarning("提示", "请输入小说标题")
                return
            
            genre_full = genre_var.get().split("-")
            genre = genre_full[0] if len(genre_full) > 0 else ""
            sub_genre = genre_full[1] if len(genre_full) > 1 else ""
            chapters = int(chapters_var.get())
            concept = ""
            
            # 收集选中的标签
            selected_tags = [tag for tag, var in self.tag_vars.items() if var.get()]
            
            # 处理模板
            template_name = template_var.get()
            if template_name != "无" and template_name in TEMPLATES:
                template_data = TEMPLATES[template_name]
                # 收集模板变量值
                vars_values = {}
                for var_name, entry in template_entries.items():
                    vars_values[var_name] = entry.get().strip()
                
                # 生成模板内容
                template_content = template_data["template"]
                for var_name, value in vars_values.items():
                    template_content = template_content.replace(f"{{{var_name}}}", value)
                
                # 使用模板的类型和标签
                if not selected_tags:
                    selected_tags = template_data.get("tags", [])
                
                # 保存模板内容到概念
                if not concept:
                    concept = template_content
            
            # 创建小说目录
            safe_name = "".join(c for c in title if c.isalnum() or c in "_ -")[:30]
            novel_dir = self.config.novels_dir / f"{safe_name}_{int(time.time())}"
            novel_dir.mkdir(exist_ok=True)
            
            # 保存小说元数据
            meta = {
                "title": title,
                "genre": genre,
                "sub_genre": sub_genre,
                "channel": channel_var.get(),
                "tags": selected_tags,
                "concept": concept,
                "chapter_count": chapters,
                "word_count_per_chapter": int(word_count_var.get()),
                "created_at": datetime.now().isoformat(),
                "template": template_name if template_name != "无" else None,
            }
            with open(novel_dir / "meta.json", 'w', encoding='utf-8') as f:
                json.dump(meta, f, indent=2, ensure_ascii=False)
            
            # 初始化
            self.current_novel_dir = novel_dir
            self.memory = MemoryManager(novel_dir)
            self.agent = NovelAgent(self.ai_client, self.memory, log_callback=self._log, config=self.config)
            self.note_manager = NoteManager(novel_dir=novel_dir, config=self.config)
            self.outline = []
            self.current_chapter = 0
            self._init_character_system()
            
            self.title_var.set(title)
            self.genre_var.set(f"{genre}-{sub_genre}")
            self.chapter_var.set(f"0/{chapters}")
            
            dialog.destroy()
            self._log(f"新建小说《{title}》({sub_genre}) 创建成功，标签: {', '.join(selected_tags)}")
        
        tk.Button(action_row, text="创建小说", command=confirm, font=('微软雅黑', 10, 'bold'),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=20, pady=3).pack(side=tk.RIGHT)
    
    def _open_novel(self):
        """打开小说 - 支持从列表选择或浏览目录"""
        C = UIStyle.COLORS
        
        # 创建选择对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("打开小说")
        dialog.geometry("500x400")
        dialog.configure(bg=C['bg_dark'])
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text="选择小说", font=('微软雅黑', 14, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=15)
        
        # 小说列表
        list_frame = tk.Frame(dialog, bg=C['bg_dark'])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # 获取已有小说列表
        novels_dir = self.config.novels_dir
        novel_list = []
        
        if novels_dir.exists():
            for d in sorted(novels_dir.iterdir(), reverse=True):
                if d.is_dir():
                    meta_file = d / 'meta.json'
                    if meta_file.exists():
                        try:
                            with open(meta_file, 'r', encoding='utf-8') as f:
                                meta = json.load(f)
                            novel_list.append({
                                'dir': d,
                                'title': meta.get('title', d.name),
                                'genre': meta.get('genre', '未知'),
                                'chapters': len(list((d / 'chapters').glob('chapter_*.txt'))) if (d / 'chapters').exists() else 0,
                            })
                        except:
                            pass
        
        if novel_list:
            tk.Label(dialog, text="已有小说:", font=('微软雅黑', 10),
                    bg=C['bg_dark'], fg=C['text_secondary']).pack(anchor=tk.W, padx=20)
            
            # 列表框
            listbox = tk.Listbox(list_frame, bg=C['bg_card'], fg=C['text_primary'],
                               font=('微软雅黑', 10), selectbackground=C['accent'],
                               selectforeground='white', relief=tk.FLAT)
            scrollbar = tk.Scrollbar(list_frame, command=listbox.yview)
            listbox.configure(yscrollcommand=scrollbar.set)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            listbox.pack(fill=tk.BOTH, expand=True)
            
            for novel in novel_list:
                listbox.insert(tk.END, f"{novel['title']} - {novel['genre']} ({novel['chapters']}章)")
            
            # 选择按钮
            def open_selected():
                selection = listbox.curselection()
                if selection:
                    idx = selection[0]
                    novel = novel_list[idx]
                    novel_dir = novel['dir']
                    dialog.destroy()
                    # 延迟执行加载，确保对话框关闭后再加载
                    self.root.after(100, lambda: self._load_novel(novel_dir))
                else:
                    messagebox.showwarning("提示", "请先选择一个小说")
            
            tk.Button(dialog, text="打开选中", command=open_selected,
                     bg=C['accent'], fg='white', font=('微软雅黑', 10),
                     relief=tk.FLAT, padx=20, pady=5).pack(pady=10)
        
        # 浏览按钮
        def browse_dir():
            dialog.destroy()
            novel_dir = filedialog.askdirectory(
                title="选择小说目录",
                initialdir=str(self.config.novels_dir)
            )
            if novel_dir:
                self._load_novel(Path(novel_dir))
        
        tk.Button(dialog, text="浏览目录...", command=browse_dir,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10),
                 relief=tk.FLAT, padx=20, pady=5).pack(pady=5)
    
    def _load_novel(self, novel_dir: Path):
        """加载小说数据"""
        self._log(f"正在加载小说: {novel_dir}")
        
        meta_file = novel_dir / "meta.json"
        
        if not meta_file.exists():
            messagebox.showerror("错误", "该目录不是有效的小说目录")
            return
        
        try:
            with open(meta_file, 'r', encoding='utf-8') as f:
                meta = json.load(f)
        except Exception as e:
            messagebox.showerror("错误", f"读取meta.json失败: {e}")
            return
        
        self.current_novel_dir = novel_dir
        self.memory = MemoryManager(novel_dir)
        self.agent = NovelAgent(self.ai_client, self.memory, log_callback=self._log, config=self.config)
        self.note_manager = NoteManager(novel_dir=novel_dir, config=self.config)
        self._init_character_system()
        self.title_var.set(meta.get("title", "未知"))
        self.genre_var.set(meta.get("genre", "未知"))
        
        # 加载大纲
        outline_file = novel_dir / "outline.json"
        if outline_file.exists():
            try:
                with open(outline_file, 'r', encoding='utf-8') as f:
                    self.outline = json.load(f)
                self._refresh_outline_list()
                self._log(f"已加载大纲: {len(self.outline)} 章")
            except Exception as e:
                self._log(f"加载大纲失败: {e}")
        
        # 更新章节选择器
        self._update_chapter_selector()
        
        # 计算进度
        chapters_dir = novel_dir / "chapters"
        total_chapters = meta.get('chapter_count', 0)
        completed_chapters = 0
        
        if chapters_dir.exists():
            chapter_files = sorted(chapters_dir.glob("chapter_*.txt"))
            completed_chapters = len(chapter_files)
            self.current_chapter = completed_chapters
            self.chapter_var.set(f"{completed_chapters}/{total_chapters or '?'}")
            
            # 加载最后一章内容到编辑器
            if chapter_files:
                last_chapter_file = chapter_files[-1]
                try:
                    content = last_chapter_file.read_text(encoding='utf-8')
                    self.content_text.delete("1.0", tk.END)
                    self.content_text.insert("1.0", content)
                    self.word_count_var.set(f"字数: {len(content)}")
                    chapter_num = int(last_chapter_file.stem.split('_')[-1])
                    self.chapter_select_var.set(f"第{chapter_num}章")
                    self._log(f"已加载第{chapter_num}章内容")
                except Exception as e:
                    self._log(f"加载最后一章失败: {e}")
        
        # 切换到章节内容标签页
        try:
            for i in range(self.notebook.index("end")):
                if self.notebook.tab(i, "text").strip() == "章节内容":
                    self.notebook.select(i)
                    break
        except Exception as e:
            self._log(f"切换标签页失败: {e}")
        
        self._log(f"已打开小说《{meta.get('title', '未知')}》")
        
        # 显示进度提示
        if total_chapters and completed_chapters > 0:
            if completed_chapters < total_chapters:
                remaining = total_chapters - completed_chapters
                msg = f"小说《{meta.get('title')}》进度：已完成 {completed_chapters}/{total_chapters} 章\n\n"
                msg += f"还剩 {remaining} 章未完成。\n\n"
                msg += "接下来可以：\n"
                msg += "1. 点击「自动创作」继续自动生成剩余章节\n"
                msg += "2. 点击「生成下一章」手动逐章创作\n"
                msg += "3. 在编辑器中手动编写"
                self.root.after(500, lambda: messagebox.showinfo("继续创作", msg))
            elif completed_chapters >= total_chapters:
                result = messagebox.askyesno("已完成",
                    f"小说《{meta.get('title')}》已全部完成！共 {completed_chapters} 章。\n\n"
                    "是否续写新章？")
                if result:
                    self._continue_novel()
    
    def _continue_novel(self):
        """续写已完成的小说 - 扩展大纲和新章节"""
        if not self.current_novel_dir:
            return
        
        # 计算当前总章数
        chapters_dir = self.current_novel_dir / "chapters"
        existing = sorted(chapters_dir.glob("chapter_*.txt")) if chapters_dir.exists() else []
        current_count = len(existing)
        
        # 询问续写多少章
        add_count = simpledialog.askinteger("续写", 
            f"当前已完成 {current_count} 章\n输入要续写的章数（10-50）:",
            minvalue=1, maxvalue=100, initialvalue=10)
        
        if not add_count:
            return
        
        self._log(f"开始续写 {add_count} 章...")
        
        # 扩展大纲
        if not self.outline:
            self.outline = []
        
        meta = self._get_meta()
        title = meta.get("title", "小说")
        genre = meta.get("genre", "未知")
        
        # 生成新的大纲章节
        try:
            context = self.memory.get_global_summary() if self.memory else ""
            new_chapters = self.agent.generate_outline_continuation(
                genre, title, add_count, context, current_count
            )
            self.outline.extend(new_chapters)
            
            # 保存更新后的大纲
            with open(self.current_novel_dir / "outline.json", 'w', encoding='utf-8') as f:
                json.dump(self.outline, f, indent=2, ensure_ascii=False)
            
            # 更新meta
            meta_file = self.current_novel_dir / "meta.json"
            if meta_file.exists():
                with open(meta_file, 'w', encoding='utf-8') as f:
                    meta['chapter_count'] = len(self.outline)
                    json.dump(meta, f, indent=2, ensure_ascii=False)
            
            self.current_chapter = current_count
            self._refresh_outline_list()
            self.chapter_var.set(f"{current_count}/{len(self.outline)}")
            self._log(f"大纲已扩展到 {len(self.outline)} 章，可以继续创作了")
            
            messagebox.showinfo("续写", 
                f"已添加 {add_count} 章新大纲\n"
                f"总章数: {len(self.outline)}\n\n"
                "点击「自动创作」或「生成下一章」继续写作")
            
        except Exception as e:
            self._log(f"续写大纲生成失败: {e}")
            messagebox.showerror("错误", f"续写失败: {e}")
    
    def _create_sequel(self):
        """基于当前小说创建续集（第二部）"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先打开一部已完成的小说")
            return
        
        meta_file = self.current_novel_dir / "meta.json"
        if not meta_file.exists():
            messagebox.showerror("错误", "当前目录不是有效的小说目录")
            return
        
        with open(meta_file, 'r', encoding='utf-8') as f:
            original_meta = json.load(f)
        
        # 检查是否已完成
        chapters_dir = self.current_novel_dir / "chapters"
        if chapters_dir.exists():
            chapter_count = len(list(chapters_dir.glob("chapter_*.txt")))
            if chapter_count < original_meta.get("chapter_count", 0):
                if not messagebox.askyesno("提示", "当前小说尚未全部完成，确定要创建续集吗？"):
                    return
        
        # 读取原始小说的全局摘要
        global_summary = ""
        summary_file = self.current_novel_dir / "memory" / "global_summary.txt"
        if summary_file.exists():
            global_summary = summary_file.read_text(encoding='utf-8')
        
        # 创建续集对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("创建续集 - 第二部")
        dialog.geometry("600x500")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"《{original_meta.get('title', '')}》续集", 
                font=('微软雅黑', 14, 'bold'), bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        # 续集标题
        title_frame = tk.Frame(dialog, bg=C['bg_dark'])
        title_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(title_frame, text="续集标题:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        title_entry = tk.Entry(title_frame, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
        title_entry.insert(0, f"{original_meta.get('title', '')} 第二部")
        title_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 续集概念
        tk.Label(dialog, text="续集概念/方向:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(10, 3))
        concept_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'], height=5)
        concept_text.pack(fill=tk.X, padx=20, pady=5)
        concept_text.insert("1.0", "延续第一部的世界观和角色，展开新的冒险...")
        
        # 原著摘要预览
        if global_summary:
            tk.Label(dialog, text="原著摘要（AI将基于此生成续集）:", bg=C['bg_dark'], fg=C['text_muted']).pack(anchor=tk.W, padx=20, pady=(10, 3))
            summary_preview = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 9), bg=C['bg_card'], fg=C['text_secondary'], height=4)
            summary_preview.pack(fill=tk.X, padx=20, pady=5)
            summary_preview.insert("1.0", global_summary[:500] + ("..." if len(global_summary) > 500 else ""))
            summary_preview.config(state=tk.DISABLED)
        
        # 章节数和字数
        params_frame = tk.Frame(dialog, bg=C['bg_dark'])
        params_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Label(params_frame, text="章节数:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        chapters_var = tk.StringVar(value=str(original_meta.get("chapter_count", 20)))
        tk.Spinbox(params_frame, from_=1, to=500, textvariable=chapters_var, width=6, font=('微软雅黑', 9)).pack(side=tk.LEFT, padx=5)
        tk.Label(params_frame, text="每章字数:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT, padx=(15, 0))
        word_count_var = tk.StringVar(value=str(original_meta.get("word_count_per_chapter", 3000)))
        ttk.Combobox(params_frame, textvariable=word_count_var, values=["1000","2000","3000","5000","8000"], width=8).pack(side=tk.LEFT, padx=5)
        
        def confirm():
            title = title_entry.get().strip()
            if not title:
                messagebox.showwarning("提示", "请输入续集标题")
                return
            
            concept = concept_text.get("1.0", tk.END).strip()
            
            # 创建续集目录
            safe_name = "".join(c for c in title if c.isalnum() or c in "_ -")[:30]
            novel_dir = self.config.novels_dir / f"{safe_name}_{int(time.time())}"
            novel_dir.mkdir(exist_ok=True)
            
            # 保存续集元数据（关联原著）
            sequel_meta = {
                "title": title,
                "genre": original_meta.get("genre", ""),
                "sub_genre": original_meta.get("sub_genre", ""),
                "channel": original_meta.get("channel", "male"),
                "tags": original_meta.get("tags", []),
                "concept": concept,
                "chapter_count": int(chapters_var.get()),
                "word_count_per_chapter": int(word_count_var.get()),
                "created_at": datetime.now().isoformat(),
                "is_sequel": True,
                "original_novel": str(self.current_novel_dir),
                "original_title": original_meta.get("title", ""),
            }
            with open(novel_dir / "meta.json", 'w', encoding='utf-8') as f:
                json.dump(sequel_meta, f, indent=2, ensure_ascii=False)
            
            # 复制原著的世界观和角色设定
            orig_memory = self.current_novel_dir / "memory"
            new_memory = novel_dir / "memory"
            new_memory.mkdir(exist_ok=True)
            
            # 复制世界观
            orig_settings = orig_memory / "settings.json"
            if orig_settings.exists():
                shutil.copy2(orig_settings, new_memory / "settings.json")
            
            # 复制角色
            orig_chars = self.current_novel_dir / "characters"
            if orig_chars.exists():
                shutil.copytree(orig_chars, novel_dir / "characters", dirs_exist_ok=True)
            
            # 保存续集概念作为参考
            with open(novel_dir / "sequel_concept.txt", 'w', encoding='utf-8') as f:
                f.write(f"原著: {original_meta.get('title', '')}\n\n")
                f.write(f"原著摘要:\n{global_summary}\n\n")
                f.write(f"续集概念:\n{concept}")
            
            # 切换到续集
            self.current_novel_dir = novel_dir
            self.memory = MemoryManager(novel_dir)
            self.agent = NovelAgent(self.ai_client, self.memory, log_callback=self._log, config=self.config)
            self.note_manager = NoteManager(novel_dir=novel_dir, config=self.config)
            self.outline = []
            self.current_chapter = 0
            self._init_character_system()
            
            self.title_var.set(title)
            self.genre_var.set(f"{original_meta.get('genre', '')}-{original_meta.get('sub_genre', '')}")
            self.chapter_var.set(f"0/{chapters_var.get()}")
            
            dialog.destroy()
            self._log(f"续集《{title}》已创建，基于原著《{original_meta.get('title', '')}》")
            messagebox.showinfo("成功", f"续集《{title}》已创建！\n世界观和角色已继承自原著。\n点击「自动创作」开始生成。")
        
        tk.Button(dialog, text="创建续集", command=confirm, bg=C['accent'], fg='white',
                 font=('微软雅黑', 11, 'bold'), padx=30, pady=8).pack(pady=15)
    
    def _create_spinoff(self):
        """基于当前小说创建同人衍生作品"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先打开一部小说作为原著")
            return
        
        meta_file = self.current_novel_dir / "meta.json"
        if not meta_file.exists():
            messagebox.showerror("错误", "当前目录不是有效的小说目录")
            return
        
        with open(meta_file, 'r', encoding='utf-8') as f:
            original_meta = json.load(f)
        
        # 读取角色信息
        characters = self.memory.get_characters() if self.memory else {}
        char_names = list(characters.keys()) if characters else []
        
        # 创建同人作品对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("创建同人衍生作品")
        dialog.geometry("650x600")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"基于《{original_meta.get('title', '')}》的同人作品", 
                font=('微软雅黑', 14, 'bold'), bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        # 同人作品标题
        title_frame = tk.Frame(dialog, bg=C['bg_dark'])
        title_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(title_frame, text="作品标题:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        title_entry = tk.Entry(title_frame, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
        title_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 衍生类型
        type_frame = tk.Frame(dialog, bg=C['bg_dark'])
        type_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(type_frame, text="衍生类型:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        spinoff_type = tk.StringVar(value="平行世界")
        ttk.Combobox(type_frame, textvariable=spinoff_type, 
                    values=["平行世界", "角色外传", "前传", "IF线", "现代AU", "古代AU", "其他"],
                    state="readonly", width=15).pack(side=tk.LEFT, padx=5)
        
        # 选择主要角色
        if char_names:
            tk.Label(dialog, text="选择主要角色（可多选）:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(10, 3))
            char_frame = tk.Frame(dialog, bg=C['bg_card'])
            char_frame.pack(fill=tk.X, padx=20, pady=5)
            
            char_vars = {}
            for name in char_names[:10]:  # 最多显示10个角色
                var = tk.BooleanVar(value=False)
                char_vars[name] = var
                tk.Checkbutton(char_frame, text=name, variable=var, bg=C['bg_card'], fg=C['text_primary'],
                             selectcolor=C['bg_dark'], font=('微软雅黑', 9)).pack(side=tk.LEFT, padx=5)
        
        # 衍生概念
        tk.Label(dialog, text="衍生概念/设定:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(10, 3))
        concept_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'], height=6)
        concept_text.pack(fill=tk.X, padx=20, pady=5)
        concept_text.insert("1.0", "在这个平行世界中...")
        
        # 章节数和字数
        params_frame = tk.Frame(dialog, bg=C['bg_dark'])
        params_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Label(params_frame, text="章节数:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        chapters_var = tk.StringVar(value="10")
        tk.Spinbox(params_frame, from_=1, to=200, textvariable=chapters_var, width=6, font=('微软雅黑', 9)).pack(side=tk.LEFT, padx=5)
        tk.Label(params_frame, text="每章字数:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT, padx=(15, 0))
        word_count_var = tk.StringVar(value="3000")
        ttk.Combobox(params_frame, textvariable=word_count_var, values=["1000","2000","3000","5000"], width=8).pack(side=tk.LEFT, padx=5)
        
        def confirm():
            title = title_entry.get().strip()
            if not title:
                messagebox.showwarning("提示", "请输入作品标题")
                return
            
            concept = concept_text.get("1.0", tk.END).strip()
            selected_chars = [name for name, var in char_vars.items() if var.get()] if char_names else []
            
            # 创建同人作品目录
            safe_name = "".join(c for c in title if c.isalnum() or c in "_ -")[:30]
            novel_dir = self.config.novels_dir / f"{safe_name}_{int(time.time())}"
            novel_dir.mkdir(exist_ok=True)
            
            # 保存同人作品元数据
            spinoff_meta = {
                "title": title,
                "genre": original_meta.get("genre", ""),
                "sub_genre": original_meta.get("sub_genre", ""),
                "channel": original_meta.get("channel", "male"),
                "tags": original_meta.get("tags", []) + ["同人", spinoff_type.get()],
                "concept": concept,
                "chapter_count": int(chapters_var.get()),
                "word_count_per_chapter": int(word_count_var.get()),
                "created_at": datetime.now().isoformat(),
                "is_spinoff": True,
                "spinoff_type": spinoff_type.get(),
                "original_novel": str(self.current_novel_dir),
                "original_title": original_meta.get("title", ""),
                "selected_characters": selected_chars,
            }
            with open(novel_dir / "meta.json", 'w', encoding='utf-8') as f:
                json.dump(spinoff_meta, f, indent=2, ensure_ascii=False)
            
            # 复制世界观设定
            orig_memory = self.current_novel_dir / "memory"
            new_memory = novel_dir / "memory"
            new_memory.mkdir(exist_ok=True)
            
            orig_settings = orig_memory / "settings.json"
            if orig_settings.exists():
                shutil.copy2(orig_settings, new_memory / "settings.json")
            
            # 复制选定的角色
            if selected_chars and characters:
                chars_dir = novel_dir / "characters"
                chars_dir.mkdir(exist_ok=True)
                for char_name in selected_chars:
                    char_file = self.current_novel_dir / "characters" / f"{char_name}.json"
                    if char_file.exists():
                        shutil.copy2(char_file, chars_dir / f"{char_name}.json")
            
            # 保存同人设定文档
            with open(novel_dir / "spinoff_concept.txt", 'w', encoding='utf-8') as f:
                f.write(f"原著: {original_meta.get('title', '')}\n")
                f.write(f"衍生类型: {spinoff_type.get()}\n")
                f.write(f"主要角色: {', '.join(selected_chars)}\n\n")
                f.write(f"衍生概念:\n{concept}")
            
            # 切换到同人作品
            self.current_novel_dir = novel_dir
            self.memory = MemoryManager(novel_dir)
            self.agent = NovelAgent(self.ai_client, self.memory, log_callback=self._log, config=self.config)
            self.note_manager = NoteManager(novel_dir=novel_dir, config=self.config)
            self.outline = []
            self.current_chapter = 0
            self._init_character_system()
            
            self.title_var.set(title)
            self.chapter_var.set(f"0/{chapters_var.get()}")
            
            dialog.destroy()
            self._log(f"同人作品《{title}》已创建，类型：{spinoff_type.get()}")
            messagebox.showinfo("成功", f"同人作品《{title}》已创建！\n类型：{spinoff_type.get()}\n角色：{', '.join(selected_chars) or '无'}\n点击「自动创作」开始生成。")
        
        tk.Button(dialog, text="创建同人作品", command=confirm, bg=C['accent'], fg='white',
                 font=('微软雅黑', 11, 'bold'), padx=30, pady=8).pack(pady=15)
    
    def _show_settings(self):
        """显示设置对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("系统配置")
        dialog.geometry("550x650")
        dialog.transient(self.root)
        dialog.grab_set()
        
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # ===== Tab 1: AI模型配置 =====
        ai_frame = ttk.Frame(notebook)
        notebook.add(ai_frame, text="AI模型")
        
        ttk.Label(ai_frame, text="AI服务商:").pack(anchor=tk.W, padx=20, pady=(15,3))
        provider_var = tk.StringVar(value=self.config.get("api_provider", "ollama"))
        provider_combo = ttk.Combobox(ai_frame, textvariable=provider_var, 
            values=["ollama", "openai", "deepseek", "claude", "custom"], state="readonly", width=50)
        provider_combo.pack(padx=20, pady=3)
        
        ttk.Label(ai_frame, text="API地址 (Ollama默认 http://localhost:11434):").pack(anchor=tk.W, padx=20, pady=(10,3))
        base_entry = ttk.Entry(ai_frame, width=52)
        base_entry.insert(0, self.config.get("api_base", "http://localhost:11434"))
        base_entry.pack(padx=20, pady=3)
        
        ttk.Label(ai_frame, text="API密钥 (Ollama不需要):").pack(anchor=tk.W, padx=20, pady=(10,3))
        key_entry = ttk.Entry(ai_frame, width=52, show="*")
        key_entry.insert(0, self.config.get("api_key", ""))
        key_entry.pack(padx=20, pady=3)
        
        ttk.Label(ai_frame, text="模型名称:").pack(anchor=tk.W, padx=20, pady=(10,3))
        model_frame = ttk.Frame(ai_frame)
        model_frame.pack(fill=tk.X, padx=20, pady=3)
        model_var = tk.StringVar(value=self.config.get("model", "qwen2.5:14b"))
        model_combo = ttk.Combobox(model_frame, textvariable=model_var, width=38)
        model_combo.pack(side=tk.LEFT)
        model_combo.set(self.config.get("model", "qwen2.5:14b"))
        
        def refresh_ollama():
            models = self.ai_client.get_ollama_models()
            if models:
                model_combo['values'] = models
                model_var.set(models[0])
                messagebox.showinfo("Ollama模型", f"发现 {len(models)} 个模型:\n" + "\n".join(models[:10]))
            else:
                messagebox.showwarning("提示", "未检测到Ollama模型，请确保Ollama已启动")
        ttk.Button(model_frame, text="检测Ollama", command=refresh_ollama).pack(side=tk.LEFT, padx=5)
        
        ttk.Label(ai_frame, text="温度 (0-1):").pack(anchor=tk.W, padx=20, pady=(10,3))
        temp_var = tk.StringVar(value=str(self.config.get("temperature", 0.8)))
        ttk.Spinbox(ai_frame, from_=0, to=1, increment=0.1, textvariable=temp_var, width=10).pack(anchor=tk.W, padx=20, pady=3)
        
        ttk.Label(ai_frame, text="上下文窗口 (字符数，影响长记忆):").pack(anchor=tk.W, padx=20, pady=(10,3))
        ctx_var = tk.StringVar(value=str(self.config.get("context_window", 32000)))
        ctx_combo = ttk.Combobox(ai_frame, textvariable=ctx_var, 
            values=["8000", "16000", "32000", "64000", "128000"], width=10)
        ctx_combo.pack(anchor=tk.W, padx=20, pady=3)
        ttk.Label(ai_frame, text="Ollama模型取决于显存，建议8K-32K；云端API可用更大窗口", 
                  foreground="gray").pack(anchor=tk.W, padx=20)
        
        # ===== Tab 2: 文生图配置 =====
        img_frame = ttk.Frame(notebook)
        notebook.add(img_frame, text="文生图")
        
        ttk.Label(img_frame, text="文生图后端:").pack(anchor=tk.W, padx=20, pady=(15,3))
        img_provider_var = tk.StringVar(value=self.config.get("img_provider", "comfyui"))
        ttk.Combobox(img_frame, textvariable=img_provider_var, 
            values=["comfyui", "sdapi", "disabled"], state="readonly", width=50).pack(padx=20, pady=3)
        
        ttk.Label(img_frame, text="API地址:").pack(anchor=tk.W, padx=20, pady=(10,3))
        img_base_entry = ttk.Entry(img_frame, width=52)
        img_base_entry.insert(0, self.config.get("img_api_base", "http://127.0.0.1:8188"))
        img_base_entry.pack(padx=20, pady=3)
        
        ttk.Label(img_frame, text="ComfyUI默认端口: 8188, SD WebUI默认端口: 7860").pack(anchor=tk.W, padx=20, pady=(3,10))
        
        ttk.Label(img_frame, text="模型文件名:").pack(anchor=tk.W, padx=20, pady=(5,3))
        img_model_entry = ttk.Entry(img_frame, width=52)
        img_model_entry.insert(0, self.config.get("img_model", "sd_xl_base_1.0.safetensors"))
        img_model_entry.pack(padx=20, pady=3)
        
        auto_detect_var = tk.BooleanVar(value=self.config.get("auto_detect_scene", True))
        ttk.Checkbutton(img_frame, text="生成章节后自动检测名场面并提醒生成插图", variable=auto_detect_var).pack(anchor=tk.W, padx=20, pady=15)
        
        ttk.Label(img_frame, text="支持的后端:\n- ComfyUI: 本地部署的ComfyUI，需启动API模式\n- SD API: Stable Diffusion WebUI的API模式\n- Disabled: 不使用文生图", 
                  justify=tk.LEFT).pack(anchor=tk.W, padx=20, pady=10)
        
        # ===== Tab 3: 云端存储配置 =====
        cloud_frame = ttk.Frame(notebook)
        notebook.add(cloud_frame, text="云端存储")
        
        ttk.Label(cloud_frame, text="云端存储配置", font=("", 11, "bold")).pack(anchor=tk.W, padx=20, pady=(15,10))
        ttk.Label(cloud_frame, text="支持: WebDAV（坚果云）、百度网盘、夸克网盘、迅雷网盘、阿里云盘").pack(anchor=tk.W, padx=20, pady=(0,10))
        
        # 云存储提供商选择
        cloud_provider_var = tk.StringVar(value="webdav")
        providers = self.cloud_storage.get_available_providers()
        provider_names = [p["name"] for p in providers]
        provider_ids = [p["id"] for p in providers]
        
        ttk.Label(cloud_frame, text="选择云存储:").pack(anchor=tk.W, padx=20, pady=(5,3))
        cloud_combo = ttk.Combobox(cloud_frame, textvariable=cloud_provider_var, 
                                   values=provider_names, state="readonly", width=50)
        cloud_combo.pack(padx=20, pady=3)
        cloud_combo.set(provider_names[0] if provider_names else "")
        
        # 配置区域
        config_frame = ttk.LabelFrame(cloud_frame, text="配置信息", padding=10)
        config_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # WebDAV配置
        webdav_frame = ttk.Frame(config_frame)
        ttk.Label(webdav_frame, text="WebDAV地址:").grid(row=0, column=0, sticky=tk.W, pady=2)
        webdav_url_entry = ttk.Entry(webdav_frame, width=40)
        webdav_url_entry.insert(0, self.cloud_storage.config.get("webdav", {}).get("url", "https://dav.jianguoyun.com/dav/"))
        webdav_url_entry.grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(webdav_frame, text="用户名:").grid(row=1, column=0, sticky=tk.W, pady=2)
        webdav_user_entry = ttk.Entry(webdav_frame, width=40)
        webdav_user_entry.insert(0, self.cloud_storage.config.get("webdav", {}).get("username", ""))
        webdav_user_entry.grid(row=1, column=1, padx=5, pady=2)
        
        ttk.Label(webdav_frame, text="密码/应用密钥:").grid(row=2, column=0, sticky=tk.W, pady=2)
        webdav_pass_entry = ttk.Entry(webdav_frame, width=40, show="*")
        webdav_pass_entry.insert(0, self.cloud_storage.config.get("webdav", {}).get("password", ""))
        webdav_pass_entry.grid(row=2, column=1, padx=5, pady=2)
        webdav_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 百度网盘配置
        baidu_frame = ttk.Frame(config_frame)
        ttk.Label(baidu_frame, text="Access Token:").grid(row=0, column=0, sticky=tk.W, pady=2)
        baidu_token_entry = ttk.Entry(baidu_frame, width=40)
        baidu_token_entry.insert(0, self.cloud_storage.config.get("baidu", {}).get("access_token", ""))
        baidu_token_entry.grid(row=0, column=1, padx=5, pady=2)
        baidu_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 夸克网盘配置
        quark_frame = ttk.Frame(config_frame)
        ttk.Label(quark_frame, text="Cookie:").grid(row=0, column=0, sticky=tk.W, pady=2)
        quark_cookie_entry = ttk.Entry(quark_frame, width=40)
        quark_cookie_entry.insert(0, self.cloud_storage.config.get("quark", {}).get("cookie", ""))
        quark_cookie_entry.grid(row=0, column=1, padx=5, pady=2)
        quark_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 迅雷网盘配置
        xunlei_frame = ttk.Frame(config_frame)
        ttk.Label(xunlei_frame, text="Access Token:").grid(row=0, column=0, sticky=tk.W, pady=2)
        xunlei_token_entry = ttk.Entry(xunlei_frame, width=40)
        xunlei_token_entry.insert(0, self.cloud_storage.config.get("xunlei", {}).get("access_token", ""))
        xunlei_token_entry.grid(row=0, column=1, padx=5, pady=2)
        xunlei_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 阿里云盘配置
        aliyun_frame = ttk.Frame(config_frame)
        ttk.Label(aliyun_frame, text="Access Token:").grid(row=0, column=0, sticky=tk.W, pady=2)
        aliyun_token_entry = ttk.Entry(aliyun_frame, width=40)
        aliyun_token_entry.insert(0, self.cloud_storage.config.get("aliyun", {}).get("access_token", ""))
        aliyun_token_entry.grid(row=0, column=1, padx=5, pady=2)
        aliyun_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 测试连接按钮
        def test_cloud_connection():
            provider_name = cloud_combo.get()
            provider_id = provider_ids[provider_names.index(provider_name)] if provider_name in provider_names else "webdav"
            
            # 保存配置
            if provider_id == "webdav":
                self.cloud_storage.configure_provider("webdav", {
                    "url": webdav_url_entry.get(),
                    "username": webdav_user_entry.get(),
                    "password": webdav_pass_entry.get()
                })
            elif provider_id == "baidu":
                self.cloud_storage.configure_provider("baidu", {
                    "access_token": baidu_token_entry.get()
                })
            elif provider_id == "quark":
                self.cloud_storage.configure_provider("quark", {
                    "cookie": quark_cookie_entry.get()
                })
            elif provider_id == "xunlei":
                self.cloud_storage.configure_provider("xunlei", {
                    "access_token": xunlei_token_entry.get()
                })
            elif provider_id == "aliyun":
                self.cloud_storage.configure_provider("aliyun", {
                    "access_token": aliyun_token_entry.get()
                })
            
            # 测试连接
            if self.cloud_storage.connect_provider(provider_id):
                messagebox.showinfo("成功", f"{provider_name} 连接成功！")
            else:
                messagebox.showwarning("失败", f"{provider_name} 连接失败，请检查配置")
        
        ttk.Button(cloud_frame, text="测试连接", command=test_cloud_connection).pack(pady=10)
        
        # ===== Tab 5: 高级设置 =====
        advanced_frame = ttk.Frame(notebook)
        notebook.add(advanced_frame, text="高级")
        
        # 18+内容开关（隐藏按钮）
        adult_frame = tk.LabelFrame(advanced_frame, text=" 内容控制 ", padx=10, pady=10)
        adult_frame.pack(fill=tk.X, padx=15, pady=10)
        
        # 三角形隐藏按钮
        secret_frame = tk.Frame(adult_frame)
        secret_frame.pack(fill=tk.X, pady=5)
        
        self._adult_toggle_visible = False
        adult_var = tk.BooleanVar(value=self.config.get("adult_content", False))
        
        def toggle_adult_visibility():
            """点击三角形显示/隐藏18+开关"""
            self._adult_toggle_visible = not self._adult_toggle_visible
            if self._adult_toggle_visible:
                adult_controls.pack(fill=tk.X, pady=5)
                secret_btn.configure(text="▼ 18+内容设置")
            else:
                adult_controls.pack_forget()
                secret_btn.configure(text="▶ 点击展开更多设置")
        
        secret_btn = tk.Button(secret_frame, text="▶ 点击展开更多设置", 
                              command=toggle_adult_visibility, relief=tk.FLAT,
                              fg='gray', font=('微软雅黑', 8), anchor=tk.W)
        secret_btn.pack(side=tk.LEFT)
        
        adult_controls = tk.Frame(adult_frame)
        
        tk.Label(adult_controls, text="⚠️ 以下功能仅供成年用户使用", 
                fg='red', font=('微软雅黑', 9, 'bold')).pack(anchor=tk.W, pady=(5, 10))
        
        adult_check = tk.Checkbutton(adult_controls, text="启用18+内容生成", variable=adult_var,
                                     font=('微软雅黑', 10))
        adult_check.pack(anchor=tk.W, pady=3)
        
        edge_var = tk.BooleanVar(value=self.config.get("edge_content", False))
        edge_check = tk.Checkbutton(adult_controls, text="允许擦边内容", variable=edge_var,
                                    font=('微软雅黑', 10))
        edge_check.pack(anchor=tk.W, pady=3)
        
        tk.Label(adult_controls, text="启用后，AI在创作时会根据剧情需要加入相关描写", 
                fg='gray', font=('微软雅黑', 8)).pack(anchor=tk.W, pady=(5, 0))
        
        # 卷管理设置
        volume_frame = tk.LabelFrame(advanced_frame, text=" 卷管理 ", padx=10, pady=10)
        volume_frame.pack(fill=tk.X, padx=15, pady=10)
        
        tk.Label(volume_frame, text="每卷默认章节数:", font=('微软雅黑', 10)).pack(anchor=tk.W, pady=3)
        vol_chapters_var = tk.StringVar(value=str(self.config.get("chapters_per_volume", 100)))
        ttk.Spinbox(volume_frame, from_=10, to=500, textvariable=vol_chapters_var, width=10).pack(anchor=tk.W, pady=3)
        
        tk.Label(volume_frame, text="角色传记默认字数:", font=('微软雅黑', 10)).pack(anchor=tk.W, pady=(10, 3))
        bio_words_var = tk.StringVar(value=str(self.config.get("biography_word_count", 100000)))
        bio_combo = ttk.Combobox(volume_frame, textvariable=bio_words_var,
                                values=["10000", "30000", "50000", "100000", "200000"], width=10)
        bio_combo.pack(anchor=tk.W, pady=3)
        tk.Label(volume_frame, text="生成角色个人传时的默认字数", fg='gray', font=('微软雅黑', 8)).pack(anchor=tk.W)
        
        # 智能体优化
        agent_frame = tk.LabelFrame(advanced_frame, text=" 智能体优化 ", padx=10, pady=10)
        agent_frame.pack(fill=tk.X, padx=15, pady=10)
        
        context_var = tk.BooleanVar(value=self.config.get("smart_context", True))
        tk.Checkbutton(agent_frame, text="智能上下文管理（防止章节过多卡死）", variable=context_var,
                      font=('微软雅黑', 10)).pack(anchor=tk.W, pady=3)
        
        summary_var = tk.BooleanVar(value=self.config.get("auto_summary", True))
        tk.Checkbutton(agent_frame, text="自动生成章节摘要（改善上下文连贯性）", variable=summary_var,
                      font=('微软雅黑', 10)).pack(anchor=tk.W, pady=3)
        
        # ===== 保存 =====
        def save():
            self.config.set("api_provider", provider_var.get())
            self.config.set("api_key", key_entry.get().strip())
            self.config.set("api_base", base_entry.get().strip())
            self.config.set("model", model_var.get().strip())
            self.config.set("temperature", float(temp_var.get()))
            self.config.set("context_window", int(ctx_var.get()))
            self.config.set("img_provider", img_provider_var.get())
            self.config.set("img_api_base", img_base_entry.get().strip())
            self.config.set("img_model", img_model_entry.get().strip())
            self.config.set("auto_detect_scene", auto_detect_var.get())
            self.config.set("adult_content", adult_var.get())
            self.config.set("edge_content", edge_var.get())
            self.config.set("chapters_per_volume", int(vol_chapters_var.get()))
            self.config.set("biography_word_count", int(bio_words_var.get()))
            self.config.set("smart_context", context_var.get())
            self.config.set("auto_summary", summary_var.get())
            
            self.ai_client = AIClient(self.config)
            self.image_gen = ImageGenerator(self.config)
            if self.memory:
                self.agent = NovelAgent(self.ai_client, self.memory, log_callback=self._log, config=self.config)
            
            self._update_status()
            dialog.destroy()
            self._log("配置已保存")
        
        ttk.Button(dialog, text="保存配置", command=save).pack(pady=10)
    
    def _gen_settings(self):
        """生成世界观"""
        if not self._check_ready():
            return
        if hasattr(self, '_gen_running') and self._gen_running:
            self._log("生成任务正在进行中，请勿重复点击")
            return
        
        self._gen_running = True
        
        def run():
            try:
                meta = self._get_meta()
                self.agent.generate_settings(meta["genre"], meta["title"], meta.get("concept", ""))
                self._log("世界观设定已保存到 memory/settings.json")
            except Exception as e:
                self._log(f"生成失败: {e}")
            finally:
                self._gen_running = False
        
        threading.Thread(target=run, daemon=True).start()
    
    def _gen_characters(self):
        """生成角色"""
        if not self._check_ready():
            return
        if hasattr(self, '_gen_running') and self._gen_running:
            self._log("生成任务正在进行中，请勿重复点击")
            return
        
        self._gen_running = True
        
        def run():
            try:
                meta = self._get_meta()
                self.agent.generate_characters(meta["genre"], meta["title"])
                self._log("角色档案已保存到 memory/characters.json")
            except Exception as e:
                self._log(f"生成失败: {e}")
            finally:
                self._gen_running = False
        
        threading.Thread(target=run, daemon=True).start()
    
    def _gen_outline(self):
        """生成大纲"""
        if not self._check_ready():
            return
        if hasattr(self, '_gen_running') and self._gen_running:
            self._log("生成任务正在进行中，请勿重复点击")
            return
        
        self._gen_running = True
        
        def run():
            try:
                meta = self._get_meta()
                new_outline = self.agent.generate_outline(
                    meta["genre"], meta["title"], meta["chapter_count"]
                )
                
                # 使用锁保护共享状态
                with self._state_lock:
                    self.outline = new_outline
                
                # 保存大纲
                with open(self.current_novel_dir / "outline.json", 'w', encoding='utf-8') as f:
                    json.dump(new_outline, f, indent=2, ensure_ascii=False)
                
                # GUI操作在主线程
                self.root.after(0, self._refresh_outline_list)
                self._log("大纲已保存到 outline.json")
            except Exception as e:
                self._log(f"生成失败: {e}")
            finally:
                self._gen_running = False
        
        threading.Thread(target=run, daemon=True).start()
    
    def _gen_chapter(self):
        """生成下一章"""
        if not self._check_ready():
            return
        if not self.outline:
            messagebox.showwarning("提示", "请先生成大纲")
            return
        
        self.current_chapter += 1
        if self.current_chapter > len(self.outline):
            messagebox.showinfo("提示", "所有章节已生成完毕")
            self.current_chapter = len(self.outline)
            return
        
        chapter_info = self.outline[self.current_chapter - 1]

        def run():
            try:
                ch_num = self.current_chapter  # 捕获当前值，避免竞态
                meta = self._get_meta()
                content = self.agent.generate_chapter(
                    ch_num,
                    chapter_info.get("title", f"第{ch_num}章"),
                    chapter_info.get("summary", ""),
                    word_count=meta.get("word_count_per_chapter", 3000)
                )
                
                # 保存章节
                chapters_dir = self.current_novel_dir / "chapters"
                chapters_dir.mkdir(exist_ok=True)
                with open(chapters_dir / f"chapter_{ch_num:04d}.txt", 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # GUI操作必须在主线程
                self.root.after(0, lambda: self._display_chapter(
                    ch_num, chapter_info.get("title", ""), content))
                
                # 自动定稿
                self.agent.finalize_chapter(ch_num, content)
                
                # 名场面检测
                if self.config.get("auto_detect_scene", True):
                    self._detect_and_prompt_image(content, ch_num)
                
                self._log(f"第{ch_num}章已保存并定稿")
            except Exception as e:
                self._log(f"生成失败: {e}")
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        
        threading.Thread(target=run, daemon=True).start()
    
    def _review_chapter(self):
        """审校当前章节"""
        if not self._check_ready():
            return
        
        content = self.content_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "没有可审校的内容")
            return
        
        with self._state_lock:
            ch_num = self.current_chapter

        def run():
            try:
                review = self.agent.review_chapter(ch_num, content)
                review_json = json.dumps(review, indent=2, ensure_ascii=False)
                self.root.after(0, lambda: self._display_review(review_json))
                self._log(f"审校完成，评分：{review.get('overall_score', 'N/A')}")
            except Exception as e:
                self._log(f"审校失败: {e}")
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        
        threading.Thread(target=run, daemon=True).start()
    
    def _style_optimize(self):
        """风格优化当前章节"""
        if not self._check_ready():
            return
        
        content = self.content_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "没有可优化的内容")
            return
        
        def run():
            try:
                self._log("正在进行风格优化...")
                
                settings = self.memory.get_settings() if self.memory else {}
                meta = self._get_meta()
                
                system = f"""你是专业小说风格优化师。小说类型：{meta.get('genre', '未知')}
世界观：{json.dumps(settings, ensure_ascii=False)[:300] if settings else '未知'}

优化要求：
1. 保持原文情节和结构
2. 增强文学性和可读性
3. 优化句式和用词
4. 增加细节描写（适当）
5. 确保风格一致
6. 直接输出优化后的内容"""
                
                prompt = f"请优化以下章节内容：\n\n{content[:3000]}"
                
                result = self.ai.chat([{"role": "user", "content": prompt}], system=system, max_tokens=4000)
                
                self.root.after(0, lambda: self._display_optimized(result))
                self._log("风格优化完成")
                
            except Exception as e:
                self._log(f"风格优化失败: {e}")
        
        threading.Thread(target=run, daemon=True).start()
    
    def _display_optimized(self, content):
        """显示优化后的内容"""
        self.content_text.delete("1.0", tk.END)
        self.content_text.insert("1.0", content)
        self.word_count_var.set(f"字数: {len(content)}")
        self._log("优化内容已加载到编辑器，请审阅后保存")
    
    def _style_imitation(self):
        """仿写风格 - 从文件夹导入多个作者的作品进行风格模仿"""
        if not self._check_ready():
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("仿写风格 - 模仿作者写作风格")
        dialog.geometry("750x650")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="仿写风格", font=('微软雅黑', 14, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        # 已导入的作者风格列表
        style_frame = tk.LabelFrame(dialog, text=" 已导入的作者风格 ", bg=C['bg_dark'], fg=C['accent_light'],
                                    font=('微软雅黑', 10))
        style_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        
        style_list = tk.Listbox(style_frame, bg=C['bg_card'], fg=C['text_primary'],
                               font=('微软雅黑', 10), height=6, selectmode=tk.EXTENDED)
        style_list.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 存储风格数据
        self._imported_styles = getattr(self, '_imported_styles', [])
        for style in self._imported_styles:
            style_list.insert(tk.END, f"{style.get('author', '未知')} - {style.get('unique_features', '')[:30]}...")
        
        # 操作按钮
        btn_frame = tk.Frame(style_frame, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        def import_author_folder():
            """导入作者作品文件夹"""
            folder = filedialog.askdirectory(title="选择包含作者作品的文件夹")
            if not folder:
                return
            
            folder = Path(folder)
            texts = []
            
            # 读取文件夹中的所有文本文件
            for f in folder.glob("*"):
                if f.suffix.lower() in ['.txt', '.md']:
                    try:
                        texts.append((f.stem, f.read_text(encoding='utf-8')))
                    except Exception as e:
                        logger.debug(f"读取文件失败 {f.name}: {e}")
                elif f.suffix.lower() == '.docx':
                    try:
                        from docx import Document
                        doc = Document(str(f))
                        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                        texts.append((f.stem, text))
                    except ImportError:
                        messagebox.showwarning("提示", "需要安装 python-docx 才能读取 Word 文档")
                        return
            
            if not texts:
                messagebox.showwarning("提示", "文件夹中没有找到可读取的文本文件")
                return
            
            # 分析每个文件的风格
            def analyze_all():
                for name, text in texts:
                    if len(text) > 100:  # 至少100字才分析
                        style = self.agent.analyze_style(text[:3000], name)
                        self._imported_styles.append(style)
                        self.root.after(0, lambda s=style: style_list.insert(tk.END, 
                            f"{s.get('author', '未知')} - {s.get('unique_features', '')[:30]}..."))
                
                self.root.after(0, lambda: messagebox.showinfo("完成", f"已导入 {len(texts)} 个作者的风格"))
                self._log(f"已导入 {len(texts)} 个作者的风格特征")
            
            self._log("正在分析作者风格...")
            threading.Thread(target=analyze_all, daemon=True).start()
        
        def import_single_file():
            """导入单个文件分析风格"""
            file_path = filedialog.askopenfilename(
                title="选择作者作品",
                filetypes=[("文本文件", "*.txt *.md"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
            )
            if not file_path:
                return
            
            file_path = Path(file_path)
            text = ""
            
            try:
                if file_path.suffix.lower() in ['.txt', '.md']:
                    text = file_path.read_text(encoding='utf-8')
                elif file_path.suffix.lower() == '.docx':
                    from docx import Document
                    doc = Document(str(file_path))
                    text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            except Exception as e:
                messagebox.showerror("错误", f"读取文件失败: {e}")
                return
            
            if len(text) < 100:
                messagebox.showwarning("提示", "文本内容太少，无法准确分析风格")
                return
            
            author_name = tk.simpledialog.askstring("作者名称", "请输入作者名称:", initialvalue=file_path.stem)
            if not author_name:
                return
            
            def analyze():
                style = self.agent.analyze_style(text[:3000], author_name)
                self._imported_styles.append(style)
                self.root.after(0, lambda: style_list.insert(tk.END, 
                    f"{style.get('author', '未知')} - {style.get('unique_features', '')[:30]}..."))
                self.root.after(0, lambda: messagebox.showinfo("完成", f"已分析 {author_name} 的风格"))
                self._log(f"已分析 {author_name} 的写作风格")
            
            self._log(f"正在分析 {author_name} 的风格...")
            threading.Thread(target=analyze, daemon=True).start()
        
        def remove_selected():
            """移除选中的风格"""
            selected = style_list.curselection()
            for idx in reversed(selected):
                style_list.delete(idx)
                if idx < len(self._imported_styles):
                    self._imported_styles.pop(idx)
        
        tk.Button(btn_frame, text="导入文件夹", command=import_author_folder,
                 bg=C['accent'], fg='white', font=('微软雅黑', 9), padx=10).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="导入单个文件", command=import_single_file,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 9), padx=10).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="移除选中", command=remove_selected,
                 bg=C['error'], fg='white', font=('微软雅黑', 9), padx=10).pack(side=tk.RIGHT, padx=3)
        
        # 创作设置
        write_frame = tk.LabelFrame(dialog, text=" 仿写创作 ", bg=C['bg_dark'], fg=C['accent_light'],
                                   font=('微软雅黑', 10))
        write_frame.pack(fill=tk.X, padx=15, pady=5)
        
        # 创作提示
        tk.Label(write_frame, text="创作提示:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=10, pady=(5, 3))
        prompt_text = tk.Text(write_frame, wrap=tk.WORD, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'], height=3)
        prompt_text.pack(fill=tk.X, padx=10, pady=3)
        prompt_text.insert("1.0", "请用模仿的风格写一段关于...")
        
        # 字数设置
        params_frame = tk.Frame(write_frame, bg=C['bg_dark'])
        params_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Label(params_frame, text="字数:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        word_count_var = tk.StringVar(value="3000")
        ttk.Combobox(params_frame, textvariable=word_count_var, 
                    values=["500", "1000", "2000", "3000", "5000"], width=8).pack(side=tk.LEFT, padx=5)
        
        # 模式选择
        mode_var = tk.StringVar(value="single")
        tk.Radiobutton(params_frame, text="单一风格", variable=mode_var, value="single",
                       bg=C['bg_dark'], fg=C['text_primary'], selectcolor=C['bg_dark']).pack(side=tk.LEFT, padx=10)
        tk.Radiobutton(params_frame, text="融合风格", variable=mode_var, value="blend",
                       bg=C['bg_dark'], fg=C['text_primary'], selectcolor=C['bg_dark']).pack(side=tk.LEFT, padx=10)
        
        def start_imitation():
            """开始仿写"""
            if not self._imported_styles:
                messagebox.showwarning("提示", "请先导入至少一个作者的风格")
                return
            
            prompt = prompt_text.get("1.0", tk.END).strip()
            if not prompt:
                messagebox.showwarning("提示", "请输入创作提示")
                return
            
            word_count = int(word_count_var.get())
            mode = mode_var.get()
            selected = style_list.curselection()
            
            if mode == "single":
                if not selected:
                    messagebox.showwarning("提示", "请选择一个作者风格")
                    return
                style = self._imported_styles[selected[0]]
                
                def generate():
                    try:
                        self._log(f"正在使用 {style.get('author', '')} 的风格创作...")
                        result = self.agent.generate_with_style(prompt, style, word_count)
                        self.root.after(0, lambda: self._display_generated(result))
                        self._log("仿写完成")
                    except Exception as e:
                        self._log(f"仿写失败: {e}")
                
                threading.Thread(target=generate, daemon=True).start()
            else:
                # 融合模式
                if len(selected) < 2:
                    messagebox.showwarning("提示", "融合模式需要选择至少2个风格")
                    return
                styles = [self._imported_styles[i] for i in selected]
                
                def generate():
                    try:
                        self._log("正在融合多个风格创作...")
                        result = self.agent.blend_styles(styles, prompt, word_count)
                        self.root.after(0, lambda: self._display_generated(result))
                        self._log("风格融合创作完成")
                    except Exception as e:
                        self._log(f"风格融合失败: {e}")
                
                threading.Thread(target=generate, daemon=True).start()
        
        def apply_to_chapter():
            """将仿写结果应用到当前章节"""
            if not self._imported_styles:
                messagebox.showwarning("提示", "请先导入风格")
                return
            
            selected = style_list.curselection()
            if not selected:
                messagebox.showwarning("提示", "请选择一个风格")
                return
            
            style = self._imported_styles[selected[0]]
            current_content = self.content_text.get("1.0", tk.END).strip()
            
            if not current_content:
                messagebox.showwarning("提示", "当前章节没有内容")
                return
            
            def rewrite():
                try:
                    self._log(f"正在用 {style.get('author', '')} 的风格重写当前章节...")
                    prompt = f"请用以下风格重写这段内容：\n\n{current_content[:2000]}"
                    result = self.agent.generate_with_style(prompt, style, len(current_content))
                    self.root.after(0, lambda: self._display_generated(result))
                    self._log("风格重写完成")
                except Exception as e:
                    self._log(f"风格重写失败: {e}")
            
            threading.Thread(target=rewrite, daemon=True).start()
        
        action_frame = tk.Frame(write_frame, bg=C['bg_dark'])
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        tk.Button(action_frame, text="开始仿写", command=start_imitation,
                 bg=C['success'], fg='white', font=('微软雅黑', 11, 'bold'), padx=20, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(action_frame, text="用选中风格重写当前章节", command=apply_to_chapter,
                 bg=C['warning'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
    
    def _display_generated(self, content):
        """显示生成的内容"""
        dialog = tk.Toplevel(self.root)
        dialog.title("仿写结果")
        dialog.geometry("700x500")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="仿写结果", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(10, 5))
        
        result_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 11),
                             bg=C['bg_card'], fg=C['text_primary'], padx=15, pady=15)
        result_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        result_text.insert("1.0", content)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def insert_to_editor():
            self.content_text.insert(tk.INSERT, "\n\n" + content)
            dialog.destroy()
            self._log("仿写内容已插入编辑器")
        
        def replace_editor():
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", content)
            self.word_count_var.set(f"字数: {len(content)}")
            dialog.destroy()
            self._log("仿写内容已替换编辑器内容")
        
        tk.Button(btn_frame, text="插入到编辑器", command=insert_to_editor,
                 bg=C['accent'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="替换编辑器内容", command=replace_editor,
                 bg=C['warning'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=dialog.destroy,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10), padx=15).pack(side=tk.RIGHT, padx=5)
    
    def _generate_character_biography(self, char_name: str = None):
        """生成角色个人传记"""
        if not self._check_ready():
            return
        
        # 如果没有指定角色，让用户选择
        if not char_name:
            characters = self.memory.get_characters() if self.memory else {}
            if not characters:
                messagebox.showwarning("提示", "请先创建角色")
                return
            
            # 角色选择对话框
            char_name = tk.simpledialog.askstring("选择角色", 
                f"请输入角色名称:\n可用角色: {', '.join(characters.keys())}")
            if not char_name:
                return
        
        # 获取字数设置
        default_words = self.config.get("biography_word_count", 100000)
        
        # 字数选择对话框
        word_dialog = tk.Toplevel(self.root)
        word_dialog.title("角色传记设置")
        word_dialog.geometry("400x250")
        word_dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(word_dialog, text=f"生成「{char_name}」个人传记", 
                font=('微软雅黑', 12, 'bold'), bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        tk.Label(word_dialog, text="传记字数:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=30)
        word_var = tk.StringVar(value=str(default_words))
        word_combo = ttk.Combobox(word_dialog, textvariable=word_var,
                                 values=["5000", "10000", "30000", "50000", "100000", "200000"],
                                 width=15)
        word_combo.pack(anchor=tk.W, padx=30, pady=5)
        
        include_mental = tk.BooleanVar(value=True)
        tk.Checkbutton(word_dialog, text="包含心理历程", variable=include_mental,
                      bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=30, pady=3)
        
        include_contrast = tk.BooleanVar(value=True)
        tk.Checkbutton(word_dialog, text="分析性格反差", variable=include_contrast,
                      bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=30, pady=3)
        
        def start_generate():
            word_count = int(word_var.get())
            word_dialog.destroy()
            
            def run():
                try:
                    self._log(f"正在生成「{char_name}」的个人传记（约{word_count}字）...")
                    
                    # 获取角色信息
                    characters = self.memory.get_characters()
                    char_info = characters.get(char_name, {})
                    
                    # 获取世界观和大纲
                    settings = self.memory.get_settings() if self.memory else {}
                    outline = self.outline if self.outline else []
                    
                    system = f"""你是专业的小说传记作家。请为角色「{char_name}」撰写一部完整的个人传记。

角色信息：{json.dumps(char_info, ensure_ascii=False)[:1000]}

世界观：{json.dumps(settings, ensure_ascii=False)[:500]}

传记要求：
1. 从角色的出生/起源开始写起
2. 详细描述角色的成长历程
3. 包含角色的心理变化过程
4. 分析角色的性格特点和反差
5. 描述角色的重要经历和转折点
6. 在结尾总结：
   - 这个角色是什么样的人
   - 他的核心性格特征
   - 他的心理发展过程
   - 他身上的反差和矛盾
   - 他对故事的意义

{"请重点描写角色的心理历程。" if include_mental.get() else ""}
{"请分析角色性格中的反差和矛盾。" if include_contrast.get() else ""}

字数要求：约{word_count}字"""
                    
                    prompt = f"请为「{char_name}」撰写个人传记。大纲参考：{json.dumps(outline[:5], ensure_ascii=False)}"
                    
                    result = self.ai.chat([{"role": "user", "content": prompt}], 
                                         system=system, max_tokens=word_count * 2)
                    
                    # 保存传记
                    bio_dir = self.current_novel_dir / "biographies"
                    bio_dir.mkdir(exist_ok=True)
                    bio_file = bio_dir / f"{char_name}_传记.txt"
                    bio_file.write_text(result, encoding='utf-8')
                    
                    # 同步到角色面板
                    if self.memory:
                        characters = self.memory.get_characters()
                        if char_name in characters:
                            characters[char_name]['biography'] = result[:500] + "..."
                            characters[char_name]['biography_file'] = str(bio_file)
                            self.memory.save_characters(characters)
                    
                    self.root.after(0, lambda: self._show_biography(result, char_name))
                    self._log(f"「{char_name}」个人传记生成完成，已保存到 {bio_file}")
                    
                except Exception as e:
                    self._log(f"传记生成失败: {e}")
            
            threading.Thread(target=run, daemon=True).start()
        
        tk.Button(word_dialog, text="开始生成", command=start_generate,
                 bg=C['accent'], fg='white', font=('微软雅黑', 11, 'bold'), padx=20, pady=5).pack(pady=15)
    
    def _show_biography(self, content: str, char_name: str):
        """显示角色传记"""
        dialog = tk.Toplevel(self.root)
        dialog.title(f"角色传记 - {char_name}")
        dialog.geometry("800x600")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"「{char_name}」个人传记", 
                font=('微软雅黑', 14, 'bold'), bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(10, 5))
        
        bio_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 11),
                          bg=C['bg_card'], fg=C['text_primary'], padx=20, pady=15)
        bio_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        bio_text.insert("1.0", content)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def insert_to_chapter():
            self.content_text.insert(tk.INSERT, "\n\n" + content)
            dialog.destroy()
        
        tk.Button(btn_frame, text="插入到当前章节", command=insert_to_chapter,
                 bg=C['accent'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=dialog.destroy,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10), padx=15).pack(side=tk.RIGHT, padx=5)
    
    def _generate_synopsis(self):
        """AI生成书籍简介"""
        if not self._check_ready():
            return
        
        def run():
            try:
                self._log("正在生成书籍简介...")
                
                meta = self._get_meta()
                settings = self.memory.get_settings() if self.memory else {}
                outline = self.outline if self.outline else []
                
                system = """你是专业的书籍简介撰写专家。根据小说的世界观、大纲和设定，撰写一段吸引读者的书籍简介。

简介要求：
1. 150-300字
2. 突出故事亮点和卖点
3. 设置悬念，吸引读者
4. 不要剧透关键情节
5. 语言精炼有力"""
                
                prompt = f"""小说信息：
标题：{meta.get('title', '未命名')}
类型：{meta.get('genre', '未知')}
概念：{meta.get('concept', '无')}

世界观：{json.dumps(settings, ensure_ascii=False)[:500]}

大纲概要：{json.dumps(outline[:10], ensure_ascii=False)[:500]}

请生成书籍简介。"""
                
                result = self.ai.chat([{"role": "user", "content": prompt}], system=system, max_tokens=1000)
                
                # 保存简介
                synopsis_file = self.current_novel_dir / "synopsis.txt"
                synopsis_file.write_text(result, encoding='utf-8')
                
                self.root.after(0, lambda: self._show_synopsis(result))
                self._log("书籍简介生成完成")
                
            except Exception as e:
                self._log(f"简介生成失败: {e}")
        
        threading.Thread(target=run, daemon=True).start()
    
    def _show_synopsis(self, content: str):
        """显示书籍简介"""
        dialog = tk.Toplevel(self.root)
        dialog.title("书籍简介")
        dialog.geometry("500x400")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="书籍简介", font=('微软雅黑', 14, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        synopsis_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 12),
                               bg=C['bg_card'], fg=C['text_primary'], padx=20, pady=15)
        synopsis_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        synopsis_text.insert("1.0", content)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def save():
            new_content = synopsis_text.get("1.0", tk.END).strip()
            synopsis_file = self.current_novel_dir / "synopsis.txt"
            synopsis_file.write_text(new_content, encoding='utf-8')
            dialog.destroy()
            self._log("书籍简介已保存")
        
        tk.Button(btn_frame, text="保存", command=save,
                 bg=C['success'], fg='white', font=('微软雅黑', 10), padx=20).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=dialog.destroy,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10), padx=15).pack(side=tk.RIGHT, padx=5)
    
    def _import_document(self):
        """导入作者文档（txt/docx/md）"""
        file_path = filedialog.askopenfilename(
            title="选择要导入的文档",
            filetypes=[
                ("所有支持格式", "*.txt *.docx *.md"),
                ("TXT文件", "*.txt"),
                ("Word文档", "*.docx"),
                ("Markdown文件", "*.md"),
                ("所有文件", "*.*"),
            ]
        )
        
        if not file_path:
            return
        
        try:
            file_path = Path(file_path)
            content = ""
            
            if file_path.suffix.lower() == '.txt' or file_path.suffix.lower() == '.md':
                content = file_path.read_text(encoding='utf-8')
            elif file_path.suffix.lower() == '.docx':
                try:
                    from docx import Document
                    doc = Document(str(file_path))
                    content = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                except ImportError:
                    messagebox.showerror("错误", "需要安装 python-docx 库才能导入 Word 文档\n请运行: pip install python-docx")
                    return
            else:
                messagebox.showerror("错误", f"不支持的文件格式: {file_path.suffix}")
                return
            
            if not content.strip():
                messagebox.showwarning("提示", "文档内容为空")
                return
            
            self._imported_content = content
            self._imported_file = file_path.name
            
            self._show_import_preview(content, file_path.name)
            
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")
    
    def _show_import_preview(self, content, filename):
        """显示导入内容预览"""
        dialog = tk.Toplevel(self.root)
        dialog.title(f"导入预览 - {filename}")
        dialog.geometry("600x500")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"已导入: {filename}", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 5))
        
        tk.Label(dialog, text=f"字数: {len(content)}", font=('微软雅黑', 10),
                bg=C['bg_dark'], fg=C['text_secondary']).pack(pady=(0, 10))
        
        preview_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 10),
                              bg=C['bg_card'], fg=C['text_primary'], height=15)
        preview_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        preview_text.insert("1.0", content[:2000] + ("..." if len(content) > 2000 else ""))
        preview_text.config(state=tk.DISABLED)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def insert_direct():
            self.content_text.insert(tk.INSERT, content)
            dialog.destroy()
            self._log(f"已插入导入内容: {filename}")
        
        def ai_analyze():
            dialog.destroy()
            self._ai_analyze_content()
        
        tk.Button(btn_frame, text="直接插入", command=insert_direct,
                 bg=C['accent'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="AI分析建议", command=ai_analyze,
                 bg=C['success'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=dialog.destroy,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10), padx=15).pack(side=tk.RIGHT, padx=5)
    
    def _ai_analyze_content(self):
        """AI分析导入的内容并给出建议"""
        if not hasattr(self, '_imported_content') or not self._imported_content:
            messagebox.showwarning("提示", "请先导入文档（创作流程 → 导入文档）")
            return
        
        if not self._check_ready():
            return
        
        content = self._imported_content
        filename = self._imported_file
        
        def run():
            try:
                self._log(f"正在AI分析导入内容: {filename}...")
                
                meta = self._get_meta()
                settings = self.memory.get_settings() if self.memory else {}
                outline = self.outline if self.outline else []
                
                system = """你是专业小说顾问。分析导入的文档内容，给出详细的创作建议。

分析维度：
1. 内容摘要（200字以内）
2. 写作风格分析
3. 优点与亮点
4. 需要改进的地方
5. 与当前小说的关联建议
6. 下一步创作建议（具体可执行的3-5条建议）

输出格式要求：
- 使用清晰的标题和分段
- 建议要具体可执行
- 语言要专业但易懂"""
                
                context = f"当前小说：《{meta.get('title', '未命名')}》\n类型：{meta.get('genre', '未知')}\n"
                if outline:
                    context += f"大纲章节数：{len(outline)}\n"
                
                prompt = f"""{context}

导入的文档内容（文件名：{filename}）：
---
{content[:3000]}
---

请分析以上内容并给出创作建议。"""
                
                result = self.ai.chat([{"role": "user", "content": prompt}], system=system, max_tokens=3000)
                
                self.root.after(0, lambda: self._display_analysis_result(result, filename))
                self._log("AI分析完成")
                
            except Exception as e:
                self._log(f"AI分析失败: {e}")
        
        threading.Thread(target=run, daemon=True).start()
    
    def _display_analysis_result(self, analysis, filename):
        """显示AI分析结果"""
        dialog = tk.Toplevel(self.root)
        dialog.title(f"AI分析报告 - {filename}")
        dialog.geometry("700x600")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="AI分析报告", font=('微软雅黑', 14, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        result_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 11),
                             bg=C['bg_card'], fg=C['text_primary'], padx=20, pady=15)
        result_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        result_text.insert("1.0", analysis)
        
        scrollbar = tk.Scrollbar(result_text, command=result_text.yview)
        result_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def insert_content():
            if hasattr(self, '_imported_content'):
                self.content_text.insert(tk.INSERT, self._imported_content)
                dialog.destroy()
                self._log(f"已插入导入内容")
        
        def use_as_reference():
            if self.note_manager:
                self.note_manager.add_note("AI分析参考", analysis)
                self._log("分析结果已保存到笔记")
            dialog.destroy()
        
        tk.Button(btn_frame, text="插入原文", command=insert_content,
                 bg=C['accent'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="保存为参考", command=use_as_reference,
                 bg=C['success'], fg='white', font=('微软雅黑', 10), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=dialog.destroy,
                 bg=C['bg_light'], fg=C['text_primary'], font=('微软雅黑', 10), padx=15).pack(side=tk.RIGHT, padx=5)
    
    def _stop_generate(self):
        """停止自动创作"""
        self._auto_running = False
        self._log("已请求停止自动创作，正在完成当前章节...")
    
    def _auto_generate(self):
        """自动创作全流程"""
        # 防重复点击
        if hasattr(self, '_auto_running') and self._auto_running:
            self._log("自动创作正在进行中，请勿重复点击")
            return
        
        if not self._check_ready():
            return
        
        self._auto_running = True
        self.auto_btn.config(state=tk.DISABLED, text="创作中...")
        
        meta = self._get_meta()
        
        def run():
            try:
                self._log("=== 开始自动创作 ===")
                
                # 1. 生成世界观（如果不存在）
                settings_file = self.current_novel_dir / "memory" / "settings.json"
                if not settings_file.exists():
                    try:
                        self.agent.generate_settings(meta["genre"], meta["title"], meta.get("concept", ""))
                    except Exception as e:
                        self._log(f"世界观生成失败，跳过: {e}")
                else:
                    self._log("世界观已存在，跳过生成")

                # 2. 生成角色（如果不存在）
                characters_dir = self.current_novel_dir / "characters"
                if not characters_dir.exists() or not list(characters_dir.glob("*.json")):
                    try:
                        self.agent.generate_characters(meta["genre"], meta["title"])
                    except Exception as e:
                        self._log(f"角色生成失败，跳过: {e}")
                else:
                    self._log("角色已存在，跳过生成")

                # 3. 生成大纲（如果不存在）
                outline_file = self.current_novel_dir / "outline.json"
                if not outline_file.exists() or not self.outline:
                    try:
                        with self._state_lock:
                            self.outline = self.agent.generate_outline(meta["genre"], meta["title"], meta["chapter_count"])
                        with open(outline_file, 'w', encoding='utf-8') as f:
                            with self._state_lock:
                                json.dump(self.outline, f, indent=2, ensure_ascii=False)
                        self.root.after(0, self._refresh_outline_list)
                    except Exception as e:
                        self._log(f"大纲生成失败: {e}")
                        self._auto_running = False
                        self.root.after(0, lambda: self.auto_btn.config(state=tk.NORMAL, text="自动创作"))
                        return
                else:
                    self._log("大纲已存在，跳过生成")

                # 4. 逐章生成（从已完成的下一章开始）
                chapters_dir = self.current_novel_dir / "chapters"
                chapters_dir.mkdir(exist_ok=True)
                existing_chapters = set()
                for f in chapters_dir.glob("chapter_*.txt"):
                    try:
                        num = int(f.stem.split('_')[-1])
                        existing_chapters.add(num)
                    except ValueError:
                        pass
                
                with self._state_lock:
                    outline_snapshot = list(self.outline)
                
                total = len(outline_snapshot)
                skipped = 0
                generated = 0
                failed = 0
                
                for i, chapter_info in enumerate(outline_snapshot):
                    ch_num = i + 1
                    
                    # 检查是否请求停止
                    if not self._auto_running:
                        self._log("自动创作已停止")
                        break
                    
                    # 跳过已完成的章节
                    if ch_num in existing_chapters:
                        skipped += 1
                        continue
                    
                    with self._state_lock:
                        self.current_chapter = ch_num
                    
                    try:
                        content = self.agent.generate_chapter(
                            ch_num,
                            chapter_info.get("title", f"第{ch_num}章"),
                            chapter_info.get("summary", ""),
                            word_count=meta.get("word_count_per_chapter", 3000)
                        )

                        # 保存
                        with open(chapters_dir / f"chapter_{ch_num:04d}.txt", 'w', encoding='utf-8') as f:
                            f.write(content)

                        # 定稿
                        self.agent.finalize_chapter(ch_num, content)

                        # 更新UI
                        self.root.after(0, lambda c=content, n=ch_num, t=chapter_info.get("title", ""): self._display_chapter(n, t, c))
                        self._log(f"第{ch_num}章创作完成")
                        generated += 1
                    except Exception as e:
                        self._log(f"第{ch_num}章创作失败，继续下一章: {e}")
                        failed += 1
                        continue
                
                # 更新进度显示
                self.root.after(0, lambda: self.chapter_var.set(f"{len(existing_chapters) + generated}/{total}"))
                
                summary = f"=== 自动创作完成 ===\n"
                summary += f"总章数: {total}\n"
                if skipped > 0:
                    summary += f"跳过已完成: {skipped} 章\n"
                summary += f"本次生成: {generated} 章\n"
                if failed > 0:
                    summary += f"失败: {failed} 章\n"
                self._log(summary)
                
                if generated > 0:
                    self.root.after(0, lambda: messagebox.showinfo("完成", f"《{meta['title']}》创作完成！\n本次生成 {generated} 章"))
                elif skipped > 0:
                    self.root.after(0, lambda: messagebox.showinfo("提示", "所有章节已完成，无需继续创作"))
                    
            except Exception as e:
                self._log(f"自动创作失败: {e}")
            finally:
                self._auto_running = False
                self.root.after(0, lambda: self.auto_btn.config(state=tk.NORMAL, text="自动创作"))
        
        threading.Thread(target=run, daemon=True).start()
    
    def _display_chapter(self, num, title, content):
        """显示章节内容（线程安全）"""
        self.content_text.delete("1.0", tk.END)
        self.content_text.insert("1.0", content)
        self.chapter_title_var.set(f"第{num}章: {title}")
        self.word_count_var.set(f"字数: {len(content)}")
        meta = self._get_meta()
        self.chapter_var.set(f"{self.current_chapter}/{meta.get('chapter_count', '?')}")
    
    def _save_chapter(self):
        """保存当前章节"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先创建或打开小说")
            return
        
        content = self.content_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "没有可保存的内容")
            return
        
        chapters_dir = self.current_novel_dir / "chapters"
        chapters_dir.mkdir(exist_ok=True)
        with open(chapters_dir / f"chapter_{self.current_chapter:04d}.txt", 'w', encoding='utf-8') as f:
            f.write(content)
        
        self._log(f"第{self.current_chapter}章已保存")
        messagebox.showinfo("成功", "章节已保存")
    
    def _prev_chapter(self):
        """加载上一章"""
        if not self.outline or not self.current_novel_dir:
            return
        if self.current_chapter <= 1:
            self._log("已经是第一章")
            return
        
        # 保存当前章节
        self._save_chapter_silent()
        
        # 切换到上一章
        self.current_chapter -= 1
        self._load_chapter_by_number(self.current_chapter)
    
    def _next_chapter(self):
        """加载下一章"""
        if not self.outline or not self.current_novel_dir:
            return
        if self.current_chapter >= len(self.outline):
            self._log("已经是最后一章")
            return
        
        # 保存当前章节
        self._save_chapter_silent()
        
        # 切换到下一章
        self.current_chapter += 1
        self._load_chapter_by_number(self.current_chapter)
    
    def _save_chapter_silent(self):
        """静默保存当前章节（不弹窗）"""
        if not self.current_novel_dir:
            return
        content = self.content_text.get("1.0", tk.END).strip()
        if not content:
            return
        chapters_dir = self.current_novel_dir / "chapters"
        chapters_dir.mkdir(exist_ok=True)
        with open(chapters_dir / f"chapter_{self.current_chapter:04d}.txt", 'w', encoding='utf-8') as f:
            f.write(content)
        self._log(f"第{self.current_chapter}章已自动保存")
    
    def _load_chapter_by_number(self, ch_num):
        """根据章节号加载内容"""
        if not self.outline or ch_num < 1 or ch_num > len(self.outline):
            return
        
        chapter_info = self.outline[ch_num - 1]
        chapters_dir = self.current_novel_dir / "chapters"
        chapter_file = chapters_dir / f"chapter_{ch_num:04d}.txt"
        
        if chapter_file.exists():
            content = chapter_file.read_text(encoding='utf-8')
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", content)
            self.chapter_title_var.set(f"第{ch_num}章: {chapter_info.get('title', '')}")
            self.word_count_var.set(f"字数: {len(content)}")
            self._log(f"已加载第{ch_num}章")
        else:
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", f"章节大纲：\n{chapter_info.get('summary', '无')}\n\n在此处编写内容...")
            self.chapter_title_var.set(f"第{ch_num}章: {chapter_info.get('title', '')} (未生成)")
            self.word_count_var.set("字数: 0")
            self._log(f"第{ch_num}章尚未生成，可手动编写")
        
        # 更新进度显示
        self.chapter_var.set(f"{ch_num}/{len(self.outline)}")
    
    def _update_chapter_selector(self):
        """更新章节选择器"""
        if not self.outline:
            self.chapter_select['values'] = []
            return
        
        chapters = [f"第{i+1}章" for i in range(len(self.outline))]
        self.chapter_select['values'] = chapters
        
        # 设置当前章节
        if self.current_chapter > 0 and self.current_chapter <= len(chapters):
            self.chapter_select_var.set(chapters[self.current_chapter - 1])
        elif chapters:
            self.chapter_select_var.set(chapters[0])
    
    def _on_chapter_select(self, event):
        """章节选择器回调 - 跳转到指定章节"""
        selection = self.chapter_select_var.get()
        if not selection:
            return
        try:
            ch_num = int(selection.replace("第", "").replace("章", "").strip())
        except ValueError:
            return
        
        self._save_chapter_silent()
        self.current_chapter = ch_num
        self._load_chapter_by_number(ch_num)
        
        # 更新大纲列表选中状态
        self.outline_list.selection_clear(0, tk.END)
        self.outline_list.selection_set(ch_num - 1)
        self.outline_list.see(ch_num - 1)
    
    def _export_txt(self):
        """导出全文TXT"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先创建或打开小说")
            return
        
        meta = self._get_meta()
        save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt")],
            initialfile=f"{meta.get('title', '小说')}.txt"
        )
        
        if not save_path:
            return
        
        chapters_dir = self.current_novel_dir / "chapters"
        chapter_files = sorted(chapters_dir.glob("chapter_*.txt"))
        
        with open(save_path, 'w', encoding='utf-8') as out:
            out.write(f"《{meta.get('title', '小说')}》\n\n")
            for cf in chapter_files:
                content = cf.read_text(encoding='utf-8')
                out.write(content + "\n\n")
        
        self._log(f"全文已导出到: {save_path}")
        messagebox.showinfo("成功", f"全文已导出到:\n{save_path}")
    
    def _refresh_outline_list(self):
        """刷新大纲列表"""
        self.outline_list.delete(0, tk.END)
        outline_type = self.outline_type_var.get()
        
        if outline_type == "章节大纲":
            chapters = []
            for item in self.outline:
                ch = item.get("chapter", "?")
                title = item.get("title", "未命名")
                self.outline_list.insert(tk.END, f"第{ch}章: {title}")
                chapters.append(f"第{ch}章")
            
            # 更新章节选择器
            if hasattr(self, 'chapter_select'):
                self.chapter_select['values'] = chapters
                if chapters:
                    current = f"第{self.current_chapter}章" if self.current_chapter > 0 else chapters[0]
                    self.chapter_select_var.set(current)
        elif outline_type == "整体大纲":
            overall = self._get_overall_outline()
            for i, item in enumerate(overall):
                self.outline_list.insert(tk.END, f"{i+1}. {item.get('title', '未命名')}")
        elif outline_type == "故事大纲":
            stories = self._get_story_outlines()
            for name, story in stories.items():
                self.outline_list.insert(tk.END, f"📖 {name}")
    
    def _on_outline_type_change(self, event=None):
        """大纲类型切换"""
        self._refresh_outline_list()
    
    def _get_overall_outline(self) -> list:
        """获取整体大纲"""
        if not self.current_novel_dir:
            return []
        outline_dir = self.current_novel_dir / "outlines"
        overall_file = outline_dir / "overall.json"
        if overall_file.exists():
            with open(overall_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def _save_overall_outline(self, data: list):
        """保存整体大纲"""
        if not self.current_novel_dir:
            return
        outline_dir = self.current_novel_dir / "outlines"
        outline_dir.mkdir(exist_ok=True)
        with open(outline_dir / "overall.json", 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def _get_story_outlines(self) -> dict:
        """获取所有故事大纲"""
        if not self.current_novel_dir:
            return {}
        outline_dir = self.current_novel_dir / "outlines"
        stories_file = outline_dir / "stories.json"
        if stories_file.exists():
            with open(stories_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _save_story_outlines(self, data: dict):
        """保存故事大纲"""
        if not self.current_novel_dir:
            return
        outline_dir = self.current_novel_dir / "outlines"
        outline_dir.mkdir(exist_ok=True)
        with open(outline_dir / "stories.json", 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def _add_outline_item(self):
        """添加大纲项"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先创建或打开小说")
            return
        
        outline_type = self.outline_type_var.get()
        
        dialog = tk.Toplevel(self.root)
        dialog.title(f"添加{outline_type}")
        dialog.geometry("500x400")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"添加{outline_type}", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(15, 10))
        
        # 标题输入
        title_frame = tk.Frame(dialog, bg=C['bg_dark'])
        title_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(title_frame, text="标题:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        title_entry = tk.Entry(title_frame, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
        title_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 内容输入
        tk.Label(dialog, text="内容:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(10, 3))
        content_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 10),
                              bg=C['bg_card'], fg=C['text_primary'], height=12)
        content_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)
        
        def confirm():
            title = title_entry.get().strip()
            content = content_text.get("1.0", tk.END).strip()
            if not title:
                messagebox.showwarning("提示", "请输入标题")
                return
            
            if outline_type == "整体大纲":
                overall = self._get_overall_outline()
                overall.append({"title": title, "content": content})
                self._save_overall_outline(overall)
            elif outline_type == "故事大纲":
                stories = self._get_story_outlines()
                stories[title] = {"title": title, "content": content, "chapters": []}
                self._save_story_outlines(stories)
            elif outline_type == "章节大纲":
                ch_num = len(self.outline) + 1
                self.outline.append({"chapter": ch_num, "title": title, "summary": content})
                outline_file = self.current_novel_dir / "outline.json"
                with open(outline_file, 'w', encoding='utf-8') as f:
                    json.dump(self.outline, f, indent=2, ensure_ascii=False)
            
            self._refresh_outline_list()
            dialog.destroy()
            self._log(f"已添加{outline_type}: {title}")
        
        tk.Button(dialog, text="确认添加", command=confirm, bg=C['accent'], fg='white',
                 font=('微软雅黑', 10), padx=20, pady=5).pack(pady=10)
    
    def _edit_outline_item(self):
        """编辑大纲项"""
        selection = self.outline_list.curselection()
        if not selection:
            messagebox.showwarning("提示", "请先选择要编辑的大纲项")
            return
        
        idx = selection[0]
        outline_type = self.outline_type_var.get()
        
        # 获取当前内容
        if outline_type == "章节大纲":
            if idx < len(self.outline):
                item = self.outline[idx]
                title = item.get("title", "")
                content = item.get("summary", "")
            else:
                return
        elif outline_type == "整体大纲":
            overall = self._get_overall_outline()
            if idx < len(overall):
                item = overall[idx]
                title = item.get("title", "")
                content = item.get("content", "")
            else:
                return
        elif outline_type == "故事大纲":
            stories = self._get_story_outlines()
            names = list(stories.keys())
            if idx < len(names):
                item = stories[names[idx]]
                title = item.get("title", "")
                content = item.get("content", "")
            else:
                return
        
        dialog = tk.Toplevel(self.root)
        dialog.title(f"编辑{outline_type}")
        dialog.geometry("500x400")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text=f"编辑{outline_type}", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(15, 10))
        
        # 标题输入
        title_frame = tk.Frame(dialog, bg=C['bg_dark'])
        title_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(title_frame, text="标题:", bg=C['bg_dark'], fg=C['text_primary']).pack(side=tk.LEFT)
        title_entry = tk.Entry(title_frame, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
        title_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        title_entry.insert(0, title)
        
        # 内容输入
        tk.Label(dialog, text="内容:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(10, 3))
        content_text = tk.Text(dialog, wrap=tk.WORD, font=('微软雅黑', 10),
                              bg=C['bg_card'], fg=C['text_primary'], height=12)
        content_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)
        content_text.insert("1.0", content)
        
        def confirm():
            new_title = title_entry.get().strip()
            new_content = content_text.get("1.0", tk.END).strip()
            if not new_title:
                messagebox.showwarning("提示", "请输入标题")
                return
            
            if outline_type == "章节大纲":
                self.outline[idx]["title"] = new_title
                self.outline[idx]["summary"] = new_content
                outline_file = self.current_novel_dir / "outline.json"
                with open(outline_file, 'w', encoding='utf-8') as f:
                    json.dump(self.outline, f, indent=2, ensure_ascii=False)
            elif outline_type == "整体大纲":
                overall = self._get_overall_outline()
                overall[idx] = {"title": new_title, "content": new_content}
                self._save_overall_outline(overall)
            elif outline_type == "故事大纲":
                stories = self._get_story_outlines()
                names = list(stories.keys())
                old_name = names[idx]
                del stories[old_name]
                stories[new_title] = {"title": new_title, "content": new_content, "chapters": []}
                self._save_story_outlines(stories)
            
            self._refresh_outline_list()
            dialog.destroy()
            self._log(f"已更新{outline_type}: {new_title}")
        
        tk.Button(dialog, text="确认修改", command=confirm, bg=C['accent'], fg='white',
                 font=('微软雅黑', 10), padx=20, pady=5).pack(pady=10)
    
    def _delete_outline_item(self):
        """删除大纲项"""
        selection = self.outline_list.curselection()
        if not selection:
            messagebox.showwarning("提示", "请先选择要删除的大纲项")
            return
        
        idx = selection[0]
        outline_type = self.outline_type_var.get()
        
        if not messagebox.askyesno("确认", f"确定要删除这个{outline_type}吗？"):
            return
        
        if outline_type == "章节大纲":
            if idx < len(self.outline):
                removed = self.outline.pop(idx)
                # 重新编号
                for i, item in enumerate(self.outline):
                    item["chapter"] = i + 1
                outline_file = self.current_novel_dir / "outline.json"
                with open(outline_file, 'w', encoding='utf-8') as f:
                    json.dump(self.outline, f, indent=2, ensure_ascii=False)
                self._log(f"已删除: {removed.get('title', '')}")
        elif outline_type == "整体大纲":
            overall = self._get_overall_outline()
            if idx < len(overall):
                removed = overall.pop(idx)
                self._save_overall_outline(overall)
                self._log(f"已删除: {removed.get('title', '')}")
        elif outline_type == "故事大纲":
            stories = self._get_story_outlines()
            names = list(stories.keys())
            if idx < len(names):
                removed_name = names[idx]
                del stories[removed_name]
                self._save_story_outlines(stories)
                self._log(f"已删除: {removed_name}")
        
        self._refresh_outline_list()
    
    # ===== 云端同步 =====
    
    def _cloud_sync(self):
        """云端同步对话框"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先新建或打开小说")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("云端同步")
        dialog.geometry("400x300")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="云端同步", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(15, 10))
        
        # 获取可用的云存储
        providers = self.cloud_storage.get_available_providers()
        configured = [p for p in providers if p.get("configured")]
        
        if not configured:
            tk.Label(dialog, text="未配置云存储\n\n请在 设置 → 云端存储 中配置",
                    font=('微软雅黑', 10), bg=C['bg_dark'], fg=C['text_muted']).pack(pady=20)
        else:
            provider_var = tk.StringVar(value=configured[0]["name"])
            for p in configured:
                tk.Radiobutton(dialog, text=p["name"], variable=provider_var, value=p["name"],
                              font=('微软雅黑', 10), bg=C['bg_dark'], fg=C['text_primary'],
                              selectcolor=C['accent']).pack(anchor=tk.W, padx=30, pady=3)
            
            def do_upload():
                provider_name = provider_var.get()
                provider_id = next((p["id"] for p in configured if p["name"] == provider_name), None)
                if provider_id:
                    def run():
                        try:
                            self._log(f"开始上传到 {provider_name}...")
                            success = self.cloud_storage.upload_novel(self.current_novel_dir, provider_id)
                            if success:
                                self._log(f"上传成功！")
                                self.root.after(0, lambda: messagebox.showinfo("成功", f"小说已上传到 {provider_name}"))
                            else:
                                self._log(f"上传失败")
                                self.root.after(0, lambda: messagebox.showerror("失败", f"上传失败，请检查网络和配置"))
                        except Exception as e:
                            self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
                    threading.Thread(target=run, daemon=True).start()
            
            def do_download():
                provider_name = provider_var.get()
                provider_id = next((p["id"] for p in configured if p["name"] == provider_name), None)
                if provider_id:
                    def run():
                        try:
                            self._log(f"开始从 {provider_name} 下载...")
                            success = self.cloud_storage.download_novel("/AI_NovelWriter", self.current_novel_dir, provider_id)
                            if success:
                                self._log(f"下载成功！")
                                self.root.after(0, lambda: messagebox.showinfo("成功", f"小说已从 {provider_name} 下载"))
                            else:
                                self._log(f"下载失败")
                                self.root.after(0, lambda: messagebox.showerror("失败", f"下载失败"))
                        except Exception as e:
                            self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
                    threading.Thread(target=run, daemon=True).start()
            
            btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
            btn_frame.pack(fill=tk.X, padx=30, pady=15)
            
            tk.Button(btn_frame, text="上传到云端", font=('微软雅黑', 10),
                     bg=C['accent'], fg='white', relief=tk.FLAT, padx=15,
                     command=do_upload).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="从云端下载", font=('微软雅黑', 10),
                     bg=C['success'], fg='white', relief=tk.FLAT, padx=15,
                     command=do_download).pack(side=tk.LEFT, padx=5)
        
        tk.Button(dialog, text="关闭", font=('微软雅黑', 10),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=15,
                 command=dialog.destroy).pack(pady=10)
    
    # ===== 格式转换 =====
    
    def _show_format_converter(self):
        """显示格式转换对话框"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先新建或打开小说")
            return
        
        if not self.format_converter:
            self.format_converter = FormatConverter(self.current_novel_dir)
        
        dialog = tk.Toplevel(self.root)
        dialog.title("格式转换")
        dialog.geometry("450x400")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="选择导出格式:", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(15, 10))
        
        # 格式列表
        formats = self.format_converter.get_formats()
        format_var = tk.StringVar(value="html")
        
        for fmt_key, fmt_info in formats.items():
            frame = tk.Frame(dialog, bg=C['bg_dark'])
            frame.pack(fill=tk.X, padx=20, pady=3)
            
            tk.Radiobutton(frame, text=f"{fmt_info['name']} ({fmt_info['ext']})", 
                          variable=format_var, value=fmt_key,
                          font=('微软雅黑', 10), bg=C['bg_dark'], fg=C['text_primary'],
                          selectcolor=C['accent']).pack(side=tk.LEFT)
            tk.Label(frame, text=fmt_info['desc'], font=('微软雅黑', 8),
                    bg=C['bg_dark'], fg=C['text_muted']).pack(side=tk.RIGHT)
        
        # 包含图片选项
        include_images_var = tk.BooleanVar(value=True)
        tk.Checkbutton(dialog, text="包含插图（如有）", variable=include_images_var,
                      font=('微软雅黑', 10), bg=C['bg_dark'], fg=C['text_primary'],
                      selectcolor=C['accent']).pack(pady=10)
        
        def do_convert():
            fmt = format_var.get()
            
            # 收集所有章节
            chapters_dir = self.current_novel_dir / "chapters"
            chapter_files = sorted(chapters_dir.glob("chapter_*.txt"))
            chapters = []
            
            for cf in chapter_files:
                num = int(cf.stem.split("_")[1])
                content = cf.read_text(encoding='utf-8')
                # 从大纲获取标题
                title = f"第{num}章"
                if self.outline and num <= len(self.outline):
                    title = self.outline[num-1].get("title", title)
                chapters.append({"num": num, "title": title, "content": content})
            
            if not chapters:
                # 如果没有章节文件，使用编辑区内容
                content = self.content_text.get("1.0", tk.END).strip()
                if not content:
                    messagebox.showwarning("提示", "没有可导出的内容")
                    return
            else:
                content = "\n\n".join(ch["content"] for ch in chapters)
            
            meta = self._get_meta() if self.current_novel_dir else {}
            
            # 获取图片数据
            images = None
            if include_images_var.get() and self.image_manager:
                images = self.image_manager.get_images_data()
            
            # 转换
            result = self.format_converter.convert(
                content=content,
                title=meta.get("title", "小说"),
                format_type=fmt,
                chapters=chapters if chapters else None,
                metadata=meta,
                images=images,
            )
            
            if result:
                self._log(f"格式转换完成: {result}")
                dialog.destroy()
                
                # 询问是否打开
                if messagebox.askyesno("成功", f"已导出为{formats[fmt]['name']}格式\n\n{result}\n\n是否打开文件？"):
                    import subprocess
                    subprocess.Popen(['start', result], shell=True)
            else:
                messagebox.showerror("错误", "格式转换失败")
        
        tk.Button(dialog, text="开始转换", font=('微软雅黑', 11, 'bold'),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=20, pady=8,
                 command=do_convert).pack(pady=15)
    
    # ===== 图片插入 =====
    
    def _insert_image(self):
        """插入图片到编辑区"""
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先新建或打开小说")
            return
        
        if not self.image_manager:
            self.image_manager = ImageManager(self.current_novel_dir)
        
        # 选择图片文件
        file_path = filedialog.askopenfilename(
            title="选择图片",
            filetypes=[
                ("图片文件", "*.png *.jpg *.jpeg *.gif *.bmp *.webp"),
                ("所有文件", "*.*"),
            ]
        )
        
        if not file_path:
            return
        
        # 导入图片
        img_path = self.image_manager.import_image(file_path)
        if not img_path:
            messagebox.showerror("错误", "导入图片失败")
            return
        
        # 在编辑区插入图片标记
        cursor_pos = self.content_text.index(tk.INSERT)
        img_name = Path(img_path).name
        
        # 插入Markdown格式的图片标记
        marker = f"\n![插图]({img_name})\n"
        self.content_text.insert(tk.INSERT, marker)
        
        # 尝试在Text widget中显示图片预览
        try:
            from PIL import Image, ImageTk
            img = Image.open(img_path)
            # 缩放图片
            max_width = 400
            ratio = max_width / img.width
            new_size = (max_width, int(img.height * ratio))
            img = img.resize(new_size, Image.LANCZOS)
            
            # 转换为Tkinter可用的格式
            photo = ImageTk.PhotoImage(img)
            
            # 先获取当前内容（不含标记）
            current = self.content_text.get("1.0", tk.END)
            current = current.replace(f"\n![插图]({img_name})\n", "")
            
            # 清空并重新插入内容
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", current)
            
            # 在Text widget中插入图片预览
            self.content_text.image_create(tk.INSERT, image=photo)
            self.content_text.insert(tk.INSERT, "\n")
            
            # 保持引用防止被垃圾回收
            if not hasattr(self, '_photo_refs'):
                self._photo_refs = []
            self._photo_refs.append(photo)
            
        except ImportError:
            # 如果没有PIL，只保留文本标记（已插入）
            pass
        except Exception as e:
            self._log(f"图片预览加载失败: {e}")
        
        self._log(f"已插入图片: {img_name}")
        
        # 更新字数
        content = self.content_text.get("1.0", tk.END).strip()
        self.word_count_var.set(str(len(content)))
    
    def _on_outline_select(self, event):
        """大纲选中事件"""
        if not self.current_novel_dir or not self.outline:
            return
        
        selection = self.outline_list.curselection()
        if not selection:
            return
        
        idx = selection[0]
        if idx < len(self.outline):
            chapter_info = self.outline[idx]
            # 使用索引+1作为章节号（与文件命名一致）
            self.current_chapter = idx + 1
            
            # 尝试加载已生成的章节
            chapters_dir = self.current_novel_dir / "chapters"
            chapter_file = chapters_dir / f"chapter_{self.current_chapter:04d}.txt"
            
            if chapter_file.exists():
                content = chapter_file.read_text(encoding='utf-8')
                self.content_text.delete("1.0", tk.END)
                self.content_text.insert("1.0", content)
                self.chapter_title_var.set(f"第{self.current_chapter}章: {chapter_info.get('title', '')}")
                self.word_count_var.set(f"字数: {len(content)}")
                self._log(f"已加载第{self.current_chapter}章，可编辑后按 Ctrl+S 保存")
            else:
                self.chapter_title_var.set(f"第{self.current_chapter}章: {chapter_info.get('title', '')} (未生成)")
                self.content_text.delete("1.0", tk.END)
                self.content_text.insert("1.0", f"章节大纲：\n{chapter_info.get('summary', '无')}\n\n在此处编写内容...")
            
            # 更新进度显示
            total = len(self.outline)
            self.chapter_var.set(f"{self.current_chapter}/{total}")
    
    def _detect_and_prompt_image(self, content: str, chapter_num: int):
        """检测名场面并生成电影级AI提示词"""
        scenes = SceneDetector.detect(content)
        if not scenes:
            return
        
        has_api = self.image_gen.is_configured()
        self._log(f"[名场面检测] 发现 {len(scenes)} 个场景")
        
        scene_type_cn = {
            "battle": "战斗场面", "beauty": "人物特写", "emotion": "情感场景",
            "epic_scene": "震撼场面", "character_closeup": "角色特写",
            "landscape": "风景全景", "confrontation": "对峙场面", "sacrifice": "牺牲时刻"
        }
        
        img_dir = self.current_novel_dir / "scene_prompts"
        img_dir.mkdir(exist_ok=True)
        
        for i, scene in enumerate(scenes):
            type_name = scene_type_cn.get(scene["type"], "名场面")
            prompt_text = scene.get("prompt", "")
            scene_text = scene.get("text", "")[:200]
            aspect_ratio = scene.get("aspect_ratio", "16:9")
            size = scene.get("size", "1024x576")
            shot_type = scene.get("shot_type", "")
            composition = scene.get("composition", "")
            style = scene.get("style", "")
            
            purpose_text = {
                "battle": "增强战斗场面的视觉冲击力",
                "beauty": "展现角色外形特征和精神面貌",
                "emotion": "捕捉情感高潮，增强读者代入",
                "epic_scene": "渲染宏大世界观和场景氛围",
                "character_closeup": "刻画角色细节表情",
                "landscape": "展示世界观和环境氛围",
                "confrontation": "表现角色对峙的张力",
                "sacrifice": "定格感动人心的瞬间",
            }.get(scene["type"], "可视化关键场景")
            
            # 保存提示词
            safe_type = scene["type"]
            prompt_file = img_dir / f"ch{chapter_num:04d}_{safe_type}_{i+1}_prompt.txt"
            prompt_content = (
                f"章节: 第{chapter_num}章\n"
                f"类型: {type_name}\n"
                f"场景: {scene_text}\n\n"
                f"画面比例: {aspect_ratio} ({size})\n"
                f"镜头: {shot_type}\n"
                f"构图: {composition}\n"
                f"质感: {style}\n\n"
                f"AI提示词:\n{prompt_text}"
            )
            prompt_file.write_text(prompt_content, encoding='utf-8')
            
            if has_api:
                self._show_image_prompt_dialog(chapter_num, i, type_name, scene_text, 
                    prompt_text, purpose_text, scene, img_dir)
            else:
                self._log(f"[提示词] 第{chapter_num}章 {type_name} 已保存: scene_prompts/{prompt_file.name}")
    
    def _show_image_prompt_dialog(self, chapter_num, idx, type_name, scene_text, prompt_text, purpose_text, scene, img_dir):
        """显示电影级图片生成提醒对话框"""
        C = UIStyle.COLORS
        dialog = tk.Toplevel(self.root)
        dialog.title(f"名场面插图 - 第{chapter_num}章")
        dialog.geometry("650x520")
        dialog.configure(bg=C['bg_dark'])
        dialog.grab_set()
        
        tk.Label(dialog, text=f"第{chapter_num}章 检测到【{type_name}】", font=('微软雅黑', 14, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 5))
        
        # 电影级参数
        aspect_ratio = scene.get("aspect_ratio", "16:9")
        size = scene.get("size", "1024x576")
        shot_type = scene.get("shot_type", "")
        composition = scene.get("composition", "")
        style = scene.get("style", "")
        
        cinematic_frame = tk.Frame(dialog, bg=C['bg_card'])
        cinematic_frame.pack(fill=tk.X, padx=20, pady=5)
        
        cinematic_info = (
            f"画面比例: {aspect_ratio} ({size})  |  "
            f"镜头: {shot_type[:30]}...\n"
            f"构图: {composition[:30]}...  |  "
            f"质感: {style[:30]}..."
        )
        tk.Label(cinematic_frame, text=cinematic_info, font=('微软雅黑', 9),
                bg=C['bg_card'], fg=C['accent_light'], wraplength=600, justify=tk.LEFT).pack(padx=10, pady=5)
        
        info_frame = tk.Frame(dialog, bg=C['bg_card'])
        info_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(info_frame, text=f"场景: {scene_text}", font=('微软雅黑', 10),
                bg=C['bg_card'], fg=C['text_primary'], wraplength=600, justify=tk.LEFT).pack(padx=10, pady=5)
        
        tk.Label(dialog, text=f"目的: {purpose_text}", font=('微软雅黑', 9),
                bg=C['bg_dark'], fg=C['text_secondary'], wraplength=600).pack(padx=20, pady=3)
        
        timer_var = tk.StringVar(value="30秒后自动生成提示词")
        tk.Label(dialog, textvariable=timer_var, font=('微软雅黑', 9),
                bg=C['bg_dark'], fg=C['warning']).pack(pady=5)
        
        prompt_file = img_dir / f"ch{chapter_num:04d}_{scene['type']}_{idx+1}_prompt.txt"
        
        def countdown(remaining=30):
            if not dialog.winfo_exists():
                return
            if remaining <= 0:
                do_save_prompt()
                return
            timer_var.set(f"{remaining}秒后自动生成提示词")
            dialog.after(1000, lambda: countdown(remaining - 1))
        
        def do_save_prompt():
            dialog.destroy()
            self._log(f"提示词已保存: {prompt_file.name}")
            self.root.after(0, lambda: messagebox.showinfo("已保存", f"插图提示词已保存到:\n{img_dir.name}/{prompt_file.name}"))
        
        def do_generate():
            dialog.destroy()
            def gen_img():
                self._log(f"[文生图] 正在生成: {type_name}...")
                img_data = self.image_gen.generate(
                    prompt=prompt_text,
                    negative_prompt="low quality, blurry, deformed, ugly",
                    width=self.config.get("img_width", 1024),
                    height=self.config.get("img_height", 1024),
                )
                if img_data:
                    filepath = self.image_gen.save_image(img_data, self.current_novel_dir,
                        f"chapter_{chapter_num:04d}_scene_{idx+1}")
                    self._log(f"插图已保存: {filepath}")
                    self.root.after(0, lambda: messagebox.showinfo("成功", f"插图已保存:\n{filepath}"))
            threading.Thread(target=gen_img, daemon=True).start()
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="生成插图", command=do_generate,
                 bg=C['success'], fg='white', font=('微软雅黑', 10, 'bold'), padx=20, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="生成提示词", command=do_save_prompt,
                 bg=C['accent'], fg='white', font=('微软雅黑', 10), padx=20, pady=5).pack(side=tk.LEFT, padx=5)
        
        countdown()
    
    def _open_fullscreen_writer(self):
        """打开全屏写作模式 - 与自动写作共享上下文"""
        # 获取当前章节内容
        current_text = self.content_text.get("1.0", tk.END).strip()
        
        # 构建共享上下文（世界观、角色、大纲等）
        shared_context = ""
        if self.memory:
            settings = self.memory.get_settings()
            if settings:
                shared_context += f"世界观: {json.dumps(settings, ensure_ascii=False)[:500]}\n"
            characters = self.memory.get_characters()
            if characters:
                shared_context += f"角色: {json.dumps(characters, ensure_ascii=False)[:500]}\n"
            global_summary = self.memory.get_global_summary()
            if global_summary:
                shared_context += f"故事进展: {global_summary[:300]}\n"
        
        # 如果有大纲，添加当前章节大纲
        if self.outline and self.current_chapter > 0 and self.current_chapter <= len(self.outline):
            ch_info = self.outline[self.current_chapter - 1]
            shared_context += f"当前章节大纲: {ch_info.get('summary', '')}\n"
        
        def save_callback(content):
            # 保存到当前章节
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", content)
            self.word_count_var.set(str(len(content)))
            self._save_chapter()
            
            # 自动更新记忆
            if self.memory and self.agent:
                with self._state_lock:
                    ch_num = self.current_chapter
                threading.Thread(target=lambda: self.agent.finalize_chapter(
                    ch_num, content), daemon=True).start()
                self._log(f"[联动] 全屏写作内容已保存并更新记忆")
        
        FullscreenWriter(
            parent=self.root,
            ai_client=self.ai_client,
            config=self.config,
            initial_text=current_text,
            save_callback=save_callback,
            shared_context=shared_context  # 传递共享上下文
        )
    
    # ===== 笔记功能 =====
    
    # ===== 角色系统 =====
    
    def _init_character_system(self):
        """初始化角色系统"""
        if self.current_novel_dir:
            self.character_system = CharacterSystem(self.current_novel_dir)
            self.character_system.load()
            self._update_char_display()
            self.format_converter = FormatConverter(self.current_novel_dir)
            self.image_manager = ImageManager(self.current_novel_dir)
    
    def _update_char_display(self):
        """更新角色面板显示 - 同时更新左侧角色卡片"""
        C = UIStyle.COLORS
        
        # 更新左侧角色卡片
        for w in self.char_cards_container.winfo_children():
            w.destroy()
        
        if not self.character_system:
            tk.Label(self.char_cards_container, text="未创建角色", font=('微软雅黑', 9),
                    bg=C['bg_medium'], fg=C['text_muted']).pack(anchor=tk.W)
            self.char_select_combo['values'] = []
            return
        
        names = self.character_system.get_character_names()
        self.char_select_combo['values'] = names
        
        # 生成角色卡片
        for name in names[:5]:  # 最多显示5个角色
            self._create_char_card(name)
        
        # 更新日志页的角色详情
        if hasattr(self, 'char_detail_frame'):
            for w in self.char_detail_frame.winfo_children():
                w.destroy()
            
            if not self.character_system.character:
                tk.Label(self.char_detail_frame, text="请选择角色", font=('微软雅黑', 9),
                        bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W, pady=2)
                return
            
            char = self.character_system.character
            self._display_char_details(char)
    
    def _create_char_card(self, name):
        """创建单个角色卡片"""
        C = UIStyle.COLORS
        card = tk.Frame(self.char_cards_container, bg=C['bg_card'], padx=8, pady=6)
        card.pack(fill=tk.X, pady=2)
        
        # 头像 (首字母)
        avatar_colors = ['#534ab7', '#0f6e56', '#3b82f6', '#ef4444', '#f59e0b']
        color_idx = hash(name) % len(avatar_colors)
        
        avatar = tk.Label(card, text=name[0], font=('微软雅黑', 10, 'bold'),
                         bg=avatar_colors[color_idx], fg='white', width=2, height=1)
        avatar.pack(side=tk.LEFT, padx=(0, 8))
        
        # 角色信息
        info_frame = tk.Frame(card, bg=C['bg_card'])
        info_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        tk.Label(info_frame, text=name, font=('微软雅黑', 10, 'bold'),
                bg=C['bg_card'], fg=C['text_primary']).pack(anchor=tk.W)
        
        # 获取角色详情
        char = self.character_system.get_character(name)
        if char:
            level = getattr(char, 'level', 1)
            title = getattr(char, 'title', '无称号')
            tk.Label(info_frame, text=f"Lv.{level} | {title}", font=('微软雅黑', 8),
                    bg=C['bg_card'], fg=C['text_muted']).pack(anchor=tk.W)
        
        # 点击事件
        def select_char(n=name):
            self.char_select_var.set(n)
            self._on_char_select()
        
        card.bind('<Button-1>', lambda e: select_char())
        avatar.bind('<Button-1>', lambda e: select_char())
    
    def _display_char_details(self, char):
        """显示角色详细信息"""
        C = UIStyle.COLORS
        
        # 基本信息
        tk.Label(self.char_detail_frame, text=f"「{char.name}」{char.title}", 
                font=('微软雅黑', 10, 'bold'), bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W, pady=2)
        tk.Label(self.char_detail_frame, text=f"等级: Lv.{char.level}  |  EXP: {char.exp}/{char.exp_to_next}", 
                font=('微软雅黑', 9), bg=C['bg_medium'], fg=C['text_primary']).pack(anchor=tk.W)
        
        # 属性
        tk.Label(self.char_detail_frame, text="─ 属性 ─", font=('微软雅黑', 9, 'bold'),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W, pady=(5, 2))
        attrs_frame = tk.Frame(self.char_detail_frame, bg=C['bg_medium'])
        attrs_frame.pack(fill=tk.X)
        for attr_name, attr_val in [("HP", f"{char.hp}/{char.max_hp}"), ("MP", f"{char.mp}/{char.max_mp}"),
                                    ("攻击", getattr(char, 'attack', '?')), ("防御", getattr(char, 'defense', '?')),
                                    ("速度", getattr(char, 'speed', '?')), ("智力", getattr(char, 'intelligence', '?'))]:
            tk.Label(attrs_frame, text=f"{attr_name}: {attr_val}", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_primary']).pack(side=tk.LEFT, padx=3)
        
        # 武器
        tk.Label(self.char_detail_frame, text="─ 武器 ─", font=('微软雅黑', 9, 'bold'),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W, pady=(5, 2))
        if char.weapon:
            w = char.weapon
            w_name = w.get('name', '无')
            w_quality = w.get('quality', '普通')
            tk.Label(self.char_detail_frame, text=f"⚔ {w_name} [{w_quality}]", font=('微软雅黑', 9),
                    bg=C['bg_medium'], fg=C['accent_light']).pack(anchor=tk.W)
        else:
            tk.Label(self.char_detail_frame, text="未装备武器", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W)
        
        # 技能
        tk.Label(self.char_detail_frame, text="─ 技能 ─", font=('微软雅黑', 9, 'bold'),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W, pady=(5, 2))
        if char.skills:
            for skill in char.skills[:5]:
                s_name = skill.get('name', '未知')
                s_lv = skill.get('level', 1)
                tk.Label(self.char_detail_frame, text=f"✦ {s_name} Lv.{s_lv}", font=('微软雅黑', 8),
                        bg=C['bg_medium'], fg=C['text_primary']).pack(anchor=tk.W)
        else:
            tk.Label(self.char_detail_frame, text="未学习技能", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W)
        
        # 性格/背景
        tk.Label(self.char_detail_frame, text="─ 性格/背景 ─", font=('微软雅黑', 9, 'bold'),
                bg=C['bg_medium'], fg=C['text_secondary']).pack(anchor=tk.W, pady=(5, 2))
        personality = getattr(char, 'personality', '')
        backstory = getattr(char, 'backstory', '')
        appearance = getattr(char, 'appearance', '')
        if personality:
            tk.Label(self.char_detail_frame, text=f"性格: {str(personality)[:100]}", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_primary'], wraplength=200).pack(anchor=tk.W)
        if backstory:
            tk.Label(self.char_detail_frame, text=f"背景: {str(backstory)[:100]}", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_primary'], wraplength=200).pack(anchor=tk.W)
        if appearance:
            tk.Label(self.char_detail_frame, text=f"外貌: {str(appearance)[:80]}", font=('微软雅黑', 8),
                    bg=C['bg_medium'], fg=C['text_primary'], wraplength=200).pack(anchor=tk.W)
        
        self.char_detail_frame.update_idletasks()
        # Update scroll region
        canvas = self.char_detail_frame.master
        canvas.configure(scrollregion=canvas.bbox("all"))
    
    def _gen_char_biography(self):
        """从角色面板按钮生成选中角色个人传"""
        name = self.char_select_var.get()
        if not name or name == "无角色":
            messagebox.showwarning("提示", "请先选择一个角色")
            return
        
        if not self._check_ready(silent=True):
            return
        
        # 直接生成传记（跳过角色选择对话框）
        self._generate_character_biography(name)
    
    def _on_char_select(self, event=None):
        """切换活跃角色"""
        name = self.char_select_var.get()
        if self.character_system and name:
            self.character_system.set_active(name)
            self._update_char_display()
            self._log(f"切换到角色: {name}")
    
    def _create_character_dialog(self):
        """创建角色对话框 - 完善版"""
        dialog = tk.Toplevel(self.root)
        dialog.title("创建角色")
        dialog.geometry("450x500")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        fields = {}
        field_list = [
            ("角色名称:", "name", ""),
            ("称号:", "title", ""),
            ("等级:", "level", "1"),
            ("性格特点:", "personality", ""),
            ("外貌描述:", "appearance", ""),
            ("背景故事:", "backstory", ""),
        ]
        
        for label, key, default in field_list:
            tk.Label(dialog, text=label, font=('微软雅黑', 10),
                    bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(8, 2))
            if key in ("backstory",):
                entry = tk.Text(dialog, width=40, height=4, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
                entry.pack(padx=20)
            else:
                entry = tk.Entry(dialog, width=40, font=('微软雅黑', 10), bg=C['bg_card'], fg=C['text_primary'])
                entry.insert(0, default)
                entry.pack(padx=20)
            fields[key] = entry
        
        def create():
            name = fields["name"].get().strip() if isinstance(fields["name"], tk.Entry) else fields["name"].get("1.0", tk.END).strip()
            if not name:
                messagebox.showwarning("提示", "请输入角色名称")
                return
            if not self.character_system:
                self.character_system = CharacterSystem(self.current_novel_dir)
            
            if name in self.character_system.get_character_names():
                messagebox.showwarning("提示", "角色名已存在")
                return
            
            def get_val(key):
                widget = fields[key]
                if isinstance(widget, tk.Text):
                    return widget.get("1.0", tk.END).strip()
                return widget.get().strip()
            
            self.character_system.create_character(
                name=name,
                title=get_val("title"),
                level=int(get_val("level") or "1"),
                backstory=get_val("backstory"),
                personality=get_val("personality"),
                appearance=get_val("appearance"),
            )
            self._update_char_display()
            self._log(f"角色「{name}」创建成功")
            dialog.destroy()
        
        tk.Button(dialog, text="创建", font=('微软雅黑', 11),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=20, pady=5,
                 command=create).pack(pady=15)
    
    def _ai_create_character(self):
        """AI自动创建角色"""
        if not self.ai_client.is_configured():
            messagebox.showwarning("提示", "请先配置AI")
            return
        
        # 获取小说上下文
        context = ""
        if self.memory:
            settings = self.memory.get_settings()
            if settings:
                context = json.dumps(settings, ensure_ascii=False)[:300]
        
        def run():
            try:
                if not self.character_system:
                    self.character_system = CharacterSystem(self.current_novel_dir)
                
                result = self.character_system.ai_create_character(self.ai_client, context)
                
                if result["success"]:
                    char = result["character"]
                    self.root.after(0, lambda: self._update_char_display())
                    self._log(f"AI创建角色「{char.name}」成功")
                    
                    # 如果有武器/技能建议，显示给用户
                    suggestions = f"角色「{char.name}」创建成功！\n\n"
                    if result.get("weapon_suggestion"):
                        suggestions += f"建议武器: {result['weapon_suggestion']}\n"
                    if result.get("skill_suggestions"):
                        suggestions += f"建议技能: {', '.join(result['skill_suggestions'])}\n"
                    self.root.after(0, lambda: messagebox.showinfo("AI创建成功", suggestions))
                else:
                    self.root.after(0, lambda: messagebox.showerror("创建失败", result.get("error", "未知错误")))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        
        self._log("AI正在创建角色...")
        threading.Thread(target=run, daemon=True).start()
    
    def _show_char_detail(self):
        """显示角色详情"""
        if not self.character_system or not self.character_system.character:
            messagebox.showinfo("提示", "请先创建角色")
            return
        
        char = self.character_system.character
        
        dialog = tk.Toplevel(self.root)
        dialog.title(f"角色详情 - {char.name}")
        dialog.geometry("500x600")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        # 使用Notebook组织信息
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 属性页
        attr_frame = tk.Frame(notebook, bg=C['bg_dark'])
        notebook.add(attr_frame, text=" 属性 ")
        
        attr_text = tk.Text(attr_frame, wrap=tk.WORD, font=('微软雅黑', 10),
                           bg=C['bg_card'], fg=C['text_primary'], relief=tk.FLAT, padx=15, pady=15)
        attr_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        attr_text.insert("1.0", char.get_summary())
        attr_text.config(state=tk.DISABLED)
        
        # 武器/技能页
        equip_frame = tk.Frame(notebook, bg=C['bg_dark'])
        notebook.add(equip_frame, text=" 装备/技能 ")
        
        equip_text = tk.Text(equip_frame, wrap=tk.WORD, font=('微软雅黑', 10),
                            bg=C['bg_card'], fg=C['text_primary'], relief=tk.FLAT, padx=15, pady=15)
        equip_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        equip_info = "═══ 装备 ═══\n"
        equip_info += f"武器: {char.weapon.get('name', '无') if char.weapon else '无'}\n"
        equip_info += f"防具: {char.armor.get('name', '无') if char.armor else '无'}\n"
        equip_info += f"饰品: {char.accessory.get('name', '无') if char.accessory else '无'}\n\n"
        equip_info += "═══ 技能 ═══\n"
        for skill in char.skills:
            equip_info += f"• {skill.get('name', '')} ({skill.get('type', '')}) - {skill.get('desc', '')}\n"
        if not char.skills:
            equip_info += "暂无技能\n"
        
        equip_text.insert("1.0", equip_info)
        equip_text.config(state=tk.DISABLED)
        
        # 统计页
        stats_frame = tk.Frame(notebook, bg=C['bg_dark'])
        notebook.add(stats_frame, text=" 统计 ")
        
        stats_text = tk.Text(stats_frame, wrap=tk.WORD, font=('微软雅黑', 10),
                            bg=C['bg_card'], fg=C['text_primary'], relief=tk.FLAT, padx=15, pady=15)
        stats_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        stats_text.insert("1.0", self.character_system.get_stats_display())
        stats_text.config(state=tk.DISABLED)
        
        # 操作按钮
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Button(btn_frame, text="重命名", font=('微软雅黑', 9),
                 bg=C['bg_light'], fg=C['text_primary'], relief=tk.FLAT, padx=8,
                 command=lambda: self._rename_character(dialog)).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="删除角色", font=('微软雅黑', 9),
                 bg=C['error'], fg='white', relief=tk.FLAT, padx=8,
                 command=lambda: self._delete_character(dialog)).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="休息恢复", font=('微软雅黑', 9),
                 bg=C['success'], fg='white', relief=tk.FLAT, padx=8,
                 command=lambda: self._rest_character()).pack(side=tk.RIGHT, padx=3)
        tk.Button(btn_frame, text="📖 故事线", font=('微软雅黑', 9),
                 bg=C['accent'], fg='white', relief=tk.FLAT, padx=8,
                 command=lambda: self._edit_character_story(char.name)).pack(side=tk.RIGHT, padx=3)
    
    def _rename_character(self, dialog):
        """重命名角色"""
        if not self.character_system or not self.character_system.character:
            return
        old_name = self.character_system.character.name
        new_name = tk.simpledialog.askstring("重命名", "输入新名称:", initialvalue=old_name)
        if new_name and new_name != old_name:
            if self.character_system.rename_character(old_name, new_name):
                self._update_char_display()
                self._log(f"角色已重命名: {old_name} → {new_name}")
                dialog.destroy()
            else:
                messagebox.showwarning("提示", "名称已存在或无效")
    
    def _delete_character(self, dialog):
        """删除角色"""
        if not self.character_system or not self.character_system.character:
            return
        name = self.character_system.character.name
        if messagebox.askyesno("确认", f"确定删除角色「{name}」？"):
            self.character_system.delete_character(name)
            self._update_char_display()
            self._log(f"已删除角色: {name}")
            dialog.destroy()
    
    def _rest_character(self):
        """角色休息恢复"""
        if self.character_system and self.character_system.character:
            self.character_system.character.rest()
            self.character_system.save_character()
            self._update_char_display()
            self._log(f"{self.character_system.character.name} 休息恢复，HP/MP已满")
    
    def _edit_character_story(self, char_name: str):
        """编辑角色故事线"""
        if not self.current_novel_dir:
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title(f"角色故事线 - {char_name}")
        dialog.geometry("600x500")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        # 加载角色故事
        stories_dir = self.current_novel_dir / "character_stories"
        stories_dir.mkdir(exist_ok=True)
        story_file = stories_dir / f"{char_name}.json"
        
        story_data = {"name": char_name, "story_arcs": [], "notes": ""}
        if story_file.exists():
            with open(story_file, 'r', encoding='utf-8') as f:
                story_data = json.load(f)
        
        tk.Label(dialog, text=f"📖 {char_name} 的故事线", font=('微软雅黑', 12, 'bold'),
                bg=C['bg_dark'], fg=C['accent_light']).pack(pady=(15, 10))
        
        # 故事线列表
        list_frame = tk.Frame(dialog, bg=C['bg_dark'])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        
        story_list = tk.Listbox(list_frame, bg=C['bg_card'], fg=C['text_primary'],
                               font=('微软雅黑', 10), selectbackground=C['accent'],
                               relief=tk.FLAT, height=8)
        story_list.pack(fill=tk.BOTH, expand=True)
        
        for arc in story_data.get("story_arcs", []):
            story_list.insert(tk.END, f"• {arc.get('title', '未命名')}")
        
        # 故事详情
        detail_frame = tk.Frame(dialog, bg=C['bg_dark'])
        detail_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        
        tk.Label(detail_frame, text="故事详情:", bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W)
        detail_text = tk.Text(detail_frame, wrap=tk.WORD, font=('微软雅黑', 10),
                             bg=C['bg_card'], fg=C['text_primary'], height=6)
        detail_text.pack(fill=tk.BOTH, expand=True)
        detail_text.insert("1.0", story_data.get("notes", ""))
        
        def on_story_select(event):
            selection = story_list.curselection()
            if selection and selection[0] < len(story_data.get("story_arcs", [])):
                arc = story_data["story_arcs"][selection[0]]
                detail_text.delete("1.0", tk.END)
                detail_text.insert("1.0", f"标题: {arc.get('title', '')}\n\n{arc.get('content', '')}")
        
        story_list.bind('<<ListboxSelect>>', on_story_select)
        
        # 操作按钮
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def add_arc():
            title = tk.simpledialog.askstring("添加故事线", "故事线标题:")
            if title:
                story_data.setdefault("story_arcs", []).append({"title": title, "content": ""})
                story_list.insert(tk.END, f"• {title}")
                self._log(f"为 {char_name} 添加故事线: {title}")
        
        def edit_arc():
            selection = story_list.curselection()
            if not selection:
                messagebox.showwarning("提示", "请先选择故事线")
                return
            idx = selection[0]
            if idx < len(story_data.get("story_arcs", [])):
                content = detail_text.get("1.0", tk.END).strip()
                story_data["story_arcs"][idx]["content"] = content
                self._log(f"更新 {char_name} 的故事线")
        
        def save_story():
            story_data["notes"] = detail_text.get("1.0", tk.END).strip()
            with open(story_file, 'w', encoding='utf-8') as f:
                json.dump(story_data, f, indent=2, ensure_ascii=False)
            self._log(f"已保存 {char_name} 的故事线")
            messagebox.showinfo("成功", "故事线已保存")
        
        def view_all_stories():
            """查看所有角色的故事线"""
            all_stories_window = tk.Toplevel(dialog)
            all_stories_window.title("所有角色故事线")
            all_stories_window.geometry("700x500")
            all_stories_window.configure(bg=C['bg_dark'])
            
            all_text = tk.Text(all_stories_window, wrap=tk.WORD, font=('微软雅黑', 10),
                              bg=C['bg_card'], fg=C['text_primary'])
            all_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # 加载所有角色故事
            for f in stories_dir.glob("*.json"):
                try:
                    with open(f, 'r', encoding='utf-8') as fp:
                        data = json.load(fp)
                    name = data.get("name", f.stem)
                    all_text.insert(tk.END, f"═══ {name} ═══\n")
                    for arc in data.get("story_arcs", []):
                        all_text.insert(tk.END, f"  📖 {arc.get('title', '')}\n")
                        if arc.get('content'):
                            all_text.insert(tk.END, f"     {arc['content'][:100]}...\n")
                    if data.get("notes"):
                        all_text.insert(tk.END, f"  备注: {data['notes'][:100]}...\n")
                    all_text.insert(tk.END, "\n")
                except Exception as e:
                    logger.debug(f"读取角色故事失败 {f.name}: {e}")
            
            all_text.config(state=tk.DISABLED)
        
        tk.Button(btn_frame, text="添加故事线", command=add_arc, bg=C['accent'], fg='white',
                 font=('微软雅黑', 9), padx=8).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="更新内容", command=edit_arc, bg=C['bg_light'], fg=C['text_primary'],
                 font=('微软雅黑', 9), padx=8).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="保存", command=save_story, bg=C['success'], fg='white',
                 font=('微软雅黑', 9), padx=8).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="查看所有角色故事", command=view_all_stories, bg=C['bg_light'], fg=C['text_primary'],
                 font=('微软雅黑', 9), padx=8).pack(side=tk.RIGHT, padx=3)
    
    def _equip_weapon(self):
        """装备武器"""
        if not self.character_system or not self.character_system.character:
            messagebox.showinfo("提示", "请先创建角色")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("选择武器")
        dialog.geometry("500x450")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="选择武器:", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(10, 5))
        
        cat_var = tk.StringVar()
        cats = self.character_system.get_weapon_categories()
        cat_combo = ttk.Combobox(dialog, textvariable=cat_var, values=cats, state="readonly", width=20)
        cat_combo.pack(pady=5)
        cat_combo.set(cats[0] if cats else "")
        
        weapon_listbox = tk.Listbox(dialog, bg=C['bg_card'], fg=C['text_primary'],
                                   font=('微软雅黑', 9), selectbackground=C['accent'], height=10)
        weapon_listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        def update_list(*args):
            weapon_listbox.delete(0, tk.END)
            weapons = self.character_system.get_weapons(cat_var.get())
            for w in weapons:
                q = w.get('quality', '')
                custom = " [自定义]" if w.get('custom') else ""
                attrs = ", ".join(f"{k}:{v}" for k, v in w.get('attributes', {}).items())
                weapon_listbox.insert(tk.END, f"[{q}]{custom} {w.get('name', '')} - {attrs}")
        
        cat_combo.bind('<<ComboboxSelected>>', update_list)
        update_list()
        
        def equip():
            sel = weapon_listbox.curselection()
            if sel:
                weapons = self.character_system.get_weapons(cat_var.get())
                if sel[0] < len(weapons):
                    weapon = weapons[sel[0]]
                    self.character_system.character.equip_weapon(weapon)
                    self.character_system.save_character()
                    self._update_char_display()
                    self._log(f"装备武器: {weapon.get('name', '')}")
                    dialog.destroy()
        
        def add_custom():
            """添加自定义武器"""
            sub = tk.Toplevel(dialog)
            sub.title("自定义武器")
            sub.geometry("350x300")
            sub.configure(bg=C['bg_dark'])
            
            fields = {}
            for label, default in [("名称:", ""), ("品质:", "凡品"), ("描述:", ""), ("力量加成:", "10"), 
                                   ("敏捷加成:", "0"), ("体质加成:", "0"), ("智力加成:", "0")]:
                tk.Label(sub, text=label, bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(5,0))
                e = tk.Entry(sub, width=30)
                e.insert(0, default)
                e.pack(padx=20)
                fields[label] = e
            
            def save_custom():
                name = fields["名称:"].get().strip()
                if not name:
                    return
                attrs = {}
                for attr_name, field_key in [("力量", "力量加成:"), ("敏捷", "敏捷加成:"), 
                                              ("体质", "体质加成:"), ("智力", "智力加成:")]:
                    try:
                        val = int(fields[field_key].get())
                        if val > 0:
                            attrs[attr_name] = val
                    except (ValueError, KeyError, tk.TclError):
                        pass
                
                self.character_system.add_custom_weapon(
                    name=name, category=cat_var.get(), quality=fields["品质:"].get(),
                    desc=fields["描述:"].get(), attributes=attrs
                )
                update_list()
                self._log(f"添加自定义武器: {name}")
                sub.destroy()
            
            tk.Button(sub, text="保存", bg=C['accent'], fg='white', command=save_custom).pack(pady=10)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(btn_frame, text="装备", font=('微软雅黑', 10),
                 bg=C['accent'], fg='white', relief=tk.FLAT, command=equip).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="+ 自定义武器", font=('微软雅黑', 9),
                 bg=C['success'], fg='white', relief=tk.FLAT, command=add_custom).pack(side=tk.RIGHT, padx=5)
    
    def _learn_skill(self):
        """学习技能"""
        if not self.character_system or not self.character_system.character:
            messagebox.showinfo("提示", "请先创建角色")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("学习技能")
        dialog.geometry("500x450")
        dialog.configure(bg=UIStyle.COLORS['bg_dark'])
        C = UIStyle.COLORS
        
        tk.Label(dialog, text="选择技能:", font=('微软雅黑', 10, 'bold'),
                bg=C['bg_dark'], fg=C['text_primary']).pack(pady=(10, 5))
        
        cat_var = tk.StringVar()
        cats = self.character_system.get_skill_categories()
        cat_combo = ttk.Combobox(dialog, textvariable=cat_var, values=cats, state="readonly", width=20)
        cat_combo.pack(pady=5)
        cat_combo.set(cats[0] if cats else "")
        
        skill_listbox = tk.Listbox(dialog, bg=C['bg_card'], fg=C['text_primary'],
                                  font=('微软雅黑', 9), selectbackground=C['accent'], height=10)
        skill_listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        def update_list(*args):
            skill_listbox.delete(0, tk.END)
            skills = self.character_system.get_skills(cat_var.get())
            for s in skills:
                custom = " [自定义]" if s.get('custom') else ""
                skill_listbox.insert(tk.END, f"{s.get('name', '')}{custom} - {s.get('desc', '')} (MP:{s.get('mp_cost', 0)})")
        
        cat_combo.bind('<<ComboboxSelected>>', update_list)
        update_list()
        
        def learn():
            sel = skill_listbox.curselection()
            if sel:
                skills = self.character_system.get_skills(cat_var.get())
                if sel[0] < len(skills):
                    skill = skills[sel[0]]
                    if self.character_system.character.learn_skill(skill):
                        self.character_system.save_character()
                        self._update_char_display()
                        self._log(f"学会技能: {skill.get('name', '')}")
                    else:
                        messagebox.showinfo("提示", "已学会该技能")
                    dialog.destroy()
        
        def add_custom():
            sub = tk.Toplevel(dialog)
            sub.title("自定义技能")
            sub.geometry("350x250")
            sub.configure(bg=C['bg_dark'])
            
            fields = {}
            for label, default in [("名称:", ""), ("类型:", cat_var.get() or "攻击"), ("描述:", ""), ("MP消耗:", "10")]:
                tk.Label(sub, text=label, bg=C['bg_dark'], fg=C['text_primary']).pack(anchor=tk.W, padx=20, pady=(5,0))
                e = tk.Entry(sub, width=30)
                e.insert(0, default)
                e.pack(padx=20)
                fields[label] = e
            
            def save_custom():
                name = fields["名称:"].get().strip()
                if not name:
                    return
                mp = int(fields["MP消耗:"].get() or 0)
                self.character_system.add_custom_skill(
                    name=name, skill_type=fields["类型:"].get(),
                    desc=fields["描述:"].get(), mp_cost=mp
                )
                update_list()
                self._log(f"添加自定义技能: {name}")
                sub.destroy()
            
            tk.Button(sub, text="保存", bg=C['accent'], fg='white', command=save_custom).pack(pady=10)
        
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(btn_frame, text="学习", font=('微软雅黑', 10),
                 bg=C['accent'], fg='white', relief=tk.FLAT, command=learn).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="+ 自定义技能", font=('微软雅黑', 9),
                 bg=C['success'], fg='white', relief=tk.FLAT, command=add_custom).pack(side=tk.RIGHT, padx=5)
    
    def _refresh_notes(self):
        """刷新笔记列表"""
        self.notes_list.delete(0, tk.END)
        note_type = self.note_type_var.get()
        
        if note_type == "project":
            notes = self.note_manager.get_project_notes()
            for n in notes:
                self.notes_list.insert(tk.END, f"{n.get('title', '无标题')}")
        elif note_type == "doc":
            notes = self.note_manager.get_doc_notes(self.current_chapter)
            for n in notes:
                self.notes_list.insert(tk.END, f"[位置{n.get('position',0)}] {n.get('content', '')[:30]}")
        elif note_type == "sticky":
            notes = self.note_manager.get_sticky_notes()
            for n in notes:
                self.notes_list.insert(tk.END, f"{n.get('content', '')[:40]}")
    
    def _on_note_select(self, event):
        """笔记选中"""
        selection = self.notes_list.curselection()
        if not selection:
            return
        idx = selection[0]
        note_type = self.note_type_var.get()
        
        notes = []
        if note_type == "project":
            notes = self.note_manager.get_project_notes()
        elif note_type == "doc":
            notes = self.note_manager.get_doc_notes(self.current_chapter)
        elif note_type == "sticky":
            notes = self.note_manager.get_sticky_notes()
        
        if idx < len(notes):
            self.note_content.delete("1.0", tk.END)
            self.note_content.insert("1.0", notes[idx].get("content", ""))
    
    def _add_note(self):
        """新建笔记"""
        note_type = self.note_type_var.get()
        content = "新笔记内容..."
        
        if note_type == "project":
            self.note_manager.add_project_note("新笔记", content)
        elif note_type == "doc":
            self.note_manager.add_doc_note(self.current_chapter, content)
        elif note_type == "sticky":
            self.note_manager.add_sticky_note(content)
        
        self._refresh_notes()
    
    def _save_note(self):
        """保存当前笔记"""
        selection = self.notes_list.curselection()
        if not selection:
            messagebox.showinfo("提示", "请先选择一个笔记")
            return
        
        idx = selection[0]
        note_type = self.note_type_var.get()
        content = self.note_content.get("1.0", tk.END).strip()
        
        if note_type == "project":
            notes = self.note_manager.get_project_notes()
            if idx < len(notes):
                self.note_manager.update_project_note(notes[idx]["id"], content=content)
        elif note_type == "doc":
            notes = self.note_manager.get_doc_notes(self.current_chapter)
            if idx < len(notes):
                notes[idx]["content"] = content
                self.note_manager.save_doc_notes(self.current_chapter, notes)
        elif note_type == "sticky":
            notes = self.note_manager.get_sticky_notes()
            if idx < len(notes):
                notes[idx]["content"] = content
                self.note_manager.save_sticky_notes(notes)
        
        self._log("笔记已保存")
    
    def _delete_note(self):
        """删除笔记"""
        selection = self.notes_list.curselection()
        if not selection:
            return
        
        if not messagebox.askyesno("确认", "确定删除此笔记？"):
            return
        
        idx = selection[0]
        note_type = self.note_type_var.get()
        
        if note_type == "project":
            notes = self.note_manager.get_project_notes()
            if idx < len(notes):
                self.note_manager.delete_project_note(notes[idx]["id"])
        elif note_type == "doc":
            notes = self.note_manager.get_doc_notes(self.current_chapter)
            if idx < len(notes):
                self.note_manager.delete_doc_note(self.current_chapter, notes[idx]["id"])
        elif note_type == "sticky":
            notes = self.note_manager.get_sticky_notes()
            if idx < len(notes):
                self.note_manager.delete_sticky_note(notes[idx]["id"])
        
        self.note_content.delete("1.0", tk.END)
        self._refresh_notes()
    
    def _send_sticky_to_project(self):
        """将便笺发送到工程笔记"""
        selection = self.notes_list.curselection()
        if not selection:
            messagebox.showinfo("提示", "请先选择一个便笺")
            return
        
        idx = selection[0]
        if self.note_type_var.get() != "sticky":
            messagebox.showinfo("提示", "请先切换到便笺本")
            return
        
        notes = self.note_manager.get_sticky_notes()
        if idx < len(notes):
            self.note_manager.send_sticky_to_project(notes[idx]["id"])
            self._log("便笺已发送到工程笔记")
            messagebox.showinfo("成功", "便笺已发送到工程笔记")
    
    # ===== 创作工具集 =====
    
    def _refresh_toolkit(self):
        """刷新工具集界面"""
        for w in self.tool_content_frame.winfo_children():
            w.destroy()
        
        tool_type = self.tool_type_var.get()
        
        if tool_type == "elements":
            self._build_elements_tool()
        elif tool_type == "bridges":
            self._build_bridges_tool()
        elif tool_type == "descriptions":
            self._build_descriptions_tool()
        elif tool_type == "dialogue":
            self._build_dialogue_tool()
        elif tool_type == "story_flow":
            self._build_story_flow_tool()
        elif tool_type == "style":
            self._build_style_tool()
        elif tool_type == "adapt":
            self._build_adapt_tool()
        elif tool_type == "websearch":
            self._build_websearch_tool()
        elif tool_type == "chapters":
            self._build_chapter_analysis_tool()
        elif tool_type == "memory_viz":
            self._build_memory_viz_tool()
        elif tool_type == "summary_mgmt":
            self._build_summary_mgmt_tool()
        elif tool_type == "batch_ops":
            self._build_batch_ops_tool()
    
    
    def _show_tool_result(self, widget, text):
        widget.delete("1.0", tk.END)
        widget.insert("1.0", text)
    
    def _insert_to_chapter(self, text_widget):
        """将工具结果智能插入到当前章节（AI润色衔接）"""
        content = text_widget.get("1.0", tk.END).strip()
        if not content:
            return
        
        # 获取当前章节内容作为上下文
        current_text = self.content_text.get("1.0", tk.END).strip()
        
        if not self.ai_client.is_configured():
            # 无AI配置时直接插入
            self.content_text.insert(tk.INSERT, "\n\n" + content)
            self._log("内容已直接插入（未配置AI润色）")
            return
        
        # 使用AI润色使内容衔接更自然
        def run():
            try:
                self._log("正在AI润色插入内容...")
                
                # 获取当前章节上下文（取最后500字作为衔接参考）
                context = current_text[-500:] if len(current_text) > 500 else current_text
                
                system = """你是专业小说编辑。任务是将新内容自然地融入现有文章中。
要求：
1. 保持原文风格和语气
2. 添加自然的过渡语句
3. 确保上下文逻辑连贯
4. 不要重复已有内容
5. 直接输出修改后的内容，不要添加解释"""
                
                prompt = f"""现有文章（末尾部分）：
---
{context}
---

需要插入的新内容：
---
{content}
---

请将新内容自然地融入到现有文章末尾，确保衔接流畅。"""
                
                result = self.ai.chat([{"role": "user", "content": prompt}], system=system, max_tokens=2000)
                
                # 在主线程中插入润色后的内容
                self.root.after(0, lambda: self._insert_polished_content(result))
                self._log("AI润色完成，内容已插入")
                
            except Exception as e:
                # 润色失败时直接插入原内容
                self.root.after(0, lambda: self.content_text.insert(tk.INSERT, "\n\n" + content))
                self._log(f"AI润色失败，直接插入原内容: {e}")
        
        threading.Thread(target=run, daemon=True).start()
    
    def _insert_polished_content(self, content):
        """插入润色后的内容"""
        self.content_text.insert(tk.INSERT, "\n\n" + content)
        # 更新字数统计
        full_text = self.content_text.get("1.0", tk.END).strip()
        self.word_count_var.set(f"字数: {len(full_text)}")
    
    def _display_review(self, review_json: str):
        """显示审校结果（主线程调用）"""
        self.review_text.delete("1.0", tk.END)
        self.review_text.insert("1.0", review_json)
        self.notebook.select(2)
    
    def _check_ready(self) -> bool:
        """检查是否就绪"""
        if not self.ai_client.is_configured():
            messagebox.showwarning("提示", "请先配置AI API（设置 → AI配置）")
            return False
        if not self.current_novel_dir:
            messagebox.showwarning("提示", "请先创建或打开小说")
            return False
        return True
    
    def _get_meta(self) -> dict:
        """获取小说元数据"""
        with open(self.current_novel_dir / "meta.json", 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def _show_help(self):
        """显示帮助"""
        help_text = """AI自动写小说系统 v2.0 使用说明

═══ 1. 配置AI模型 ═══
菜单 → 设置 → AI模型

支持的AI后端：
- Ollama（本地免费）: 默认 http://localhost:11434
- OpenAI: 需要API密钥
- DeepSeek: 需要API密钥
- Claude: 需要API密钥
- 自定义API: 兼容OpenAI格式

点击"检测Ollama"可自动发现本地模型。

═══ 2. 配置文生图（可选）═══
菜单 → 设置 → 文生图

支持的后端：
- ComfyUI: 默认端口 8188
- SD WebUI API: 默认端口 7860
- Disabled: 不使用文生图

勾选"自动检测名场面"后，每章生成完毕会自动
检测适合插图的场景并提醒生成图片。

═══ 3. 新建小说 ═══
点击"新建小说"，填写：
- 标题、类型、核心概念
- 目标章节数

═══ 4. 创作流程 ═══
分步模式：
  1. 生成世界观 → 创建世界设定
  2. 生成角色 → 创建角色档案
  3. 生成大纲 → 规划故事结构
  4. 生成下一章 → 逐章创作
  5. 审校当前章 → 检查质量

自动模式：
  点击"自动创作"一键完成全流程

═══ 5. 长上下文记忆 ═══
系统自动维护：
- 世界观设定
- 角色档案
- 全局故事摘要
- 最近5章摘要
- 关键词索引

确保长篇小说的连贯性。

═══ 6. 导出 ═══
点击"导出全文"生成TXT文件
"""
        dialog = tk.Toplevel(self.root)
        dialog.title("使用说明")
        dialog.geometry("500x600")
        text = scrolledtext.ScrolledText(dialog, wrap=tk.WORD, font=("微软雅黑", 11))
        text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text.insert("1.0", help_text)
        text.config(state=tk.DISABLED)
    
    def _show_about(self):
        """显示关于"""
        messagebox.showinfo("关于", "AI自动写小说系统 v2.0\n\n功能：\n- AI API（Ollama/OpenAI/DeepSeek/Claude）\n- 长上下文记忆\n- 智能体自动创作\n- 文生图（ComfyUI/SD WebUI）\n- 名场面自动检测")
    
    def _on_close(self):
        """关闭应用"""
        if messagebox.askyesno("确认", "确定要退出吗？"):
            self.root.destroy()
    
    def run(self):
        """运行应用"""
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.mainloop()


# ==================== 入口 ====================

if __name__ == "__main__":
    app = NovelWriterApp()
    app.run()
