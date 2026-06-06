"""
阅读管理模块 - 支持多种格式的小说阅读
"""

import json
import time
import shutil
import re
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
from loguru import logger

from .config import AppConfig

# 文件格式支持（条件导入）
try:
    from ebooklib import epub
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


class ReadingManager:
    """阅读管理器 - 支持多种格式的小说阅读"""
    
    SUPPORTED_FORMATS = {
        '.txt': 'TXT文本文件',
        '.epub': 'EPUB电子书',
        '.pdf': 'PDF文档',
        '.docx': 'Word文档',
        '.md': 'Markdown文件',
    }
    
    def __init__(self, config: AppConfig):
        self.config = config
        self.books_dir = config.config_dir / "books"
        self.books_dir.mkdir(exist_ok=True)
        self.bookmarks_file = config.config_dir / "bookmarks.json"
        self.reading_history_file = config.config_dir / "reading_history.json"
        
        # 阅读设置
        self.font_size = 16
        self.font_family = "微软雅黑"
        self.line_spacing = 1.5
        self.bg_color = "#f5f0e8"  # 米白色背景
        self.text_color = "#2c2c2c"
        self.theme = "light"  # light/dark/sepia
        
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return list(self.SUPPORTED_FORMATS.keys())
    
    def import_book(self, file_path: str) -> Optional[Dict]:
        """导入书籍到书库"""
        path = Path(file_path)
        if not path.exists():
            return None
        
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            return None
        
        # 复制到书库
        dest = self.books_dir / path.name
        shutil.copy2(path, dest)
        
        # 提取元数据
        meta = self._extract_metadata(dest, ext)
        meta['file_path'] = str(dest)
        meta['import_date'] = datetime.now().isoformat()
        meta['last_read'] = None
        meta['progress'] = 0
        
        return meta
    
    def _extract_metadata(self, file_path: Path, ext: str) -> Dict:
        """提取书籍元数据"""
        meta = {
            'title': file_path.stem,
            'author': '未知',
            'format': ext,
            'size': file_path.stat().st_size,
            'pages': 0,
            'chapters': [],
        }
        
        try:
            if ext == '.txt':
                # 尝试读取前几行获取标题
                with open(file_path, 'r', encoding='utf-8') as f:
                    lines = [f.readline().strip() for _ in range(10)]
                    # 寻找标题行
                    for line in lines:
                        if line and len(line) < 50 and not line.startswith(('第', '章', '卷')):
                            meta['title'] = line
                            break
            
            elif ext == '.epub' and EPUB_SUPPORT:
                book = epub.read_epub(str(file_path))
                meta['title'] = book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else file_path.stem
                meta['author'] = book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else '未知'
                # 获取章节列表
                for item in book.get_items():
                    if item.get_type() == 9:  # ITEM_DOCUMENT
                        meta['chapters'].append(item.get_name())
            
            elif ext == '.pdf' and PDF_SUPPORT:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    meta['pages'] = len(reader.pages)
                    if reader.metadata:
                        meta['title'] = reader.metadata.title or file_path.stem
                        meta['author'] = reader.metadata.author or '未知'
            
            elif ext == '.docx' and DOCX_SUPPORT:
                doc = Document(str(file_path))
                # 获取段落数作为页数估计
                meta['pages'] = len(doc.paragraphs) // 50  # 估计每页50段落
        
        except Exception as e:
            logger.error(f"提取元数据失败: {e}")
        
        return meta
    
    def read_book(self, file_path: str, page: int = 0) -> Optional[str]:
        """读取书籍内容"""
        path = Path(file_path)
        if not path.exists():
            return None
        
        ext = path.suffix.lower()
        
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
            
            elif ext == '.epub' and EPUB_SUPPORT:
                book = epub.read_epub(str(file_path))
                content = []
                for item in book.get_items():
                    if item.get_type() == 9:  # ITEM_DOCUMENT
                        html = item.get_content().decode('utf-8', errors='ignore')
                        text = re.sub(r'<[^>]+>', '', html)
                        text = re.sub(r'\s+', ' ', text).strip()
                        if text:
                            content.append(text)
                return '\n\n'.join(content)
            
            elif ext == '.pdf' and PDF_SUPPORT:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = []
                    for i, page in enumerate(reader.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                content.append(text)
                        except Exception:
                            pass
                    return '\n\n'.join(content)
            
            elif ext == '.docx' and DOCX_SUPPORT:
                doc = Document(str(file_path))
                content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)
                return '\n\n'.join(content)
            
            elif ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        
        except Exception as e:
            logger.error(f"读取书籍失败: {e}")
            return None
        
        return None
    
    def get_bookmarks(self) -> List[Dict]:
        """获取所有书签"""
        if self.bookmarks_file.exists():
            with open(self.bookmarks_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def add_bookmark(self, file_path: str, position: int, title: str = ""):
        """添加书签"""
        bookmarks = self.get_bookmarks()
        bookmark = {
            'id': int(time.time() * 1000),
            'file_path': file_path,
            'position': position,
            'title': title or f"书签 {len(bookmarks) + 1}",
            'created_at': datetime.now().isoformat(),
        }
        bookmarks.append(bookmark)
        with open(self.bookmarks_file, 'w', encoding='utf-8') as f:
            json.dump(bookmarks, f, indent=2, ensure_ascii=False)
        return bookmark
    
    def delete_bookmark(self, bookmark_id: int):
        """删除书签"""
        bookmarks = self.get_bookmarks()
        bookmarks = [b for b in bookmarks if b.get('id') != bookmark_id]
        with open(self.bookmarks_file, 'w', encoding='utf-8') as f:
            json.dump(bookmarks, f, indent=2, ensure_ascii=False)
    
    def get_reading_history(self) -> List[Dict]:
        """获取阅读历史"""
        if self.reading_history_file.exists():
            with open(self.reading_history_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def update_reading_progress(self, file_path: str, position: int, progress: float):
        """更新阅读进度"""
        history = self.get_reading_history()
        
        # 查找现有记录
        found = False
        for record in history:
            if record.get('file_path') == file_path:
                record['position'] = position
                record['progress'] = progress
                record['last_read'] = datetime.now().isoformat()
                found = True
                break
        
        # 添加新记录
        if not found:
            history.append({
                'file_path': file_path,
                'position': position,
                'progress': progress,
                'last_read': datetime.now().isoformat(),
            })
        
        # 保存
        with open(self.reading_history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, indent=2, ensure_ascii=False)
    
    def get_library_books(self) -> List[Dict]:
        """获取书库中的所有书籍"""
        books = []
        for file in self.books_dir.iterdir():
            if file.is_file() and file.suffix.lower() in self.SUPPORTED_FORMATS:
                meta = self._extract_metadata(file, file.suffix.lower())
                meta['file_path'] = str(file)
                books.append(meta)
        return books
    
    def search_in_book(self, file_path: str, keyword: str) -> List[Dict]:
        """在书籍中搜索关键词"""
        content = self.read_book(file_path)
        if not content:
            return []
        
        results = []
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if keyword.lower() in line.lower():
                results.append({
                    'line_number': i + 1,
                    'content': line.strip(),
                    'context': lines[max(0, i-1):i+2],  # 上下文
                })
        
        return results
    
    def export_bookmarks(self, file_path: str):
        """导出书签到文件"""
        bookmarks = self.get_bookmarks()
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(bookmarks, f, indent=2, ensure_ascii=False)
    
    def import_bookmarks(self, file_path: str):
        """从文件导入书签"""
        with open(file_path, 'r', encoding='utf-8') as f:
            bookmarks = json.load(f)
        
        existing = self.get_bookmarks()
        existing_ids = {b['id'] for b in existing}
        
        # 合并书签，避免重复
        for bookmark in bookmarks:
            if bookmark['id'] not in existing_ids:
                existing.append(bookmark)
        
        with open(self.bookmarks_file, 'w', encoding='utf-8') as f:
            json.dump(existing, f, indent=2, ensure_ascii=False)
